import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime, date, time, timedelta
import uuid
import random
from io import BytesIO
import base64

st.set_page_config(page_title="CLC Behaviour Support", page_icon="📊", layout="wide", initial_sidebar_state="collapsed")

# MINIMALIST PROFESSIONAL STYLING - LIGHT GRAY/WHITE
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    * { font-family: 'Inter', sans-serif; }
    
    .stApp { background: #f8fafc; }
    
    .stButton>button {
        background: #334155 !important; color: white !important;
        border: none !important; border-radius: 6px !important;
        padding: 0.5rem 1.2rem !important; font-weight: 600 !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        transition: all 0.2s !important;
    }
    .stButton>button:hover { background: #1e293b !important; transform: translateY(-1px) !important; }
    
    button[kind="primary"] {
        background: #0ea5e9 !important; color: white !important;
    }
    button[kind="primary"]:hover { background: #0284c7 !important; }
    
    [data-testid="stVerticalBlock"] > div[style*="border"] {
        background: white !important; border-radius: 8px !important;
        padding: 1.5rem !important; box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        border: 1px solid #e2e8f0 !important;
    }
    
    [data-testid="stMetricValue"] {
        font-size: 1.875rem !important; font-weight: 700 !important; color: #0f172a !important;
    }
    [data-testid="stMetricLabel"] {
        color: #64748b !important; font-weight: 600 !important; font-size: 0.875rem !important;
        text-transform: uppercase !important; letter-spacing: 0.05em !important;
    }
    
    .stTextInput>div>div>input, .stSelectbox>div>div>select, 
    .stTextArea>div>div>textarea, .stNumberInput>div>div>input,
    .stDateInput>div>div>input, .stTimeInput>div>div>input {
        border: 1px solid #cbd5e1 !important; background: white !important;
        color: #0f172a !important; font-weight: 500 !important; border-radius: 6px !important;
    }
    .stTextInput>div>div>input:focus, .stSelectbox>div>div>select:focus,
    .stTextArea>div>div>textarea:focus {
        border-color: #0ea5e9 !important; box-shadow: 0 0 0 3px rgba(14, 165, 233, 0.1) !important;
    }
    
    h1 { color: #0f172a !important; font-weight: 700 !important; }
    h2 { color: #0f172a !important; font-weight: 700 !important; }
    h3 { color: #0f172a !important; font-weight: 600 !important; }
    
    label { color: #334155 !important; font-weight: 600 !important; font-size: 0.875rem !important; }
    
    .stSuccess { background: #ecfdf5 !important; color: #065f46 !important; 
                 border-left: 4px solid #10b981 !important; }
    .stInfo { background: #f0f9ff !important; color: #075985 !important; 
              border-left: 4px solid #0ea5e9 !important; }
    .stWarning { background: #fffbeb !important; color: #92400e !important; 
                 border-left: 4px solid #f59e0b !important; }
    .stError { background: #fef2f2 !important; color: #991b1b !important;
               border-left: 4px solid #ef4444 !important; }
    
    .stMarkdown p, .stMarkdown li { color: #334155 !important; }
    
    .streamlit-expanderHeader {
        background: #f8fafc !important; color: #0f172a !important; 
        font-weight: 600 !important; border: 1px solid #e2e8f0 !important;
    }
    
    [data-testid="stHorizontalBlock"] { gap: 1rem !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div style='background: white; padding: 1.25rem; border-radius: 8px; margin-bottom: 1.5rem;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1); border-left: 4px solid #0ea5e9;'>
    <div style='color: #0f172a; font-weight: 700; font-size: 1.125rem; margin-bottom: 0.25rem;'>
        🎭 SANDBOX MODE
    </div>
    <div style='color: #64748b; font-size: 0.875rem; font-weight: 500;'>
        This demonstration uses synthetic data only. No real student information is included.
    </div>
</div>
""", unsafe_allow_html=True)

# MOCK DATA
MOCK_STAFF = [
    {"id": "s1", "name": "Emily Jones", "role": "JP", "email": "emily.jones@example.com", "password": "demo123"},
    {"id": "s2", "name": "Daniel Lee", "role": "PY", "email": "daniel.lee@example.com", "password": "demo123"},
    {"id": "s3", "name": "Sarah Chen", "role": "SY", "email": "sarah.chen@example.com", "password": "demo123"},
    {"id": "s4", "name": "Admin User", "role": "ADM", "email": "admin@example.com", "password": "admin123"},
]

MOCK_STUDENTS = [
    {"id": "stu_jp1", "name": "Emma T.", "grade": "R", "dob": "2018-05-30", "program": "JP"},
    {"id": "stu_jp2", "name": "Oliver S.", "grade": "Y1", "dob": "2017-09-12", "program": "JP"},
    {"id": "stu_jp3", "name": "Sophie M.", "grade": "Y2", "dob": "2016-03-20", "program": "JP"},
    {"id": "stu_py1", "name": "Liam C.", "grade": "Y3", "dob": "2015-06-15", "program": "PY"},
    {"id": "stu_py2", "name": "Ava R.", "grade": "Y4", "dob": "2014-11-08", "program": "PY"},
    {"id": "stu_py3", "name": "Noah B.", "grade": "Y6", "dob": "2012-02-28", "program": "PY"},
    {"id": "stu_sy1", "name": "Isabella G.", "grade": "Y7", "dob": "2011-04-17", "program": "SY"},
    {"id": "stu_sy2", "name": "Ethan D.", "grade": "Y9", "dob": "2009-12-03", "program": "SY"},
    {"id": "stu_sy3", "name": "Mia A.", "grade": "Y11", "dob": "2007-08-20", "program": "SY"},
]

PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}
BEHAVIOUR_TYPES = ["Verbal Refusal", "Elopement", "Property Destruction", "Aggression (Peer)", 
                   "Aggression (Adult)", "Self-Harm", "Verbal Aggression", "Other"]
ANTECEDENTS = ["Requested to transition", "Given instruction/demand", "Peer conflict", 
               "Staff attention shifted", "Unstructured time", "Sensory overload", 
               "Access denied", "Change in routine", "Difficult task"]
INTERVENTIONS = ["CPI Supportive stance", "Offered break", "Reduced demand", "Provided choices", 
                "Removed audience", "Visual supports", "Co-regulation", "Prompted coping skill", "Redirection"]
LOCATIONS = ["JP Classroom", "PY Classroom", "SY Classroom", "Playground", "Library", "Admin", "Gate", "Toilets"]
VALID_PAGES = ["login", "landing", "program_students", "incident_log", "critical_incident", "student_analysis"]

# PROFESSIONAL SEVERITY GUIDE - GRAYSCALE
def show_severity_guide():
    st.markdown("""
    <div style='background: white; padding: 1.25rem; border-radius: 8px; margin: 1rem 0; 
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1); border: 1px solid #e2e8f0;'>
        <div style='color: #0f172a; font-weight: 700; margin-bottom: 1rem; font-size: 1rem;'>
            📊 Severity Level Guide
        </div>
        <div style='display: grid; grid-template-columns: repeat(5, 1fr); gap: 0.75rem;'>
            <div style='background: #f8fafc; padding: 1rem; border-radius: 6px; border: 2px solid #cbd5e1;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>1 - Low</div>
                <div style='color: #64748b; font-size: 0.8rem; line-height: 1.3;'>Persistent minor behaviours</div>
            </div>
            <div style='background: #f1f5f9; padding: 1rem; border-radius: 6px; border: 2px solid #94a3b8;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>2 - Disruptive</div>
                <div style='color: #64748b; font-size: 0.8rem; line-height: 1.3;'>Impacts others</div>
            </div>
            <div style='background: #e2e8f0; padding: 1rem; border-radius: 6px; border: 2px solid #64748b;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>3 - Concerning</div>
                <div style='color: #475569; font-size: 0.8rem; line-height: 1.3;'>Verbal aggression</div>
            </div>
            <div style='background: #cbd5e1; padding: 1rem; border-radius: 6px; border: 2px solid #475569;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>4 - Serious</div>
                <div style='color: #334155; font-size: 0.8rem; line-height: 1.3;'>Physical aggression</div>
            </div>
            <div style='background: #94a3b8; padding: 1rem; border-radius: 6px; border: 2px solid #1e293b;'>
                <div style='color: #fff; font-weight: 700; margin-bottom: 0.5rem;'>5 - Critical</div>
                <div style='color: #f1f5f9; font-size: 0.8rem; line-height: 1.3;'>Severe violence</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def send_critical_incident_email(incident_data, student, staff_email):
    st.info(f"""📧 **Email Notification Sent**
    
**To:** manager@clc.sa.edu.au, {staff_email}  
**Subject:** CRITICAL INCIDENT - {student['name']}

**Student:** {student['name']} | **Programme:** {student['program']} | **Grade:** {student['grade']}  
**Primary Behaviour:** {incident_data.get('ABCH_primary', {}).get('B', 'N/A')}

*(In production, this sends via SMTP)*
    """)

def create_graph_base64(fig):
    """Convert plotly figure to base64 PNG for Word doc"""
    img_bytes = fig.to_image(format="png", width=600, height=400)
    return base64.b64encode(img_bytes).decode()

def generate_behaviour_analysis_plan_docx(student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level):
    """Generate Word doc WITH graphs as embedded images"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Behaviour Analysis Plan', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Student Info
        doc.add_heading('Student Information', 1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.rows[0].cells[0].text = 'Student:'
        info_table.rows[0].cells[1].text = student['name']
        info_table.rows[1].cells[0].text = 'Program:'
        info_table.rows[1].cells[1].text = student['program']
        info_table.rows[2].cells[0].text = 'Grade:'
        info_table.rows[2].cells[1].text = student['grade']
        info_table.rows[3].cells[0].text = 'Date:'
        info_table.rows[3].cells[1].text = datetime.now().strftime('%d/%m/%Y')
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        summary = doc.add_paragraph()
        summary.add_run('Total Incidents: ').bold = True
        summary.add_run(f"{len(full_df)}\n")
        summary.add_run('Critical Incidents: ').bold = True
        summary.add_run(f"{len(full_df[full_df['incident_type'] == 'Critical'])}\n")
        summary.add_run('Average Severity: ').bold = True
        summary.add_run(f"{full_df['severity'].mean():.2f}\n")
        summary.add_run('Risk Level: ').bold = True
        summary.add_run(f"{risk_level} ({risk_score}/100)")
        
        doc.add_paragraph()
        
        # Key Findings
        doc.add_heading('Key Findings', 1)
        findings = doc.add_paragraph()
        findings.add_run('Primary Behaviour: ').bold = True
        findings.add_run(f"{top_beh}\n\n")
        findings.add_run('Most Common Trigger: ').bold = True
        findings.add_run(f"{top_ant}\n\n")
        findings.add_run('Hotspot Location: ').bold = True
        findings.add_run(f"{top_loc} during {top_session}")
        
        doc.add_paragraph()
        
        # GRAPHS SECTION
        doc.add_heading('Visual Analytics', 1)
        
        try:
            # Graph 1: Daily Frequency
            doc.add_heading('Daily Incident Frequency', 2)
            daily = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
            fig1 = go.Figure()
            fig1.add_trace(go.Scatter(
                x=daily["date_parsed"], y=daily["count"],
                mode='lines+markers', line=dict(color='#334155', width=2),
                marker=dict(size=6), fill='tozeroy', fillcolor='rgba(51, 65, 85, 0.1)'
            ))
            fig1.update_layout(
                height=300, width=600, showlegend=False,
                plot_bgcolor='white', paper_bgcolor='white',
                margin=dict(l=40, r=40, t=40, b=40)
            )
            img_path1 = "/tmp/daily_freq.png"
            fig1.write_image(img_path1)
            doc.add_picture(img_path1, width=Inches(5.5))
            doc.add_paragraph("Shows the pattern of incidents over time.")
            doc.add_paragraph()
            
            # Graph 2: Top Behaviours
            doc.add_heading('Most Common Behaviours', 2)
            beh_counts = full_df["behaviour_type"].value_counts().head(5)
            fig2 = go.Figure()
            fig2.add_trace(go.Bar(
                y=beh_counts.index, x=beh_counts.values,
                orientation='h', marker=dict(color='#334155'),
                text=beh_counts.values, textposition='outside'
            ))
            fig2.update_layout(
                height=300, width=600, showlegend=False,
                plot_bgcolor='white', paper_bgcolor='white',
                margin=dict(l=40, r=40, t=40, b=40)
            )
            img_path2 = "/tmp/top_behaviours.png"
            fig2.write_image(img_path2)
            doc.add_picture(img_path2, width=Inches(5.5))
            doc.add_paragraph(f"Primary behaviour: {beh_counts.index[0]} ({beh_counts.values[0]} incidents)")
            doc.add_paragraph()
            
            # Graph 3: Top Triggers
            doc.add_heading('Most Common Triggers', 2)
            ant_counts = full_df["antecedent"].value_counts().head(5)
            fig3 = go.Figure()
            fig3.add_trace(go.Bar(
                y=ant_counts.index, x=ant_counts.values,
                orientation='h', marker=dict(color='#475569'),
                text=ant_counts.values, textposition='outside'
            ))
            fig3.update_layout(
                height=300, width=600, showlegend=False,
                plot_bgcolor='white', paper_bgcolor='white',
                margin=dict(l=40, r=40, t=40, b=40)
            )
            img_path3 = "/tmp/top_triggers.png"
            fig3.write_image(img_path3)
            doc.add_picture(img_path3, width=Inches(5.5))
            doc.add_paragraph(f"Key trigger: {ant_counts.index[0]}")
            
        except Exception as e:
            doc.add_paragraph(f"Note: Graphs could not be generated. Error: {str(e)}")
        
        doc.add_paragraph()
        
        # Clinical Interpretation
        doc.add_heading('Clinical Interpretation', 1)
        interp = doc.add_paragraph()
        interp.add_run(f"Data indicates {student['name']} is most vulnerable when '{top_ant}' occurs in {top_loc} during {top_session}. ")
        interp.add_run("This behaviour serves as a safety strategy. CPI principles emphasize Supportive stance. ")
        interp.add_run("Berry Street Model suggests strengthening Body (regulation) and Relationship (connection).")
        
        doc.add_paragraph()
        
        # Recommendations
        doc.add_heading('Recommendations', 1)
        doc.add_heading('1. Proactive Strategies', 2)
        doc.add_paragraph(f"Provide check-in before '{top_ant}'", style='List Bullet')
        doc.add_paragraph(f"Offer regulated start before {top_session}", style='List Bullet')
        
        doc.add_heading('2. Co-regulation (CPI)', 2)
        doc.add_paragraph("Use Supportive stance, low slow voice", style='List Bullet')
        doc.add_paragraph("Reduce audience, one key adult", style='List Bullet')
        
        doc.add_heading('3. Teaching Skills', 2)
        doc.add_paragraph("Link to Personal & Social Capability", style='List Bullet')
        doc.add_paragraph("Teach help-seeking routines", style='List Bullet')
        
        doc.add_heading('4. SMART Goal', 2)
        doc.add_paragraph("Over 5 weeks, use help-seeking strategy in 4/5 opportunities with support.", style='List Bullet')
        
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('\n\nGenerated by CLC Behaviour Support\n')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)
        footer.add_run(datetime.now().strftime('%d %B %Y'))
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
        
    except Exception as e:
        st.error(f"Error generating document: {e}")
        return None

def init_state():
    ss = st.session_state
    if "logged_in" not in ss: ss.logged_in = False
    if "current_user" not in ss: ss.current_user = None
    if "current_page" not in ss: ss.current_page = "login"
    if "students" not in ss: ss.students = MOCK_STUDENTS
    if "staff" not in ss: ss.staff = MOCK_STAFF
    if "incidents" not in ss: ss.incidents = generate_mock_incidents(70)
    if "critical_incidents" not in ss: ss.critical_incidents = []
    if "selected_program" not in ss: ss.selected_program = "JP"
    if "selected_student_id" not in ss: ss.selected_student_id = None
    if "current_incident_id" not in ss: ss.current_incident_id = None
    if "abch_rows" not in ss: ss.abch_rows = []

def login_user(email: str, password: str) -> bool:
    email = (email or "").strip().lower()
    password = (password or "").strip()
    if not email or not password: return False
    for staff in st.session_state.staff:
        if staff.get("email", "").lower() == email and staff.get("password", "") == password:
            st.session_state.logged_in = True
            st.session_state.current_user = staff
            st.session_state.current_page = "landing"
            return True
    return False

def go_to(page: str, **kwargs):
    if page not in VALID_PAGES: return
    st.session_state.current_page = page
    for k, v in kwargs.items():
        setattr(st.session_state, k, v)
    st.rerun()

def get_student(sid): return next((s for s in st.session_state.students if s["id"] == sid), None)
def get_session_from_time(t): return "Morning" if t.hour < 11 else "Middle" if t.hour < 13 else "Afternoon"

def generate_mock_incidents(n=70):
    incidents = []
    weights = {"stu_sy1": 12, "stu_py1": 10, "stu_sy2": 9, "stu_jp1": 8, "stu_py2": 7}
    pool = []
    for stu in MOCK_STUDENTS:
        pool.extend([stu] * weights.get(stu["id"], 5))
    for _ in range(n):
        stu = random.choice(pool)
        sev = random.choices([1, 2, 3, 4, 5], weights=[20, 35, 25, 15, 5])[0]
        dt = datetime.now() - timedelta(days=random.randint(0, 90))
        dt = dt.replace(hour=random.choices([9,10,11,12,13,14,15], weights=[10,15,12,8,12,18,10])[0], 
                       minute=random.randint(0,59), second=0)
        incidents.append({
            "id": str(uuid.uuid4()), "student_id": stu["id"], "student_name": stu["name"],
            "date": dt.date().isoformat(), "time": dt.time().strftime("%H:%M:%S"),
            "day": dt.strftime("%A"), "session": get_session_from_time(dt.time()),
            "location": random.choice(LOCATIONS), "behaviour_type": random.choice(BEHAVIOUR_TYPES),
            "antecedent": random.choice(ANTECEDENTS), "intervention": random.choice(INTERVENTIONS),
            "severity": sev, "reported_by": random.choice(MOCK_STAFF)["name"],
            "description": "Mock incident", "is_critical": sev >= 4, "duration_minutes": random.randint(2, 25)
        })
    return incidents


# PAGES
def render_login_page():
    st.markdown("## 🔐 Staff Login")
    
    with st.container(border=True):
        st.markdown("**Demo Credentials:**")
        st.code("Email: emily.jones@example.com\nPassword: demo123")
    
    email = st.text_input("Email Address", placeholder="your.email@example.com", key="login_email")
    password = st.text_input("Password", type="password", placeholder="Enter password", key="login_pass")
    
    if st.button("Login", type="primary", use_container_width=True):
        if login_user(email, password):
            st.success(f"Welcome {st.session_state.current_user['name']}!")
            st.rerun()
        else:
            st.error("Invalid credentials")

def render_landing_page():
    user = st.session_state.current_user or {}
    st.markdown(f"### 👋 Welcome, {user.get('name', 'User')}")
    
    if st.button("Logout", key="logout_btn"):
        st.session_state.logged_in = False
        st.session_state.current_page = "login"
        st.rerun()
    
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1: st.metric("Students", len(st.session_state.students))
    with col2: st.metric("Total Incidents", len(st.session_state.incidents))
    with col3: st.metric("Critical", len([i for i in st.session_state.incidents if i.get("is_critical")]))
    
    st.markdown("---")
    st.markdown("### Select Program")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Junior Primary**")
        if st.button("Enter JP", use_container_width=True, type="primary", key="btn_jp"):
            go_to("program_students", selected_program="JP")
    with col2:
        st.markdown("**Primary Years**")
        if st.button("Enter PY", use_container_width=True, type="primary", key="btn_py"):
            go_to("program_students", selected_program="PY")
    with col3:
        st.markdown("**Senior Years**")
        if st.button("Enter SY", use_container_width=True, type="primary", key="btn_sy"):
            go_to("program_students", selected_program="SY")

def render_program_students_page():
    program = st.session_state.get("selected_program", "JP")
    st.markdown(f"## {PROGRAM_NAMES.get(program)} — Students")
    if st.button("⬅ Back", key="back_students"):
        go_to("landing")
    
    students = [s for s in st.session_state.students if s["program"] == program]
    for stu in students:
        stu_incidents = [i for i in st.session_state.incidents if i["student_id"] == stu["id"]]
        
        with st.container(border=True):
            col1, col2, col3 = st.columns([3, 2, 2])
            with col1:
                st.markdown(f"**{stu['name']}**")
                st.caption(f"Grade {stu['grade']}")
            with col2:
                st.metric("Incidents", len(stu_incidents))
            with col3:
                if st.button("📝 Log", key=f"log_{stu['id']}", use_container_width=True):
                    go_to("incident_log", selected_student_id=stu["id"])
                if st.button("📊 Analysis", key=f"ana_{stu['id']}", use_container_width=True):
                    go_to("student_analysis", selected_student_id=stu["id"])

def render_incident_log_page():
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## 📝 Incident Log — {student['name']}")
    show_severity_guide()
    
    # EMPTY FORM - NO PREFILLS
    with st.form("incident_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            inc_date = st.date_input("Date *", date.today(), key="inc_date")
            inc_time = st.time_input("Time *", datetime.now().time(), key="inc_time")
            location = st.selectbox("Location *", [""] + LOCATIONS, key="inc_loc")
        with col2:
            behaviour = st.selectbox("Behaviour Type *", [""] + BEHAVIOUR_TYPES, key="inc_beh")
            antecedent = st.selectbox("Antecedent/Trigger *", [""] + ANTECEDENTS, key="inc_ant")
            intervention = st.selectbox("Intervention Used *", [""] + INTERVENTIONS, key="inc_int")
        
        duration = st.number_input("Duration (minutes) *", min_value=1, value=1, key="inc_dur")
        severity = st.slider("Severity Level (see guide above) *", 1, 5, 1, key="inc_sev")
        description = st.text_area("Brief Description *", placeholder="Factual, objective description of what occurred...", key="inc_desc")
        
        submitted = st.form_submit_button("Submit Incident", type="primary")
    
    if submitted:
        if not location or not behaviour or not antecedent or not intervention or not description:
            st.error("Please complete all required fields marked with *")
        else:
            new_id = str(uuid.uuid4())
            rec = {
                "id": new_id, "student_id": student_id, "student_name": student["name"],
                "date": inc_date.isoformat(), "time": inc_time.strftime("%H:%M:%S"),
                "day": inc_date.strftime("%A"), "session": get_session_from_time(inc_time),
                "location": location, "behaviour_type": behaviour, "antecedent": antecedent,
                "intervention": intervention, "severity": severity,
                "reported_by": st.session_state.current_user["name"],
                "duration_minutes": duration, "description": description, "is_critical": severity >= 4
            }
            st.session_state.incidents.append(rec)
            st.success("✅ Incident logged successfully")
            
            if severity >= 4:
                st.warning("⚠️ **Critical Incident Detected** (Severity 4 or 5)")
                st.info("A Critical Incident ABCH form must be completed for this incident.")
                st.session_state.current_incident_id = new_id
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Complete Critical Form Now", type="primary", key="crit_now", use_container_width=True):
                        go_to("critical_incident", current_incident_id=new_id)
                with col2:
                    if st.button("Complete Later", key="crit_later", use_container_width=True):
                        go_to("program_students", selected_program=student["program"])
            else:
                if st.button("↩️ Back to Students", key="back_after_log"):
                    go_to("program_students", selected_program=student["program"])

def render_critical_incident_page():
    """PROPER CHRONOLOGY STRUCTURE - Multiple incident rows"""
    inc_id = st.session_state.get("current_incident_id")
    quick_inc = next((i for i in st.session_state.incidents if i["id"] == inc_id), None)
    
    if not quick_inc:
        st.error("No incident found")
        return
    
    student = get_student(quick_inc["student_id"])
    st.markdown(f"## 🚨 Critical Incident ABCH Form")
    st.markdown(f"**Student:** {student['name']} | **Date:** {quick_inc['date']} | **Time:** {quick_inc['time']}")
    
    st.markdown("---")
    st.markdown("### Incident Chronology")
    st.caption("Document the sequence of events using ABC format. Add rows as needed for complex incidents.")
    
    # PRIMARY ABCH ROW
    st.markdown("#### Primary Incident")
    colA1, colB1, colC1 = st.columns(3)
    with colA1:
        st.markdown("**A — Antecedent**")
        st.caption("What happened immediately before?")
        A1_text = st.text_area("", placeholder="e.g., Given instruction to transition to next activity", 
                              key="A1", height=150, label_visibility="collapsed")
    with colB1:
        st.markdown("**B — Behaviour**")
        st.caption("What did the student do?")
        B1_text = st.text_area("", placeholder="e.g., Verbal refusal, left classroom", 
                              key="B1", height=150, label_visibility="collapsed")
    with colC1:
        st.markdown("**C — Consequence**")
        st.caption("What happened as a result?")
        C1_text = st.text_area("", placeholder="e.g., Staff followed, provided space", 
                              key="C1", height=150, label_visibility="collapsed")
    
    # Hypothesis for primary
    st.markdown("**H — Hypothesis (Function)**")
    st.caption("Why did this behaviour occur? What was the student trying to achieve?")
    H1_text = st.text_area("", placeholder="e.g., To avoid the transition demand / To gain control", 
                          key="H1", height=80, label_visibility="collapsed")
    
    st.markdown("---")
    
    # ADDITIONAL ROWS
    if "abch_rows" not in st.session_state:
        st.session_state.abch_rows = []
    
    if st.button("➕ Add Another Incident Row", key="add_row"):
        st.session_state.abch_rows.append({"A": "", "B": "", "C": "", "H": ""})
        st.rerun()
    
    for idx, row in enumerate(st.session_state.abch_rows):
        st.markdown(f"#### Incident {idx + 2}")
        colA, colB, colC = st.columns(3)
        with colA:
            st.markdown("**A — Antecedent**")
            row["A"] = st.text_area("", key=f"A{idx+2}", height=120, label_visibility="collapsed")
        with colB:
            st.markdown("**B — Behaviour**")
            row["B"] = st.text_area("", key=f"B{idx+2}", height=120, label_visibility="collapsed")
        with colC:
            st.markdown("**C — Consequence**")
            row["C"] = st.text_area("", key=f"C{idx+2}", height=120, label_visibility="collapsed")
        
        st.markdown("**H — Hypothesis**")
        row["H"] = st.text_area("", key=f"H{idx+2}", height=60, label_visibility="collapsed")
        st.markdown("---")
    
    # SAFETY RESPONSES
    st.markdown("### Safety Responses Implemented")
    safety = st.multiselect("Select all actions taken (CPI-aligned, non-restraint)",
        ["CPI Supportive stance maintained", "Area cleared of other students", "Student moved to safer location",
         "Additional staff attended", "Safety plan enacted", "Continued monitoring", "First aid provided"],
        key="safety_resp")
    
    # NOTIFICATIONS
    st.markdown("### Notifications Made")
    notifications = st.multiselect("Who was notified?",
        ["Parent/Carer", "Line Manager", "Safety & Wellbeing / SSS", "DCP", "SAPOL", 
         "First Aid Officer", "Injury report completed", "Transport arranged"],
        key="notif")
    
    # OUTCOME
    st.markdown("### Outcome")
    col1, col2 = st.columns(2)
    with col1:
        removed = st.checkbox("Student removed from learning", key="outcome_removed")
        family_contact = st.checkbox("Family contacted", key="outcome_family")
    with col2:
        safety_updated = st.checkbox("Safety plan updated", key="outcome_safety")
        transport = st.checkbox("Transport home required", key="outcome_transport")
    
    other_actions = st.text_area("Other actions / follow-up required", 
                                 placeholder="Any additional actions taken or required...",
                                 key="outcome_other")
    
    st.markdown("---")
    
    if st.button("Save Critical Incident Form", type="primary", use_container_width=True, key="save_crit"):
        if not A1_text or not B1_text or not C1_text or not H1_text:
            st.error("Please complete all ABCH fields for the primary incident")
        else:
            record = {
                "id": str(uuid.uuid4()),
                "created_at": datetime.now().isoformat(),
                "quick_incident_id": inc_id,
                "student_id": quick_inc["student_id"],
                "ABCH_primary": {"A": A1_text, "B": B1_text, "C": C1_text, "H": H1_text},
                "ABCH_additional": st.session_state.abch_rows.copy(),
                "safety_responses": safety,
                "notifications": notifications,
                "outcomes": {
                    "removed": removed,
                    "family_contact": family_contact,
                    "safety_updated": safety_updated,
                    "transport": transport,
                    "other": other_actions
                }
            }
            st.session_state.critical_incidents.append(record)
            st.session_state.abch_rows = []  # Clear rows
            
            st.success("✅ Critical incident form saved successfully")
            
            staff_email = st.session_state.current_user.get("email", "staff@example.com")
            send_critical_incident_email(record, student, staff_email)
            
            st.markdown("---")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📊 View Student Analysis", type="primary", use_container_width=True, key="view_analysis"):
                    go_to("student_analysis", selected_student_id=quick_inc["student_id"])
            with col2:
                if st.button("↩️ Back to Students", use_container_width=True, key="back_crit"):
                    go_to("program_students", selected_program=student["program"])


def render_student_analysis_page():
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## 📊 Data Analysis — {student['name']}")
    
    quick = [i for i in st.session_state.incidents if i["student_id"] == student_id]
    crit = [c for c in st.session_state.critical_incidents if c["student_id"] == student_id]
    
    if not quick and not crit:
        st.info("No incident data available yet.")
        if st.button("↩️ Back", key="back_no_data"):
            go_to("program_students", selected_program=student["program"])
        return
    
    # Build dataframe
    quick_df = pd.DataFrame(quick) if quick else pd.DataFrame()
    crit_df = pd.DataFrame(crit) if crit else pd.DataFrame()
    
    if not quick_df.empty:
        quick_df["incident_type"] = "Quick"
        quick_df["date_parsed"] = pd.to_datetime(quick_df["date"])
    
    if not crit_df.empty:
        crit_df["incident_type"] = "Critical"
        crit_df["date_parsed"] = pd.to_datetime(crit_df.get("created_at", datetime.now().isoformat()))
        crit_df["severity"] = 5
        crit_df["antecedent"] = crit_df["ABCH_primary"].apply(lambda d: d.get("A","") if isinstance(d, dict) else "")
        crit_df["behaviour_type"] = crit_df["ABCH_primary"].apply(lambda d: d.get("B","") if isinstance(d, dict) else "")
    
    full_df = pd.concat([quick_df, crit_df], ignore_index=True).sort_values("date_parsed")
    
    # OVERVIEW
    st.markdown("### Overview")
    col1, col2, col3, col4 = st.columns(4)
    with col1: st.metric("Total", len(full_df))
    with col2: st.metric("Critical", len(full_df[full_df["incident_type"] == "Critical"]))
    with col3: st.metric("Avg Severity", f"{full_df['severity'].mean():.1f}")
    with col4:
        days = max((full_df["date_parsed"].max() - full_df["date_parsed"].min()).days, 1)
        st.metric("Per Day", f"{len(full_df) / days:.1f}")
    
    st.markdown("---")
    
    # GRAPH 1: Daily Frequency
    st.markdown("### 📅 Incident Frequency")
    daily = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
    fig1 = go.Figure()
    fig1.add_trace(go.Scatter(
        x=daily["date_parsed"], y=daily["count"],
        mode='lines+markers', line=dict(color='#334155', width=2),
        marker=dict(size=7, color='#334155'),
        fill='tozeroy', fillcolor='rgba(51, 65, 85, 0.1)'
    ))
    fig1.update_layout(
        height=280, showlegend=False, xaxis_title="Date", yaxis_title="Incidents",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig1, use_container_width=True)
    
    with st.expander("💡 Interpretation"):
        st.markdown("Look for patterns (e.g., Mondays, after breaks). Schedule extra support during high-frequency periods.")
    
    st.markdown("---")
    
    # GRAPH 2: Top Behaviours
    st.markdown("### 🎯 Most Common Behaviours")
    beh_counts = full_df["behaviour_type"].value_counts().head(5)
    fig2 = go.Figure()
    fig2.add_trace(go.Bar(
        y=beh_counts.index, x=beh_counts.values,
        orientation='h', marker=dict(color='#475569'),
        text=beh_counts.values, textposition='outside'
    ))
    fig2.update_layout(
        height=280, showlegend=False, xaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig2, use_container_width=True)
    
    with st.expander("💡 Interpretation"):
        st.markdown(f"**Primary:** {beh_counts.index[0]} ({beh_counts.values[0]} incidents). Focus intervention planning on top 2-3 behaviours.")
    
    st.markdown("---")
    
    # GRAPH 3: Top Triggers
    st.markdown("### 🔍 Most Common Triggers")
    ant_counts = full_df["antecedent"].value_counts().head(5)
    fig3 = go.Figure()
    fig3.add_trace(go.Bar(
        y=ant_counts.index, x=ant_counts.values,
        orientation='h', marker=dict(color='#64748b'),
        text=ant_counts.values, textposition='outside'
    ))
    fig3.update_layout(
        height=280, showlegend=False, xaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig3, use_container_width=True)
    
    with st.expander("💡 Interpretation"):
        st.markdown(f"**Key trigger:** {ant_counts.index[0]}. Plan proactive supports before this occurs.")
    
    st.markdown("---")
    
    # GRAPH 4: Severity Trend
    st.markdown("### 📊 Severity Trend")
    fig4 = go.Figure()
    fig4.add_trace(go.Scatter(
        x=full_df["date_parsed"], y=full_df["severity"],
        mode='markers', marker=dict(size=8, color='#334155', opacity=0.6)
    ))
    if len(full_df) >= 2:
        z = np.polyfit(range(len(full_df)), full_df["severity"], 1)
        p = np.poly1d(z)
        fig4.add_trace(go.Scatter(
            x=full_df["date_parsed"], y=p(range(len(full_df))),
            mode='lines', line=dict(color='#94a3b8', width=2, dash='dash'),
            name='Trend'
        ))
    fig4.update_layout(
        height=280, yaxis=dict(range=[0, 6]), xaxis_title="Date", yaxis_title="Severity",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig4, use_container_width=True)
    
    trend_dir = "increasing" if len(full_df) >= 2 and full_df.tail(5)["severity"].mean() > full_df.head(5)["severity"].mean() else "decreasing"
    
    with st.expander("💡 Interpretation"):
        st.markdown(f"Severity appears **{trend_dir}** over time. {'Review strategies if increasing' if trend_dir == 'increasing' else 'Continue approach if decreasing'}.")
    
    st.markdown("---")
    
    # RISK SCORE
    st.markdown("### 🎲 Risk Assessment")
    recent = full_df.tail(7)
    risk_score = min(100, int(
        (len(recent) / 7 * 10) +
        (recent["severity"].mean() * 8) +
        (len(full_df[full_df["incident_type"] == "Critical"]) / len(full_df) * 50)
    ))
    
    risk_level = "LOW" if risk_score < 30 else "MODERATE" if risk_score < 60 else "HIGH"
    risk_color = "#10b981" if risk_score < 30 else "#f59e0b" if risk_score < 60 else "#ef4444"
    
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown(f"""
        <div style='background: white; padding: 2rem; border-radius: 8px; text-align: center; 
                    border: 3px solid {risk_color}; box-shadow: 0 1px 3px rgba(0,0,0,0.1);'>
            <div style='font-size: 3rem; font-weight: 700; color: {risk_color};'>{risk_score}</div>
            <div style='font-size: 1rem; color: #64748b; font-weight: 600;'>Risk Score</div>
            <div style='font-size: 1.1rem; font-weight: 700; color: {risk_color}; margin-top: 0.5rem;'>{risk_level}</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        **Action Guide:**
        - **LOW (0-29):** Maintain supports, monitor weekly
        - **MODERATE (30-59):** Increase check-ins, review triggers
        - **HIGH (60-100):** Urgent meeting, intensive supports
        """)
    
    st.markdown("---")
    
    # CLINICAL SUMMARY
    st.markdown("### 🧠 Clinical Summary")
    
    top_beh = full_df["behaviour_type"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_ant = full_df["antecedent"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_loc = full_df["location"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_session = full_df["session"].mode()[0] if len(full_df) > 0 else "Unknown"
    
    st.info(f"""
    **Key Patterns:**
    - Primary behaviour: **{top_beh}**
    - Main trigger: **{top_ant}**
    - Hotspot: **{top_loc}** during **{top_session}**
    
    **Interpretation:** {student['name']} is most vulnerable when "{top_ant}" occurs in {top_loc} during {top_session}. 
    This behaviour is a safety strategy. Use CPI Supportive stance and co-regulation.
    """)
    
    st.success(f"""
    **Recommendations:**
    1. Proactive check-in before "{top_ant}", regulated start before {top_session}
    2. CPI Supportive stance, low voice, reduce audience
    3. Teach help-seeking linked to Personal & Social Capability
    4. SMART Goal: Over 5 weeks, use help-seeking in 4/5 opportunities
    """)
    
    st.markdown("---")
    
    # EXPORT
    st.markdown("### 📄 Export Data")
    
    col1, col2 = st.columns(2)
    
    with col1:
        csv = full_df.to_csv(index=False)
        st.download_button(
            "📥 Download CSV",
            csv,
            file_name=f"{student['name']}_data.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    with col2:
        docx_file = generate_behaviour_analysis_plan_docx(
            student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level
        )
        if docx_file:
            st.download_button(
                "📄 Behaviour Analysis Plan (Word)",
                docx_file,
                file_name=f"BAP_{student['name'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    st.markdown("---")
    
    if st.button("⬅ Back to Students", type="primary", key="back_analysis"):
        go_to("program_students", selected_program=student["program"])

def main():
    init_state()
    
    if not st.session_state.logged_in:
        render_login_page()
        return
    
    page = st.session_state.current_page
    
    if page == "landing": render_landing_page()
    elif page == "program_students": render_program_students_page()
    elif page == "incident_log": render_incident_log_page()
    elif page == "critical_incident": render_critical_incident_page()
    elif page == "student_analysis": render_student_analysis_page()
    else: render_landing_page()

if __name__ == "__main__":
    main()

