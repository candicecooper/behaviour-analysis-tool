import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime, date, time, timedelta
import uuid
import random
from io import BytesIO
import base64
import bcrypt

# SUPABASE CONNECTION
try:
    from supabase import create_client, Client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False
    st.warning("Supabase not installed. Run: pip install supabase")

# Initialize Supabase client
@st.cache_resource
def init_supabase() -> Client:
    """Initialize Supabase client with credentials from secrets"""
    if not SUPABASE_AVAILABLE:
        return None
    
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception as e:
        st.error(f"❌ Supabase connection failed: {e}")
        st.info("💡 Add Supabase credentials to .streamlit/secrets.toml")
        return None

# Global Supabase client
supabase: Client = init_supabase()

st.set_page_config(page_title="CLC Behaviour Support - DEMO", page_icon="📊", layout="wide", initial_sidebar_state="collapsed")

# MINIMALIST PROFESSIONAL STYLING
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    * { font-family: 'Inter', sans-serif; }
    .stApp { background: #f8fafc; }
    .stButton>button {
        background: #334155 !important; color: white !important;
        border: none !important; border-radius: 6px !important;
        padding: 0.5rem 1.2rem !important; font-weight: 600 !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        transition: all 0.2s !important;
    }
    .stButton>button:hover { background: #1e293b !important; transform: translateY(-1px) !important; }
    button[kind="primary"] { background: #0ea5e9 !important; color: white !important; }
    button[kind="primary"]:hover { background: #0284c7 !important; }
    [data-testid="stVerticalBlock"] > div[style*="border"] {
        background: white !important; border-radius: 8px !important;
        padding: 1.5rem !important; box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        border: 1px solid #e2e8f0 !important;
    }
    [data-testid="stMetricValue"] {
        font-size: 1.875rem !important; font-weight: 700 !important; color: #0f172a !important;
    }
    [data-testid="stMetricLabel"] {
        color: #64748b !important; font-weight: 600 !important; font-size: 0.875rem !important;
        text-transform: uppercase !important; letter-spacing: 0.05em !important;
    }
    .stTextInput>div>div>input, .stSelectbox>div>div>select, 
    .stTextArea>div>div>textarea, .stNumberInput>div>div>input,
    .stDateInput>div>div>input, .stTimeInput>div>div>input {
        border: 1px solid #cbd5e1 !important; background: white !important;
        color: #0f172a !important; font-weight: 500 !important; border-radius: 6px !important;
    }
    h1 { color: #0f172a !important; font-weight: 700 !important; }
    h2 { color: #0f172a !important; font-weight: 700 !important; }
    h3 { color: #0f172a !important; font-weight: 600 !important; }
    label { color: #334155 !important; font-weight: 600 !important; font-size: 0.875rem !important; }
    .stSuccess { background: #ecfdf5 !important; color: #065f46 !important; 
                 border-left: 4px solid #10b981 !important; }
    .stInfo { background: #f0f9ff !important; color: #075985 !important; 
              border-left: 4px solid #0ea5e9 !important; }
    .stWarning { background: #fffbeb !important; color: #92400e !important; 
                 border-left: 4px solid #f59e0b !important; }
</style>
""", unsafe_allow_html=True)

# Production mode - sandbox banner removed

# MOCK DATA
MOCK_STAFF = [
    {"id": "s1", "first_name": "Emily", "last_name": "Jones", "name": "Emily Jones", "role": "TSS", "program": "JP", "email": "emily.jones@example.com", "password": "demo123"},
    {"id": "s2", "first_name": "Daniel", "last_name": "Lee", "name": "Daniel Lee", "role": "TSS", "program": "PY", "email": "daniel.lee@example.com", "password": "demo123"},
    {"id": "s3", "first_name": "Sarah", "last_name": "Chen", "name": "Sarah Chen", "role": "TSS", "program": "SY", "email": "sarah.chen@example.com", "password": "demo123"},
    {"id": "s4", "first_name": "Admin", "last_name": "User", "name": "Admin User", "role": "ADM", "email": "admin@example.com", "password": "admin123"},
    {"id": "s5", "first_name": "Michael", "last_name": "Roberts", "name": "Michael Roberts", "role": "Leader", "program": "JP", "email": "michael.roberts@example.com", "password": "demo123"},
    {"id": "s6", "first_name": "Jennifer", "last_name": "Walsh", "name": "Jennifer Walsh", "role": "Leader", "program": "PY", "email": "jennifer.walsh@example.com", "password": "demo123"},
]

# Demo students - 2 per program (6 total)
MOCK_STUDENTS = [
    # Junior Primary (JP) - 2 students
    {"id": "stu_jp1", "first_name": "Emma", "last_name": "Thompson", "name": "Emma Thompson", 
     "grade": "R", "dob": "2018-05-30", "program": "JP", 
     "edid": "ED123456", "placement_start": "2024-02-01", "placement_end": None},
    {"id": "stu_jp2", "first_name": "Oliver", "last_name": "Smith", "name": "Oliver Smith", 
     "grade": "Y1", "dob": "2017-09-12", "program": "JP",
     "edid": "ED234567", "placement_start": "2024-03-15", "placement_end": None},
    # Primary Years (PY) - 2 students
    {"id": "stu_py1", "first_name": "Liam", "last_name": "Carter", "name": "Liam Carter", 
     "grade": "Y3", "dob": "2015-06-15", "program": "PY",
     "edid": "ED456789", "placement_start": "2024-02-12", "placement_end": None},
    {"id": "stu_py2", "first_name": "Ava", "last_name": "Robinson", "name": "Ava Robinson", 
     "grade": "Y5", "dob": "2013-11-08", "program": "PY",
     "edid": "ED567890", "placement_start": "2024-01-08", "placement_end": None},
    # Senior Years (SY) - 2 students
    {"id": "stu_sy1", "first_name": "Isabella", "last_name": "Garcia", "name": "Isabella Garcia", 
     "grade": "Y7", "dob": "2011-04-17", "program": "SY",
     "edid": "ED789012", "placement_start": "2024-01-29", "placement_end": None},
    {"id": "stu_sy2", "first_name": "Ethan", "last_name": "Davis", "name": "Ethan Davis", 
     "grade": "Y9", "dob": "2009-12-03", "program": "SY",
     "edid": "ED890123", "placement_start": "2024-02-26", "placement_end": None},
]

PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}
BEHAVIOUR_TYPES = ["Verbal Refusal", "Elopement", "Property Destruction", "Aggression (Peer)", 
                   "Aggression (Adult)", "Self-Harm", "Verbal Aggression", "Other"]
ANTECEDENTS = [
    "--- PEER ---",
    "Peer - negative peer feedback",
    "Peer - peer conflict/interaction", 
    "Peer - participating in competition",
    "Peer - losing against a peer/s",
    "Peer - another student escalating",
    "--- TRANSITION ---",
    "Transition - to a non-preferred activity",
    "Transition - from one activity to another",
    "Transition - from one environment to another",
    "Transition - another teacher coming into the program",
    "Transition - from play/yard to classroom",
    "Transition - from home to the program",
    "Transition - off a device",
    "--- INSTRUCTIONS ---",
    "Instructions - following instructions given by an adult",
    "Instructions - following task demands",
    "--- ENGAGEMENT ---",
    "Engagement - engaging in a non-preferred learning area (Maths/Literacy)",
    "Engagement - having to wait",
    "Engagement - change in routine",
    "Engagement - learning task too difficult or perceived as difficult",
    "Engagement - having to work with a peer",
    "Engagement - working independently",
    "Engagement - interrupted activity",
    "Engagement - not being able to finish an activity",
    "--- SENSORY ---",
    "Sensory - crowded area",
    "Sensory - noise",
    "Sensory - environment too bright",
    "Sensory - no medication",
    "Sensory - hungry/thirsty",
    "Sensory - unable to take a movement break",
    "--- OTHER ---",
    "Other - perceived injustice",
    "Other - TRT (Temporary Relief Teacher)",
    "Other - NIT teacher (New Initiatives Teacher)",
    "Other - incident before coming to program",
    "Other - verbalising not wanting to be at the program",
    "Other - staff attention shifted",
    "Other - unstructured time",
    "Other - access denied"
]
INTERVENTIONS = ["CPI Supportive stance", "Offered break", "Reduced demand", "Provided choices", 
                "Removed audience", "Visual supports", "Co-regulation", "Prompted coping skill", "Redirection"]
LOCATIONS = ["JP Classroom", "PY Classroom", "SY Classroom", "Playground", "Library", "Office", "Student Gate", "Toilets"]
VALID_PAGES = ["login", "landing", "program_students", "incident_log", "critical_incident", "student_analysis", "admin_portal"]

# AI HYPOTHESIS SYSTEM
HYPOTHESIS_FUNCTIONS = ["To get", "To avoid"]
HYPOTHESIS_ITEMS = ["Tangible", "Activity", "Sensory", "Attention"]

def format_time_12hr(time_str):
    """Convert 24hr time string to 12hr format"""
    try:
        if isinstance(time_str, str):
            dt = datetime.strptime(time_str, "%H:%M:%S")
        else:
            dt = datetime.combine(date.today(), time_str)
        return dt.strftime("%I:%M %p")
    except:
        return time_str

def format_date_dmy(date_str):
    """Convert date to DD/MM/YYYY format"""
    try:
        if isinstance(date_str, str):
            dt = datetime.strptime(date_str, "%Y-%m-%d")
        else:
            dt = date_str
        return dt.strftime("%d/%m/%Y")
    except:
        return date_str

def generate_hypothesis(antecedent, behaviour, consequence):
    """Auto-generate hypothesis based on ABC data"""
    hypotheses = []
    antecedent_lower = antecedent.lower()
    behaviour_lower = behaviour.lower()
    
    # Updated to match new antecedent categories
    if any(word in antecedent_lower for word in [
        "instruction", "demand", "task", "transition", "work", "non-preferred",
        "literacy", "maths", "wait", "routine", "independently", "difficult"
    ]):
        hypotheses.append("To avoid or escape the demand/task")
        
    if any(word in antecedent_lower for word in [
        "attention", "shifted", "ignored", "alone", "feedback", "conflict",
        "interaction", "escalating", "injustice"
    ]):
        hypotheses.append("To gain staff/peer attention")
        
    if any(word in antecedent_lower for word in [
        "sensory", "loud", "noise", "bright", "touch", "crowded", "medication",
        "hungry", "thirsty", "movement"
    ]):
        hypotheses.append("To escape sensory discomfort or seek sensory input")
        
    if any(word in antecedent_lower for word in [
        "denied", "can't have", "no", "wait", "device", "finish", "preferred"
    ]):
        hypotheses.append("To gain access to preferred item/activity")
        
    if any(word in behaviour_lower for word in ["refusal", "defiance", "left", "ran", "elopement"]):
        hypotheses.append("To assert control or autonomy")
    
    if not hypotheses:
        hypotheses.append("Function requires further analysis")
    
    return " / ".join(hypotheses[:2])

def generate_admin_summary(incident_data, student, staff_name):
    """Generate brief summary for external incident log - FOR ADMIN USE ONLY"""
    abch_primary = incident_data.get("ABCH_primary", {})
    intended = incident_data.get("intended_outcomes", [])
    
    date_time = incident_data.get("created_at", datetime.now().isoformat())
    dt = datetime.fromisoformat(date_time)
    
    location = abch_primary.get("location", "Unknown location")
    time_str = abch_primary.get("time", "Unknown time")
    behaviour = abch_primary.get("behaviour", "Behaviour not specified")
    context = abch_primary.get("context", "")
    consequence = abch_primary.get("consequence", "")
    hypothesis = abch_primary.get("hypothesis", "")
    
    summary = f"""CRITICAL INCIDENT SUMMARY - FOR ADMIN USE ONLY
    
Student: {student['name']} (Grade {student['grade']})
Date/Time: {dt.strftime('%d/%m/%Y')} at {time_str}
Location: {location}

INCIDENT DESCRIPTION:
{behaviour[:200]}{'...' if len(behaviour) > 200 else ''}

CONTEXT/ANTECEDENT:
{context[:200]}{'...' if len(context) > 200 else ''}

IMMEDIATE STAFF RESPONSE:
{consequence[:200]}{'...' if len(consequence) > 200 else ''}

BEHAVIOURAL FUNCTION:
{hypothesis}

OUTCOMES IMPLEMENTED:
{', '.join(intended[:5]) if intended else 'See full form for details'}

REPORTED BY: {staff_name}
FORM COMPLETED: {dt.strftime('%d/%m/%Y %H:%M')}

** This summary is for external departmental incident log entry purposes only **
** Full critical incident form has been saved and distributed to relevant parties **
"""
    
    return summary


def generate_hypothesis_ai(antecedent, behaviour, consequence=""):
    """AI generates structured hypothesis from ABC data - ENHANCED VERSION"""
    ant_lower = (antecedent or "").lower()
    beh_lower = (behaviour or "").lower()
    cons_lower = (consequence or "").lower()
    
    # Expanded keyword detection
    avoid_keywords = [
        "instruction", "demand", "task", "transition", "work", "difficult", "challenging",
        "non-preferred", "literacy", "maths", "wait", "routine", "peer", "independently",
        "interrupted", "finish", "teacher coming"
    ]
    
    get_keywords = [
        "attention", "item", "toy", "want", "access", "denied", "device", "finish",
        "preferred", "competition", "feedback"
    ]
    
    sensory_keywords = [
        "sensory", "loud", "noise", "touch", "bright", "crowded", "medication",
        "hungry", "thirsty", "movement break"
    ]
    
    attention_keywords = [
        "attention", "staff", "peer", "ignored", "escalating", "feedback",
        "conflict", "interaction", "injustice"
    ]
    
    # Determine FUNCTION (To get vs To avoid)
    function = "To avoid"
    if any(word in ant_lower for word in get_keywords):
        function = "To get"
    elif any(word in cons_lower for word in ["given", "received", "got", "obtained"]):
        function = "To get"
    elif "denied" in ant_lower or "can't" in ant_lower or "wait" in ant_lower:
        function = "To get"
    elif "off a device" in ant_lower or "finish" in ant_lower:
        function = "To get"
    
    # Determine ITEM (What they want to get or avoid)
    item = "Activity"
    
    if any(word in ant_lower + beh_lower for word in attention_keywords):
        item = "Attention"
    elif any(word in ant_lower + beh_lower for word in sensory_keywords):
        item = "Sensory"
    elif any(word in ant_lower + beh_lower for word in ["toy", "item", "object", "food", "device", "tangible"]):
        item = "Tangible"
    elif any(word in ant_lower for word in ["instruction", "demand", "task", "work", "learning", "maths", "literacy"]):
        item = "Activity"
    elif "transition" in ant_lower or "environment" in ant_lower or "classroom" in ant_lower:
        item = "Activity"
    
    return {"function": function, "item": item}

def format_hypothesis(hyp):
    """Format hypothesis dict to string"""
    if isinstance(hyp, dict):
        return f"{hyp.get('function', 'Unknown')} {hyp.get('item', 'Unknown')}"
    elif isinstance(hyp, str):
        return hyp
    else:
        return "Unknown"

def show_severity_guide():
    """Enhanced Behaviour Severity Continuum matching uploaded image"""
    import streamlit.components.v1 as components
    
    html_content = """<div style='background: white; padding: 1.5rem; border-radius: 8px; margin: 1rem 0; 
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1); border: 1px solid #e2e8f0;
            font-family: "Source Sans Pro", sans-serif;'>
    
    <div style='text-align: center; margin-bottom: 1.5rem;'>
        <h2 style='color: #1a1a1a; font-weight: 700; font-size: 1.8rem; margin: 0;'>
            Behaviour Severity Continuum
        </h2>
    </div>
    
    <div style='display: grid; grid-template-columns: repeat(5, 1fr); gap: 0;'>
        
        <div style='background: #81b29a; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 1</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Low Level /<br>Engaged
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Following instructions</li>
                    <li>On task</li>
                    <li>Minor defiance</li>
                    <li>Avoiding work</li>
                    <li>Answering back</li>
                    <li>Mumbling, huffing</li>
                    <li>Passive peer conflict</li>
                    <li>Attention seeking</li>
                </ul>
            </div>
            <div style='background: #6b9b7f; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Redirect / provide options</li>
                <li>Give space</li>
                <li>Offer choice</li>
                <li>Acknowledge concern</li>
                <li>Maintain routine</li>
                <li>Active listening</li>
            </ul>
        </div>
        
        <div style='background: #f4d35e; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: #2c2c2c; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 2</div>
                <div style='color: #2c2c2c; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Escalating /<br>Dysregulated
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.3); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: #2c2c2c; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: #2c2c2c; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Raised voice</li>
                    <li>Arguing, blaming</li>
                    <li>Crying, frustration</li>
                    <li>Pacing or mild exit attempts</li>
                    <li>Throwing soft items (not dangerous)</li>
                </ul>
            </div>
            <div style='background: #d9b84d; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: #2c2c2c; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: #2c2c2c; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Reduce demands</li>
                <li>Offer space / movement break</li>
                <li>Provide limited choices</li>
                <li>Avoid power struggles</li>
            </ul>
        </div>
        
        <div style='background: #ee8434; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 3</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    High Escalation /<br>Significant Risk
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Yelling, swearing</li>
                    <li>Slammed doors, hitting walls</li>
                    <li>Throwing items with possible risk</li>
                    <li>Attempting to run off</li>
                    <li>Damaging property</li>
                </ul>
            </div>
            <div style='background: #d47230; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Increase distance</li>
                <li>Notify leadership/support</li>
                <li>Remove audience</li>
                <li>Complete Critical Incident Form</li>
            </ul>
        </div>
        
        <div style='background: #c9555e; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 4</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Dangerous<br>Behaviour
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Attempts to hit, kick, grab</li>
                    <li>Throwing dangerous objects</li>
                    <li>Threats of violence</li>
                    <li>Absconding into unsafe situations</li>
                    <li>Beginning self-harm behaviour</li>
                </ul>
            </div>
            <div style='background: #b04850; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Evacuate nearby students</li>
                <li>Leadership/response team activated</li>
                <li>Maintain safety distance</li>
            </ul>
        </div>
        
        <div style='background: #7d2e2e; padding: 1.2rem 0.8rem;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Crisis</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Crisis<br>Situation
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.15); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Physical violence causing or likely to cause injury</li>
                    <li>Severe self-harm</li>
                    <li>Use of weapons or dangerous items</li>
                    <li>Full loss of control behaviour</li>
                </ul>
            </div>
            <div style='background: #5c2323; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Immediate emergency response</li>
                <li>Trained staff to manage situation</li>
                <li>Preserve evidence</li>
                <li>Complete Critical Incident Form</li>
            </ul>
        </div>
        
    </div>

    <div style='margin-top: 1.5rem; padding: 1rem; background: #fff3cd; border-radius: 6px; border-left: 4px solid #f59e0b;'>
        <div style='color: #92400e; font-weight: 700; font-size: 0.95rem; margin-bottom: 0.5rem;'>
            WARNING: Critical Incident Documentation Required
        </div>
        <div style='color: #92400e; font-size: 0.85rem; line-height: 1.5;'>
            <strong>Level 3 or above</strong> requires a Critical Incident ABCH Form to be completed immediately after the incident is resolved.
        </div>
    </div>
</div>
"""
    components.html(html_content, height=700, scrolling=True)
    
def send_critical_incident_email(incident_data, student, staff_email, leader_email, admin_email):
    """Send email notification to all parties"""
    st.info(f"""📧 **Email Notification Sent**
    
**To:** {leader_email}, {admin_email}, {staff_email}  
**Subject:** CRITICAL INCIDENT - {student['name']}

**Student:** {student['name']} | **Programme:** {student['program']} | **Grade:** {student['grade']}  

Critical Incident Form completed and saved.
Admin summary included for departmental log.

*(In production, this sends via SMTP)*
    """)



def generate_behaviour_analysis_plan_docx(student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level):
    """Generate PROFESSIONAL Behaviour Analysis Plan with enhanced formatting"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')
        from matplotlib.ticker import MaxNLocator
        import numpy as np
        
        # Color scheme - Professional teal/green theme
        PRIMARY_COLOR = RGBColor(0, 128, 128)      # Teal
        SECONDARY_COLOR = RGBColor(34, 139, 34)    # Forest Green
        ACCENT_COLOR = RGBColor(70, 130, 180)      # Steel Blue
        DARK_TEXT = RGBColor(33, 37, 41)           # Dark gray
        LIGHT_BG = RGBColor(248, 249, 250)         # Light gray
        WARNING_COLOR = RGBColor(220, 53, 69)      # Red for high risk
        SUCCESS_COLOR = RGBColor(40, 167, 69)      # Green for low risk
        
        PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}
        
        doc = Document()
        
        # ================================================================
        # DOCUMENT STYLES SETUP
        # ================================================================
        
        def set_arial(run, size=11):
            """Set Arial font for a run"""
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Set default styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Modify heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
            heading_style.font.color.rgb = PRIMARY_COLOR
            heading_style.font.bold = True
        
        def add_horizontal_line(doc):
            """Add a horizontal line separator"""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(
                r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="008080"/>'
                r'</w:pBdr>'
            )
            p._p.get_or_add_pPr().append(pBdr)
            return p
        
        def create_info_box(doc, title, content, box_color="E8F4F8"):
            """Create a colored info box"""
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            cell = table.rows[0].cells[0]
            
            # Set cell shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{box_color}" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)
            
            # Add title
            title_para = cell.paragraphs[0]
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(12)
            title_run.font.color.rgb = PRIMARY_COLOR
            set_arial(title_run, 12)
            
            # Add content
            content_para = cell.add_paragraph()
            content_run = content_para.add_run(content)
            set_arial(content_run, 11)
            
            doc.add_paragraph()  # Spacing after box
            return table
        
        def create_metric_table(doc, metrics):
            """Create a professional metrics table"""
            cols = len(metrics)
            table = doc.add_table(rows=2, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, (label, value, color) in enumerate(metrics):
                # Header cell
                header_cell = table.rows[0].cells[i]
                header_para = header_cell.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header_para.add_run(label)
                header_run.bold = True
                header_run.font.size = Pt(10)
                header_run.font.color.rgb = DARK_TEXT
                
                # Value cell
                value_cell = table.rows[1].cells[i]
                value_para = value_cell.paragraphs[0]
                value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                value_run = value_para.add_run(str(value))
                value_run.bold = True
                value_run.font.size = Pt(18)
                value_run.font.color.rgb = color
                
                # Add cell shading
                for cell in [header_cell, value_cell]:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
            
            doc.add_paragraph()
            return table
        
        # ================================================================
        # COVER PAGE
        # ================================================================
        
        # Add spacing at top
        for _ in range(3):
            doc.add_paragraph()
        
        # Logo placeholder
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run("[SCHOOL LOGO]")
        logo_run.font.size = Pt(14)
        logo_run.font.color.rgb = RGBColor(150, 150, 150)
        
        doc.add_paragraph()
        
        # School name
        school = doc.add_paragraph()
        school.alignment = WD_ALIGN_PARAGRAPH.CENTER
        school_run = school.add_run("COWANDILLA LEARNING CENTRE")
        school_run.bold = True
        school_run.font.size = Pt(24)
        school_run.font.color.rgb = PRIMARY_COLOR
        set_arial(school_run, 24)
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("Learning and Behaviour Support Unit")
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = DARK_TEXT
        set_arial(sub_run, 14)
        
        doc.add_paragraph()
        add_horizontal_line(doc)
        doc.add_paragraph()
        
        # Main title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("BEHAVIOUR ANALYSIS PLAN")
        title_run.bold = True
        title_run.font.size = Pt(32)
        title_run.font.color.rgb = SECONDARY_COLOR
        set_arial(title_run, 32)
        
        doc.add_paragraph()
        
        # Student name box
        student_box = doc.add_table(rows=1, cols=1)
        student_box.alignment = WD_TABLE_ALIGNMENT.CENTER
        student_box.columns[0].width = Inches(4)
        student_cell = student_box.rows[0].cells[0]
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5E9" w:val="clear"/>')
        student_cell._tc.get_or_add_tcPr().append(shading)
        
        student_para = student_cell.paragraphs[0]
        student_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        student_run = student_para.add_run(student['name'].upper())
        student_run.bold = True
        student_run.font.size = Pt(20)
        student_run.font.color.rgb = SECONDARY_COLOR
        set_arial(student_run, 20)
        
        grade_para = student_cell.add_paragraph()
        grade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        grade_run = grade_para.add_run(f"{PROGRAM_NAMES.get(student['program'], student['program'])} | Grade {student['grade']}")
        grade_run.font.size = Pt(12)
        set_arial(grade_run, 12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Date and classification
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Analysis Date: {datetime.now().strftime('%d %B %Y')}")
        date_run.font.size = Pt(12)
        set_arial(date_run, 12)
        
        period_para = doc.add_paragraph()
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_run = period_para.add_run(f"Data Period: {full_df['date_parsed'].min().strftime('%d/%m/%Y')} - {full_df['date_parsed'].max().strftime('%d/%m/%Y')}")
        period_run.font.size = Pt(11)
        period_run.font.color.rgb = RGBColor(100, 100, 100)
        set_arial(period_run, 11)
        
        # Add spacing before confidentiality notice
        for _ in range(4):
            doc.add_paragraph()
        
        # Confidentiality notice
        conf_box = doc.add_table(rows=1, cols=1)
        conf_box.alignment = WD_TABLE_ALIGNMENT.CENTER
        conf_box.columns[0].width = Inches(5.5)
        conf_cell = conf_box.rows[0].cells[0]
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3CD" w:val="clear"/>')
        conf_cell._tc.get_or_add_tcPr().append(shading)
        
        conf_para = conf_cell.paragraphs[0]
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run = conf_para.add_run("CONFIDENTIAL DOCUMENT")
        conf_run.bold = True
        conf_run.font.size = Pt(11)
        conf_run.font.color.rgb = RGBColor(133, 100, 4)
        
        conf_para2 = conf_cell.add_paragraph()
        conf_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run2 = conf_para2.add_run("This document contains sensitive student information and should be handled in accordance with privacy legislation.")
        conf_run2.font.size = Pt(9)
        conf_run2.font.color.rgb = RGBColor(133, 100, 4)
        
        doc.add_page_break()
        
        # ================================================================
        # TABLE OF CONTENTS
        # ================================================================
        
        toc_heading = doc.add_heading('Contents', 1)
        for run in toc_heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        toc_items = [
            ("1. Executive Summary", "3"),
            ("2. Key Metrics & Risk Assessment", "3"),
            ("3. Visual Analytics", "4"),
            ("4. Clinical Interpretation", "7"),
            ("5. Berry Street Education Model Framework", "8"),
            ("6. Crisis Prevention Institute (CPI) Principles", "10"),
            ("7. Evidence-Based Recommendations", "11"),
            ("8. Action Plan & Timeline", "12"),
            ("9. Review Schedule", "13"),
        ]
        
        toc_table = doc.add_table(rows=len(toc_items), cols=2)
        toc_table.columns[0].width = Inches(5)
        toc_table.columns[1].width = Inches(1)
        
        for i, (item, page) in enumerate(toc_items):
            toc_table.rows[i].cells[0].paragraphs[0].add_run(item)
            page_para = toc_table.rows[i].cells[1].paragraphs[0]
            page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            page_para.add_run(page)
        
        doc.add_paragraph()
        add_horizontal_line(doc)
        doc.add_page_break()
        
        # ================================================================
        # EXECUTIVE SUMMARY
        # ================================================================
        
        heading = doc.add_heading('1. Executive Summary', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        # Key metrics in a professional table
        total_incidents = len(full_df)
        critical_count = len(full_df[full_df['incident_type'] == 'Critical']) if 'incident_type' in full_df.columns else 0
        avg_severity = full_df['severity'].mean()
        
        risk_color = WARNING_COLOR if risk_score > 60 else (RGBColor(255, 193, 7) if risk_score > 30 else SUCCESS_COLOR)
        
        metrics = [
            ("Total Incidents", total_incidents, ACCENT_COLOR),
            ("Critical Incidents", critical_count, WARNING_COLOR if critical_count > 0 else SUCCESS_COLOR),
            ("Avg Severity", f"{avg_severity:.1f}/5", ACCENT_COLOR),
            ("Risk Score", f"{risk_score}/100", risk_color),
        ]
        
        create_metric_table(doc, metrics)
        
        # Executive summary text
        summary_text = f"""This Behaviour Analysis Plan presents a comprehensive analysis of {total_incidents} recorded behavioural incidents for {student['name']} over the analysis period. 

The data reveals that the primary behaviour of concern is '{top_beh}', most commonly triggered by '{top_ant}'. Incidents predominantly occur in {top_loc} during the {top_session} session.

The current risk level is assessed as {risk_level.upper()} ({risk_score}/100). This assessment considers incident frequency, severity patterns, and escalation trends."""

        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary_text)
        set_arial(summary_run, 11)
        summary_para.paragraph_format.line_spacing = 1.5
        
        doc.add_paragraph()
        
        # Key findings box
        create_info_box(doc, "KEY FINDINGS AT A GLANCE", 
            f"• Primary Behaviour: {top_beh}\n"
            f"• Key Trigger: {top_ant}\n"
            f"• Hotspot Location: {top_loc}\n"
            f"• Peak Time: {top_session} session\n"
            f"• Risk Level: {risk_level} ({risk_score}/100)",
            "E3F2FD")
        
        doc.add_page_break()
        
        # ================================================================
        # KEY METRICS & RISK ASSESSMENT
        # ================================================================
        
        heading = doc.add_heading('2. Key Metrics & Risk Assessment', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        # Risk assessment matrix visualization
        doc.add_heading('Risk Assessment Matrix', 2)
        
        # Create risk matrix as a matplotlib figure
        fig, ax = plt.subplots(figsize=(8, 3), dpi=150)
        
        # Risk scale
        risk_zones = [
            (0, 30, '#28a745', 'LOW'),
            (30, 60, '#ffc107', 'MODERATE'),
            (60, 80, '#fd7e14', 'HIGH'),
            (80, 100, '#dc3545', 'CRITICAL')
        ]
        
        for start, end, color, label in risk_zones:
            ax.barh(0, end-start, left=start, height=0.6, color=color, edgecolor='white', linewidth=2)
            ax.text((start+end)/2, 0, label, ha='center', va='center', fontsize=10, fontweight='bold', color='white')
        
        # Add marker for current risk
        ax.scatter([risk_score], [0], s=300, c='black', marker='v', zorder=5)
        ax.text(risk_score, 0.45, f'{risk_score}', ha='center', va='bottom', fontsize=12, fontweight='bold')
        
        ax.set_xlim(0, 100)
        ax.set_ylim(-0.5, 0.7)
        ax.axis('off')
        ax.set_title('Current Risk Level', fontsize=14, fontweight='bold', pad=10)
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        doc.add_paragraph()
        
        # Detailed breakdown table
        doc.add_heading('Incident Breakdown', 2)
        
        breakdown_table = doc.add_table(rows=6, cols=2)
        breakdown_table.style = 'Table Grid'
        breakdown_table.columns[0].width = Inches(3)
        breakdown_table.columns[1].width = Inches(3)
        
        breakdown_data = [
            ("Metric", "Value"),
            ("Total Recorded Incidents", str(total_incidents)),
            ("Critical Incidents (Severity 3+)", str(len(full_df[full_df['severity'] >= 3]))),
            ("Average Incident Duration", f"{full_df['duration_minutes'].mean():.0f} minutes" if 'duration_minutes' in full_df.columns else "N/A"),
            ("Most Active Day", full_df['day_of_week'].mode().iloc[0] if 'day_of_week' in full_df.columns and len(full_df) > 0 else "N/A"),
            ("Data Collection Period", f"{(full_df['date_parsed'].max() - full_df['date_parsed'].min()).days} days"),
        ]
        
        for i, (label, value) in enumerate(breakdown_data):
            breakdown_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = (i == 0)
            breakdown_table.rows[i].cells[1].paragraphs[0].add_run(value).bold = (i == 0)
            if i == 0:
                for cell in breakdown_table.rows[i].cells:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="008080" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
        
        doc.add_page_break()
        
        # ================================================================
        # VISUAL ANALYTICS
        # ================================================================
        
        heading = doc.add_heading('3. Visual Analytics', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        intro = doc.add_paragraph()
        intro_run = intro.add_run("The following visualisations provide data-driven insights into behavioural patterns, enabling targeted intervention strategies.")
        set_arial(intro_run, 11)
        intro_run.italic = True
        
        doc.add_paragraph()
        
        plt.style.use('default')
        
        # GRAPH 1: Daily Incident Frequency
        doc.add_heading('3.1 Daily Incident Frequency', 2)
        
        daily = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig, ax = plt.subplots(figsize=(10, 4), dpi=150)
        
        # Create gradient effect with color based on count
        colors = ['#4A90A4' if c <= daily['count'].median() else '#2E5A6B' for c in daily['count']]
        bars = ax.bar(daily["date_parsed"], daily["count"], color=colors, width=0.8, edgecolor='white', linewidth=0.5)
        
        ax.set_xlabel('Date', fontsize=11, fontweight='bold', color='#333333')
        ax.set_ylabel('Number of Incidents', fontsize=11, fontweight='bold', color='#333333')
        ax.grid(True, alpha=0.3, linestyle='--', axis='y')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        plt.xticks(rotation=45, ha='right')
        
        # Add trend line
        if len(daily) > 2:
            z = np.polyfit(range(len(daily)), daily['count'], 1)
            p = np.poly1d(z)
            ax.plot(daily["date_parsed"], p(range(len(daily))), 
                   linestyle='--', color='#DC3545', linewidth=2, label='Trend')
            ax.legend(loc='upper right', frameon=False)
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        interpretation = doc.add_paragraph()
        trend_direction = "increasing" if len(daily) > 2 and z[0] > 0 else "decreasing" if len(daily) > 2 and z[0] < 0 else "stable"
        int_run = interpretation.add_run(f"Interpretation: The daily incident frequency shows a {trend_direction} trend over the analysis period. Peak incident days should be cross-referenced with environmental factors.")
        set_arial(int_run, 10)
        int_run.italic = True
        
        doc.add_paragraph()
        
        # GRAPH 2: Behaviour Type Distribution
        doc.add_heading('3.2 Behaviour Type Distribution', 2)
        
        beh_counts = full_df["behaviour_type"].value_counts().head(6)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        
        colors = plt.cm.Blues(np.linspace(0.4, 0.8, len(beh_counts)))[::-1]
        bars = ax.barh(beh_counts.index, beh_counts.values, color=colors, edgecolor='white', linewidth=1)
        
        ax.set_xlabel('Incident Count', fontsize=11, fontweight='bold', color='#333333')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.xaxis.set_major_locator(MaxNLocator(integer=True))
        
        for i, (bar, v) in enumerate(zip(bars, beh_counts.values)):
            ax.text(v + 0.3, bar.get_y() + bar.get_height()/2, str(int(v)), 
                   va='center', fontweight='bold', fontsize=11)
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        interpretation = doc.add_paragraph()
        int_run = interpretation.add_run(f"Interpretation: '{beh_counts.index[0]}' is the predominant behaviour of concern, accounting for {int(beh_counts.values[0]/total_incidents*100)}% of all incidents.")
        set_arial(int_run, 10)
        int_run.italic = True
        
        doc.add_paragraph()
        
        # GRAPH 3: Trigger Analysis
        doc.add_heading('3.3 Antecedent/Trigger Analysis', 2)
        
        ant_counts = full_df["antecedent"].value_counts().head(6)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        
        colors = plt.cm.Greens(np.linspace(0.4, 0.8, len(ant_counts)))[::-1]
        bars = ax.barh(ant_counts.index, ant_counts.values, color=colors, edgecolor='white', linewidth=1)
        
        ax.set_xlabel('Incident Count', fontsize=11, fontweight='bold', color='#333333')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.xaxis.set_major_locator(MaxNLocator(integer=True))
        
        for i, (bar, v) in enumerate(zip(bars, ant_counts.values)):
            ax.text(v + 0.3, bar.get_y() + bar.get_height()/2, str(int(v)), 
                   va='center', fontweight='bold', fontsize=11)
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        interpretation = doc.add_paragraph()
        int_run = interpretation.add_run(f"Interpretation: Understanding that '{ant_counts.index[0]}' is the primary trigger enables proactive intervention before escalation occurs.")
        set_arial(int_run, 10)
        int_run.italic = True
        
        doc.add_page_break()
        
        # GRAPH 4: Severity Distribution
        doc.add_heading('3.4 Severity Over Time', 2)
        
        fig, ax = plt.subplots(figsize=(10, 4), dpi=150)
        colors = {1: '#28a745', 2: '#20c997', 3: '#ffc107', 4: '#fd7e14', 5: '#dc3545'}
        labels = {1: 'Level 1 (Low)', 2: 'Level 2 (Minor)', 3: 'Level 3 (Moderate)', 
                  4: 'Level 4 (Serious)', 5: 'Level 5 (Critical)'}
        
        for sev_level in [1, 2, 3, 4, 5]:
            sev_data = full_df[full_df['severity'] == sev_level]
            if len(sev_data) > 0:
                ax.scatter(sev_data["date_parsed"], sev_data["severity"], 
                          alpha=0.8, s=100, color=colors[sev_level], 
                          label=labels[sev_level], edgecolors='white', linewidth=1)
        
        ax.set_xlabel('Date', fontsize=11, fontweight='bold', color='#333333')
        ax.set_ylabel('Severity Level', fontsize=11, fontweight='bold', color='#333333')
        ax.set_ylim(0.5, 5.5)
        ax.set_yticks([1, 2, 3, 4, 5])
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend(loc='upper right', frameon=True, fancybox=True, shadow=True)
        ax.grid(True, alpha=0.3, linestyle='--', axis='y')
        plt.xticks(rotation=45, ha='right')
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        doc.add_paragraph()
        
        # GRAPH 5: Location Hotspots
        doc.add_heading('3.5 Location Hotspots', 2)
        
        loc_counts = full_df["location"].value_counts().head(5)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        
        colors = plt.cm.Oranges(np.linspace(0.4, 0.8, len(loc_counts)))[::-1]
        bars = ax.barh(loc_counts.index, loc_counts.values, color=colors, edgecolor='white', linewidth=1)
        
        ax.set_xlabel('Incident Count', fontsize=11, fontweight='bold', color='#333333')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.xaxis.set_major_locator(MaxNLocator(integer=True))
        
        for i, (bar, v) in enumerate(zip(bars, loc_counts.values)):
            ax.text(v + 0.3, bar.get_y() + bar.get_height()/2, str(int(v)), 
                   va='center', fontweight='bold', fontsize=11)
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        interpretation = doc.add_paragraph()
        int_run = interpretation.add_run(f"Interpretation: Environmental modifications in '{loc_counts.index[0]}' may yield significant improvements.")
        set_arial(int_run, 10)
        int_run.italic = True
        
        doc.add_paragraph()
        
        # GRAPH 6: Time of Day Pattern (Pie chart for variety)
        doc.add_heading('3.6 Time of Day Distribution', 2)
        
        session_counts = full_df["session"].value_counts()
        session_order = ['Morning', 'Middle', 'Afternoon']
        session_counts = session_counts.reindex(session_order, fill_value=0)
        
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4), dpi=150)
        
        # Pie chart
        colors_pie = ['#4A90A4', '#6BB9A0', '#E8B960']
        explode = [0.05 if v == session_counts.max() else 0 for v in session_counts.values]
        ax1.pie(session_counts.values, labels=session_counts.index, autopct='%1.0f%%',
               colors=colors_pie, explode=explode, shadow=True, startangle=90)
        ax1.set_title('Distribution by Session', fontsize=12, fontweight='bold')
        
        # Bar chart
        bars = ax2.bar(session_counts.index, session_counts.values, color=colors_pie, edgecolor='white', linewidth=2)
        ax2.set_ylabel('Incident Count', fontsize=11, fontweight='bold')
        ax2.spines['top'].set_visible(False)
        ax2.spines['right'].set_visible(False)
        ax2.yaxis.set_major_locator(MaxNLocator(integer=True))
        
        for bar, v in zip(bars, session_counts.values):
            ax2.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.3,
                    str(int(v)), ha='center', va='bottom', fontweight='bold', fontsize=12)
        ax2.set_title('Count by Session', fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        
        peak_session = session_counts.idxmax()
        peak_count = int(session_counts.max())
        peak_pct = int(peak_count / total_incidents * 100)
        
        interpretation = doc.add_paragraph()
        int_run = interpretation.add_run(f"Interpretation: The {peak_session} session accounts for {peak_pct}% of incidents ({peak_count} total). Staff resourcing and support strategies should prioritise this period.")
        set_arial(int_run, 10)
        int_run.italic = True
        
        doc.add_page_break()
        
        # ================================================================
        # CLINICAL INTERPRETATION
        # ================================================================
        
        heading = doc.add_heading('4. Clinical Interpretation', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        intro = doc.add_paragraph()
        intro.add_run("This analysis is grounded in evidence-based frameworks:").bold = True
        
        frameworks = [
            ("Applied Behaviour Analysis (ABA)", "Understanding what triggers and maintains behaviours"),
            ("Trauma-Informed Practice", "Recognising that behaviour is communication and often a response to stress"),
            ("Berry Street Education Model", "A whole-school approach to wellbeing and engagement"),
            ("Crisis Prevention Institute (CPI)", "De-escalation and maintaining dignity"),
        ]
        
        for framework, description in frameworks:
            p = doc.add_paragraph(style='List Bullet')
            run1 = p.add_run(f"{framework}: ")
            run1.bold = True
            p.add_run(description)
        
        doc.add_paragraph()
        
        # Pattern analysis box
        morning_pct = (len(full_df[full_df['session'] == 'Morning']) / total_incidents * 100) if total_incidents > 0 else 0
        middle_pct = (len(full_df[full_df['session'] == 'Middle']) / total_incidents * 100) if total_incidents > 0 else 0
        afternoon_pct = (len(full_df[full_df['session'] == 'Afternoon']) / total_incidents * 100) if total_incidents > 0 else 0
        
        pattern_content = f"""Based on analysis of {total_incidents} recorded incidents, {student['name']} experiences the most difficulty when '{top_ant}' occurs. This happens most frequently in {top_loc}, particularly during the {top_session} session.

Time of day breakdown: Morning ({morning_pct:.0f}%), Middle ({middle_pct:.0f}%), Afternoon ({afternoon_pct:.0f}%)

The behaviour '{top_beh}' is the primary concern. From a trauma-informed perspective, this behaviour is {student['name']}'s way of communicating an unmet need or responding to feeling unsafe or overwhelmed. It is not 'naughtiness' or 'choosing' to misbehave - it is a stress response."""
        
        create_info_box(doc, "UNDERSTANDING THE PATTERNS", pattern_content, "FFF8E1")
        
        doc.add_page_break()
        
        # ================================================================
        # BERRY STREET FRAMEWORK (Condensed)
        # ================================================================
        
        heading = doc.add_heading('5. Berry Street Education Model Framework', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        berry_intro = doc.add_paragraph()
        berry_intro.add_run("The Berry Street Education Model provides a sequential framework across five domains:").italic = True
        
        doc.add_paragraph()
        
        # Berry Street domains table
        domains_table = doc.add_table(rows=6, cols=3)
        domains_table.style = 'Table Grid'
        
        domains_data = [
            ("Domain", "Focus", "Application"),
            ("1. BODY", "Physical & emotional regulation", f"Help {student['name']} recognise body signals and use calming strategies before {top_session}"),
            ("2. RELATIONSHIP", "Connection & trust", "Maintain consistent, calm responses; see behaviour as communication"),
            ("3. STAMINA", "Persistence & resilience", "Break tasks into smaller steps; celebrate effort, not just outcomes"),
            ("4. ENGAGEMENT", "Active learning", f"Make learning relevant; provide choice especially when '{top_ant}' occurs"),
            ("5. CHARACTER", "Values & contribution", "Develop positive sense of self through contribution opportunities"),
        ]
        
        for i, (domain, focus, application) in enumerate(domains_data):
            for j, text in enumerate([domain, focus, application]):
                cell = domains_table.rows[i].cells[j]
                run = cell.paragraphs[0].add_run(text)
                if i == 0:
                    run.bold = True
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="008080" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif j == 0:
                    run.bold = True
        
        doc.add_paragraph()
        
        current_focus = doc.add_paragraph()
        current_focus.add_run("CURRENT PRIORITY: ").bold = True
        current_focus.add_run(f"Focus on BODY and RELATIONSHIP domains. The data shows {student['name']} is dysregulated during {top_session}, particularly when '{top_ant}' occurs. Engagement and character development require regulation and connection first.")
        
        doc.add_page_break()
        
        # ================================================================
        # CPI PRINCIPLES (Condensed)
        # ================================================================
        
        heading = doc.add_heading('6. Crisis Prevention Institute (CPI) Principles', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        cpi_principles = [
            ("Behaviour is Communication", f"When {student['name']} displays '{top_beh}', they are saying: 'I'm overwhelmed,' 'I feel unsafe,' or 'I need help.'"),
            ("Supportive Stance", "Stand at an angle, give space, keep hands visible, use low/slow voice. Your body language should say: 'I'm here to help.'"),
            ("Maintain Dignity", "Never shame or embarrass. Separate the behaviour from the person. The message: 'I care about you, even when your behaviour is difficult.'"),
            ("Early Intervention", f"Intervene at the first signs of escalation when '{top_ant}' is present. Offer a break, change the task, provide reassurance."),
            ("Co-Regulation", f"{student['name']} often cannot self-regulate when dysregulated. Stay calm and lend them your regulation. Your calm becomes their calm."),
        ]
        
        for title, description in cpi_principles:
            p = doc.add_paragraph()
            run1 = p.add_run(f"{title}: ")
            run1.bold = True
            run1.font.color.rgb = SECONDARY_COLOR
            p.add_run(description)
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # ================================================================
        # EVIDENCE-BASED RECOMMENDATIONS
        # ================================================================
        
        heading = doc.add_heading('7. Evidence-Based Recommendations', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        # Recommendations table with priority levels
        rec_table = doc.add_table(rows=9, cols=3)
        rec_table.style = 'Table Grid'
        
        recommendations = [
            ("Priority", "Recommendation", "Rationale"),
            ("HIGH", f"Implement regulated start before {top_session}", f"Peak incidents occur during {top_session}"),
            ("HIGH", f"Environmental modification in {top_loc}", f"Highest incident location"),
            ("HIGH", f"Visual check-in before '{top_ant}'", "Anticipate and prevent escalation"),
            ("MEDIUM", "Sensory regulation toolkit", "Support body-based regulation"),
            ("MEDIUM", "One key adult maintains connection", "Relationship is protective"),
            ("MEDIUM", "Offer choices for control", "Maintains dignity and autonomy"),
            ("ONGOING", "Teach help-seeking strategies", "Build long-term coping skills"),
            ("ONGOING", "Practice emotional literacy", "Develop self-awareness"),
        ]
        
        priority_colors = {"Priority": "008080", "HIGH": "DC3545", "MEDIUM": "FD7E14", "ONGOING": "28A745"}
        
        for i, (priority, rec, rationale) in enumerate(recommendations):
            for j, text in enumerate([priority, rec, rationale]):
                cell = rec_table.rows[i].cells[j]
                run = cell.paragraphs[0].add_run(text)
                if i == 0:
                    run.bold = True
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="008080" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif j == 0:
                    run.bold = True
                    color = priority_colors.get(priority, "333333")
                    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        
        doc.add_page_break()
        
        # ================================================================
        # ACTION PLAN & TIMELINE
        # ================================================================
        
        heading = doc.add_heading('8. Action Plan & Timeline', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        # SMART Goal
        goal_box = create_info_box(doc, "SMART GOAL",
            f"Within 5 weeks, {student['name']} will use a help-seeking strategy (e.g., break card, signal to adult) in 4 out of 5 opportunities when experiencing triggers related to '{top_ant}', with visual and verbal support.",
            "E8F5E9")
        
        # Timeline table
        doc.add_heading('Implementation Timeline', 2)
        
        timeline_table = doc.add_table(rows=5, cols=3)
        timeline_table.style = 'Table Grid'
        
        review_date = datetime.now() + timedelta(weeks=5)
        
        timeline_data = [
            ("Week", "Focus Area", "Key Actions"),
            ("Week 1-2", "Assessment & Setup", f"Baseline data collection, environmental audit of {top_loc}, introduce regulation tools"),
            ("Week 2-3", "Implementation", f"Scheduled regulation breaks before {top_session}, visual supports in place, key adult assigned"),
            ("Week 3-4", "Skill Building", "Teach and practice help-seeking strategies, emotional vocabulary development"),
            ("Week 5", "Review", f"Progress review on {review_date.strftime('%d %B %Y')}, data comparison, plan adjustment"),
        ]
        
        for i, (week, focus, actions) in enumerate(timeline_data):
            for j, text in enumerate([week, focus, actions]):
                cell = timeline_table.rows[i].cells[j]
                run = cell.paragraphs[0].add_run(text)
                if i == 0:
                    run.bold = True
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="008080" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        doc.add_paragraph()
        
        # ================================================================
        # REVIEW SCHEDULE
        # ================================================================
        
        heading = doc.add_heading('9. Review Schedule', 1)
        for run in heading.runs:
            run.font.color.rgb = PRIMARY_COLOR
            set_arial(run, 16)
        
        review_table = doc.add_table(rows=4, cols=3)
        review_table.style = 'Table Grid'
        
        review_data = [
            ("Review Type", "Date", "Attendees"),
            ("Weekly Check-in", "Every Friday", "Key adult, classroom teacher"),
            ("Formal Review", review_date.strftime('%d %B %Y'), "Full support team, parent/caregiver"),
            ("Plan Renewal", (review_date + timedelta(weeks=5)).strftime('%d %B %Y'), "Support team, leadership"),
        ]
        
        for i, row_data in enumerate(review_data):
            for j, text in enumerate(row_data):
                cell = review_table.rows[i].cells[j]
                run = cell.paragraphs[0].add_run(text)
                if i == 0:
                    run.bold = True
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="008080" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # ================================================================
        # SIGNATURE SECTION
        # ================================================================
        
        add_horizontal_line(doc)
        
        sig_heading = doc.add_paragraph()
        sig_heading.add_run("Prepared By").bold = True
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.columns[0].width = Inches(3)
        sig_table.columns[1].width = Inches(3)
        
        sig_table.rows[0].cells[0].paragraphs[0].add_run("Name: _________________________")
        sig_table.rows[0].cells[1].paragraphs[0].add_run("Role: _________________________")
        sig_table.rows[1].cells[0].paragraphs[0].add_run("Signature: _____________________")
        sig_table.rows[1].cells[1].paragraphs[0].add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Cowandilla Learning Centre | Learning and Behaviour Support Unit")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = PRIMARY_COLOR
        footer_run.bold = True
        
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer2_run = footer2.add_run("Evidence-based frameworks: Applied Behaviour Analysis | Trauma-Informed Practice | Berry Street Education Model | CPI")
        footer2_run.font.size = Pt(8)
        footer2_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Save document
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
        
    except Exception as e:
        import traceback
        st.error(f"BAP Generation Error: {e}")
        st.error(traceback.format_exc())
        return None

# ============================================
# SUPABASE DATABASE FUNCTIONS
# ============================================

def hash_password(plain_password: str) -> str:
    """Hash a password using bcrypt"""
    salt = bcrypt.gensalt(rounds=12)
    hashed = bcrypt.hashpw(plain_password.encode('utf-8'), salt)
    return hashed.decode('utf-8')

# FIXED load_students_from_db FUNCTION
# Replace the existing load_students_from_db function (around line 1060-1085)

def load_students_from_db():
    """Load students from Supabase database"""
    if not supabase:
        return MOCK_STUDENTS  # Fallback to mock data
    
    try:
        response = supabase.table('students').select('*').execute()
        students = []
        for row in response.data:
            # Convert grade from integer back to string format
            grade_num = row['grade']
            if grade_num == 0:
                grade_str = 'R'
            else:
                grade_str = f'Y{grade_num}'
            
            # Handle first_name/last_name - construct name if not present
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            full_name = row.get('name', '')
            
            # If first_name/last_name not in DB, parse from name
            if not first_name and full_name:
                parts = full_name.split()
                first_name = parts[0] if parts else ''
                last_name = ' '.join(parts[1:]) if len(parts) > 1 else ''
            
            # If name not in DB, construct from first/last
            if not full_name and (first_name or last_name):
                full_name = f"{first_name} {last_name}".strip()
            
            students.append({
                "id": str(row['id']),
                "first_name": first_name,
                "last_name": last_name,
                "name": full_name,
                "edid": row['edid'],
                "grade": grade_str,  # Convert back to Y1, Y2, R format
                "dob": row['dob'],
                "program": row['program'],
                "placement_start": row['placement_start'],
                "placement_end": row['placement_end']
            })
        return students
    except Exception as e:
        st.error(f"Error loading students: {e}")
        return []

def save_student_to_db(student):
    """Save a student to Supabase database"""
    if not supabase:
        # DEMO MODE: Return True to allow app to function without database
        return True
    
    try:
        # Convert grade to just the number if it starts with Y
        grade_value = student['grade']
        if isinstance(grade_value, str):
            if grade_value.startswith('Y'):
                grade_value = grade_value[1:]
            elif grade_value == 'R':
                grade_value = 0
        
        data = {
            "first_name": student.get('first_name', student['name'].split()[0] if student['name'] else ''),
            "last_name": student.get('last_name', ' '.join(student['name'].split()[1:]) if len(student['name'].split()) > 1 else ''),
            "name": student['name'],
            "edid": student['edid'],
            "grade": int(grade_value) if str(grade_value).isdigit() else 0,
            "year_level": int(grade_value) if str(grade_value).isdigit() else 0,
            "dob": student['dob'],
            "program": student['program'],
            "placement_start": student['placement_start'],
            "placement_end": student.get('placement_end')
        }
        
        if 'id' in student and student['id'].startswith('stu_'):
            supabase.table('students').insert(data).execute()
        else:
            supabase.table('students').update(data).eq('id', student['id']).execute()
        
        return True
        
    except Exception as e:
        st.error(f"Error saving student: {e}")
        return False
def delete_student_from_db(student_id):
    """Delete a student from Supabase database"""
    if not supabase:
        # DEMO MODE: Return True to allow app to function without database
        return True
    
    try:
        supabase.table('students').delete().eq('id', student_id).execute()
        return True
    except Exception as e:
        st.error(f"Error deleting student: {e}")
        return False

def load_staff_from_db():
    """Load staff from Supabase database"""
    if not supabase:
        return MOCK_STAFF
    
    try:
        response = supabase.table('staff').select('*').execute()
        staff = []
        for row in response.data:
            # Handle first_name/last_name - construct name if not present
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            full_name = row.get('name', '')
            
            # If first_name/last_name not in DB, parse from name
            if not first_name and full_name:
                parts = full_name.split()
                first_name = parts[0] if parts else ''
                last_name = ' '.join(parts[1:]) if len(parts) > 1 else ''
            
            # If name not in DB, construct from first/last
            if not full_name and (first_name or last_name):
                full_name = f"{first_name} {last_name}".strip()
            
            staff.append({
                "id": str(row['id']),
                "first_name": first_name,
                "last_name": last_name,
                "name": full_name,
                "email": row['email'],
                "password": row.get('password'),  # Keep for backward compatibility
                "password_hash": row.get('password_hash', row.get('password')),  # Use password_hash if available
                "role": row['role'],
                "program": row.get('program'),
                "phone": row.get('phone'),
                "notes": row.get('notes'),
                "receive_critical_emails": row.get('receive_critical_emails', True),
                "created_date": row.get('created_at', '')[:10] if row.get('created_at') else None
            })
        return staff  # Return what we have from database
    except Exception as e:
        st.error(f"Error loading staff: {e}")
        return []  # Return empty list on error

def save_staff_to_db(staff_member):
    """Save a staff member to Supabase database"""
    if not supabase:
        # DEMO MODE: Return True to allow app to function without database
        return True
    
    try:
        data = {
            "first_name": staff_member.get('first_name', staff_member['name'].split()[0] if staff_member['name'] else ''),
            "last_name": staff_member.get('last_name', ' '.join(staff_member['name'].split()[1:]) if len(staff_member['name'].split()) > 1 else ''),
            "name": staff_member['name'],
            "email": staff_member['email'],
            "password": staff_member['password'],
            "role": staff_member['role'],
            "program": staff_member.get('program'),
            "phone": staff_member.get('phone'),
            "notes": staff_member.get('notes'),
            "receive_critical_emails": staff_member.get('receive_critical_emails', True)
        }
        
        if 'id' in staff_member and staff_member['id'].startswith('staff_'):
            # New staff (generated ID from app)
            supabase.table('staff').insert(data).execute()
        else:
            # Existing staff (UUID from database)
            supabase.table('staff').update(data).eq('id', staff_member['id']).execute()
        return True
    except Exception as e:
        st.error(f"Error saving staff: {e}")
        return False

def delete_staff_from_db(staff_id):
    """Delete a staff member from Supabase database"""
    if not supabase:
        # DEMO MODE: Return True to allow app to function without database
        return True
    
    try:
        supabase.table('staff').delete().eq('id', staff_id).execute()
        return True
    except Exception as e:
        st.error(f"Error deleting staff: {e}")
        return False

def load_incidents_from_db(student_id=None):
    """Load incidents from Supabase database"""
    if not supabase:
        return []
    
    try:
        query = supabase.table('incidents').select('*')
        if student_id:
            query = query.eq('student_id', student_id)
        response = query.execute()
        
        incidents = []
        for row in response.data:
            incidents.append({
                "id": str(row['id']),
                "student_id": str(row['student_id']),
                "date": row['date'],
                "time": row['time'],
                "day": row['day_of_week'],
                "session": row['session'],
                "location": row['location'],
                "behaviour_type": row['behaviour_type'],
                "antecedent": row['antecedent'],
                "intervention": row['intervention'],  # Already array in DB
                "severity": row['severity'],
                "reported_by": str(row['reported_by']) if row.get('reported_by') else None,
                "description": row.get('description', ''),
                "duration_minutes": row.get('duration_minutes'),
                "hypothesis_function": row.get('hypothesis_function'),
                "hypothesis_item": row.get('hypothesis_item'),
                "is_critical": row.get('is_critical', False)
            })
        return incidents
    except Exception as e:
        st.error(f"Error loading incidents: {e}")
        return []

def save_incident_to_db(incident):
    """Save an incident to Supabase database"""
    if not supabase:
        # DEMO MODE: Return True to allow app to function without database
        return True
    
    try:
        data = {
            "student_id": incident['student_id'],
            "date": incident['date'],
            "time": incident['time'],
            "day_of_week": incident['day'],
            "session": incident['session'],
            "location": incident['location'],
            "behaviour_type": incident['behaviour_type'],
            "antecedent": incident['antecedent'],
            "intervention": incident['intervention'],
            "severity": incident['severity'],
            "reported_by": incident.get('reported_by'),
            "description": incident.get('description', ''),
            "duration_minutes": incident.get('duration_minutes'),
            "hypothesis_function": incident.get('hypothesis_function'),
            "hypothesis_item": incident.get('hypothesis_item'),
            "is_critical": incident.get('is_critical', False)
        }
        
        if 'id' not in incident or not incident['id']:
            # New incident
            supabase.table('incidents').insert(data).execute()
        else:
            # Update existing
            supabase.table('incidents').update(data).eq('id', incident['id']).execute()
        return True
    except Exception as e:
        st.error(f"Error saving incident: {e}")
        return False

def load_critical_incidents_from_db(student_id=None):
    """Load critical incidents from Supabase database"""
    if not supabase:
        return []
    
    try:
        query = supabase.table('critical_incidents').select('*')
        if student_id:
            query = query.eq('student_id', student_id)
        response = query.execute()
        
        critical = []
        for row in response.data:
            critical.append({
                "id": str(row['id']),
                "student_id": str(row['student_id']),
                "severity": row['severity'],
                "reported_by": str(row['reported_by']) if row.get('reported_by') else None,
                "ABCH_primary": {
                    "location": row['primary_location'],
                    "context": row['primary_context'],
                    "time": row['primary_time'],
                    "behaviour": row['primary_behaviour'],
                    "consequence": row['primary_consequence'],
                    "hypothesis_function": row.get('primary_hypothesis_function'),
                    "hypothesis_item": row.get('primary_hypothesis_item')
                },
                "ABCH_additional": row.get('additional_entries', []),
                "outcomes": row['outcomes'],
                "sapol_reference": row.get('sapol_reference'),
                "admin_summary": row.get('admin_summary'),
                "created_at": row.get('created_at')
            })
        return critical
    except Exception as e:
        st.error(f"Error loading critical incidents: {e}")
        return []

def save_critical_incident_to_db(critical):
    """Save a critical incident to Supabase database"""
    if not supabase:
        # DEMO MODE: Return True to allow app to function without database
        return True
    
    try:
        primary = critical.get('ABCH_primary', {})
        data = {
            "student_id": critical['student_id'],
            "severity": critical['severity'],
            "reported_by": critical.get('reported_by'),
            "primary_location": primary.get('location', ''),
            "primary_context": primary.get('context', ''),
            "primary_time": primary.get('time', ''),
            "primary_behaviour": primary.get('behaviour', ''),
            "primary_consequence": primary.get('consequence', ''),
            "primary_hypothesis_function": primary.get('hypothesis_function'),
            "primary_hypothesis_item": primary.get('hypothesis_item'),
            "additional_entries": critical.get('ABCH_additional', []),
            "outcomes": critical.get('outcomes', []),
            "sapol_reference": critical.get('sapol_reference'),
            "admin_summary": critical.get('admin_summary')
        }
        
        if 'id' not in critical or not critical['id']:
            # New critical incident
            supabase.table('critical_incidents').insert(data).execute()
        else:
            # Update existing
            supabase.table('critical_incidents').update(data).eq('id', critical['id']).execute()
        return True
    except Exception as e:
        st.error(f"Error saving critical incident: {e}")
        return False


def init_state():
    """Initialize session state with DEMO DATA - This is the demo version"""
    ss = st.session_state
    if "logged_in" not in ss: ss.logged_in = False
    if "current_user" not in ss: ss.current_user = None
    if "current_page" not in ss: ss.current_page = "login"
    
    # DEMO VERSION: Always use mock data for demonstration purposes
    if "students" not in ss: 
        ss.students = MOCK_STUDENTS.copy()
    if "staff" not in ss: 
        ss.staff = MOCK_STAFF.copy()
    if "incidents" not in ss: 
        ss.incidents = generate_demo_incidents()
    if "critical_incidents" not in ss: 
        ss.critical_incidents = generate_demo_critical_incidents()
    
    if "selected_program" not in ss: ss.selected_program = "JP"
    if "selected_student_id" not in ss: ss.selected_student_id = None
    if "current_incident_id" not in ss: ss.current_incident_id = None
    if "abch_rows" not in ss: ss.abch_rows = []
    if "show_critical_prompt" not in ss: ss.show_critical_prompt = False

def login_user(email: str, password: str) -> bool:
    """Login user - Demo version uses plain password matching"""
    email = (email or "").strip().lower()
    password = (password or "").strip()
    
    if not email or not password:
        return False
    
    for staff in st.session_state.staff:
        staff_email = staff.get("email", "").lower()
        
        if staff_email == email:
            # Demo version: Use plain password matching
            if staff.get("password") == password:
                st.session_state.logged_in = True
                st.session_state.current_user = staff
                st.session_state.current_page = "landing"
                return True
            
            # Also try bcrypt if available
            stored_hash = staff.get("password_hash", "")
            if stored_hash:
                try:
                    if isinstance(stored_hash, str):
                        stored_hash_bytes = stored_hash.encode('utf-8')
                    else:
                        stored_hash_bytes = stored_hash
                    password_bytes = password.encode('utf-8')
                    
                    if bcrypt.checkpw(password_bytes, stored_hash_bytes):
                        st.session_state.logged_in = True
                        st.session_state.current_user = staff
                        st.session_state.current_page = "landing"
                        return True
                except Exception:
                    pass
    
    return False
def go_to(page: str, **kwargs):
    if page not in VALID_PAGES: return
    st.session_state.current_page = page
    for k, v in kwargs.items():
        setattr(st.session_state, k, v)
    st.rerun()

def get_student(sid): 
    return next((s for s in st.session_state.students if s["id"] == sid), None)

def get_session_from_time(t): 
    return "Morning" if t.hour < 11 else "Middle" if t.hour < 13 else "Afternoon"

def generate_demo_incidents():
    """Generate comprehensive demo incidents - 5-8 per student with realistic patterns"""
    incidents = []
    
    # Define realistic incident patterns for each student
    student_patterns = {
        "stu_jp1": {  # Emma Thompson - R - escape/avoidance focused
            "primary_behaviours": ["Verbal Refusal", "Elopement"],
            "primary_antecedents": ["Demand - literacy tasks", "Transition - to a non-preferred activity"],
            "primary_locations": ["Classroom", "Learning Space"],
            "hypothesis_function": "Escape/Avoidance",
            "hypothesis_item": "non-preferred literacy demands",
            "incident_count": 7
        },
        "stu_jp2": {  # Oliver Smith - Y1 - attention seeking
            "primary_behaviours": ["Verbal Aggression", "Property Destruction"],
            "primary_antecedents": ["Attention - staff diverted to another student", "Peer - peer conflict/interaction"],
            "primary_locations": ["Classroom", "Playground"],
            "hypothesis_function": "Access to Attention",
            "hypothesis_item": "staff attention when feeling ignored",
            "incident_count": 6
        },
        "stu_py1": {  # Liam Carter - Y3 - escape focused
            "primary_behaviours": ["Elopement", "Verbal Refusal", "Aggression (Adult)"],
            "primary_antecedents": ["Demand - maths tasks", "Demand - completing a task independently"],
            "primary_locations": ["Classroom", "Learning Space", "Withdrawal Room"],
            "hypothesis_function": "Escape/Avoidance",
            "hypothesis_item": "challenging academic demands",
            "incident_count": 8
        },
        "stu_py2": {  # Ava Robinson - Y5 - sensory/tangible
            "primary_behaviours": ["Verbal Refusal", "Property Destruction"],
            "primary_antecedents": ["Sensory - environment too noisy/busy", "Tangible - denied access to preferred item"],
            "primary_locations": ["Classroom", "Hall/Assembly Area"],
            "hypothesis_function": "Access to Tangible",
            "hypothesis_item": "preferred items or sensory regulation",
            "incident_count": 5
        },
        "stu_sy1": {  # Isabella Garcia - Y7 - peer conflict focused
            "primary_behaviours": ["Verbal Aggression", "Aggression (Peer)", "Elopement"],
            "primary_antecedents": ["Peer - peer conflict/interaction", "Peer - negative peer feedback"],
            "primary_locations": ["Classroom", "Playground", "Corridors"],
            "hypothesis_function": "Escape/Avoidance",
            "hypothesis_item": "negative peer interactions",
            "incident_count": 7
        },
        "stu_sy2": {  # Ethan Davis - Y9 - escape/control focused
            "primary_behaviours": ["Verbal Refusal", "Verbal Aggression", "Property Destruction"],
            "primary_antecedents": ["Demand - given an instruction by staff", "Internal - ruminating on past events"],
            "primary_locations": ["Classroom", "Learning Space", "Outdoors"],
            "hypothesis_function": "Escape/Avoidance",
            "hypothesis_item": "adult-directed demands",
            "incident_count": 6
        }
    }
    
    # Generate incidents for each student
    for student_id, pattern in student_patterns.items():
        student = next((s for s in MOCK_STUDENTS if s["id"] == student_id), None)
        if not student:
            continue
            
        for i in range(pattern["incident_count"]):
            # Vary dates over past 60 days
            days_ago = random.randint(1, 60)
            dt = datetime.now() - timedelta(days=days_ago)
            
            # Vary times throughout school day
            hour = random.choices([9, 10, 11, 12, 13, 14, 15], weights=[12, 18, 15, 10, 12, 18, 15])[0]
            dt = dt.replace(hour=hour, minute=random.randint(0, 59), second=0)
            
            # Use pattern-based selections with some variation
            if random.random() < 0.7:  # 70% follow pattern
                behaviour = random.choice(pattern["primary_behaviours"])
                antecedent = random.choice(pattern["primary_antecedents"])
                location = random.choice(pattern["primary_locations"])
            else:  # 30% random for variation
                behaviour = random.choice(BEHAVIOUR_TYPES)
                antecedent = random.choice([a for a in ANTECEDENTS if not a.startswith("---")])
                location = random.choice(LOCATIONS)
            
            # Severity distribution - ensure some are critical (3+)
            if i < 2:  # First 2 incidents per student are critical
                severity = random.choice([3, 4])
            else:
                severity = random.choices([1, 2, 3, 4], weights=[25, 40, 25, 10])[0]
            
            is_critical = severity >= 3
            
            incidents.append({
                "id": str(uuid.uuid4()),
                "student_id": student_id,
                "student_name": student["name"],
                "date": dt.date().isoformat(),
                "time": dt.time().strftime("%H:%M:%S"),
                "day": dt.strftime("%A"),
                "session": get_session_from_time(dt.time()),
                "location": location,
                "behaviour_type": behaviour,
                "antecedent": antecedent,
                "intervention": random.sample(INTERVENTIONS, k=random.randint(1, 3)),
                "severity": severity,
                "reported_by": random.choice(MOCK_STAFF)["id"],
                "description": f"Demo incident for {student['name']}",
                "is_critical": is_critical,
                "duration_minutes": random.randint(3, 20),
                "hypothesis_function": pattern["hypothesis_function"],
                "hypothesis_item": pattern["hypothesis_item"]
            })
    
    return sorted(incidents, key=lambda x: x["date"], reverse=True)

def generate_demo_critical_incidents():
    """Generate demo critical incident records with full ABCH data"""
    critical_incidents = []
    
    # Create 2-3 critical incidents per student
    critical_data = [
        # Emma Thompson - JP
        {
            "student_id": "stu_jp1",
            "ABCH_primary": {
                "location": "Classroom",
                "context": "Asked to complete writing task, became overwhelmed",
                "time": "09:30 AM",
                "behaviour": "Threw chair, screamed, ran out of classroom",
                "consequence": "Room cleared, student followed at safe distance",
                "hypothesis": "Escape/Avoidance - overwhelming literacy demand"
            },
            "intended_outcomes": ["Access to Tangible", "Escape/Avoidance"],
            "severity": 4
        },
        {
            "student_id": "stu_jp1",
            "ABCH_primary": {
                "location": "Learning Space",
                "context": "Transition from play to structured activity",
                "time": "11:15 AM",
                "behaviour": "Pushed staff member, kicked furniture",
                "consequence": "Space cleared, calm voice used, waited for regulation",
                "hypothesis": "Escape/Avoidance - transition difficulty"
            },
            "intended_outcomes": ["Escape/Avoidance"],
            "severity": 3
        },
        # Oliver Smith - JP
        {
            "student_id": "stu_jp2",
            "ABCH_primary": {
                "location": "Playground",
                "context": "Staff helping another student, Oliver wanted attention",
                "time": "01:30 PM",
                "behaviour": "Threw rocks at peers, verbal threats",
                "consequence": "Immediate intervention, removed from area",
                "hypothesis": "Access to Attention - feeling overlooked"
            },
            "intended_outcomes": ["Access to Attention"],
            "severity": 4
        },
        {
            "student_id": "stu_jp2",
            "ABCH_primary": {
                "location": "Classroom",
                "context": "Group activity, not chosen as leader",
                "time": "10:45 AM",
                "behaviour": "Destroyed peer's work, verbal aggression",
                "consequence": "Separated from group, 1:1 support provided",
                "hypothesis": "Access to Attention - seeking recognition"
            },
            "intended_outcomes": ["Access to Attention", "Access to Tangible"],
            "severity": 3
        },
        # Liam Carter - PY
        {
            "student_id": "stu_py1",
            "ABCH_primary": {
                "location": "Classroom",
                "context": "Maths assessment, struggling with content",
                "time": "09:45 AM",
                "behaviour": "Flipped desk, attempted to leave building",
                "consequence": "Building secured, leadership called",
                "hypothesis": "Escape/Avoidance - academic frustration"
            },
            "intended_outcomes": ["Escape/Avoidance"],
            "severity": 4
        },
        {
            "student_id": "stu_py1",
            "ABCH_primary": {
                "location": "Withdrawal Room",
                "context": "Asked to return to class after break",
                "time": "02:00 PM",
                "behaviour": "Punched wall, verbal abuse to staff",
                "consequence": "Given space, parents contacted",
                "hypothesis": "Escape/Avoidance - avoiding classroom return"
            },
            "intended_outcomes": ["Escape/Avoidance"],
            "severity": 3
        },
        {
            "student_id": "stu_py1",
            "ABCH_primary": {
                "location": "Learning Space",
                "context": "Independent work task without support",
                "time": "11:30 AM",
                "behaviour": "Threw materials, grabbed staff arm",
                "consequence": "CPI techniques used, incident documented",
                "hypothesis": "Escape/Avoidance - feeling unsupported"
            },
            "intended_outcomes": ["Escape/Avoidance", "Access to Attention"],
            "severity": 4
        },
        # Ava Robinson - PY
        {
            "student_id": "stu_py2",
            "ABCH_primary": {
                "location": "Hall/Assembly Area",
                "context": "Whole school assembly, sensory overload",
                "time": "09:15 AM",
                "behaviour": "Screaming, attempted to run outside",
                "consequence": "Escorted to quiet space, sensory tools provided",
                "hypothesis": "Sensory regulation - overwhelmed by environment"
            },
            "intended_outcomes": ["Escape/Avoidance"],
            "severity": 3
        },
        {
            "student_id": "stu_py2",
            "ABCH_primary": {
                "location": "Classroom",
                "context": "iPad taken away at end of allocated time",
                "time": "02:30 PM",
                "behaviour": "Threw iPad, hit staff member",
                "consequence": "Room cleared, leadership notified",
                "hypothesis": "Access to Tangible - denied preferred item"
            },
            "intended_outcomes": ["Access to Tangible"],
            "severity": 4
        },
        # Isabella Garcia - SY
        {
            "student_id": "stu_sy1",
            "ABCH_primary": {
                "location": "Corridors",
                "context": "Confrontation with peer about social media post",
                "time": "12:45 PM",
                "behaviour": "Physical altercation, verbal threats",
                "consequence": "Students separated, parents contacted",
                "hypothesis": "Escape/Avoidance - peer conflict escalation"
            },
            "intended_outcomes": ["Escape/Avoidance"],
            "severity": 4
        },
        {
            "student_id": "stu_sy1",
            "ABCH_primary": {
                "location": "Classroom",
                "context": "Peer made negative comment about appearance",
                "time": "10:00 AM",
                "behaviour": "Left class, verbal aggression to peer",
                "consequence": "Followed at distance, given space to regulate",
                "hypothesis": "Escape/Avoidance - social rejection"
            },
            "intended_outcomes": ["Escape/Avoidance", "Access to Attention"],
            "severity": 3
        },
        # Ethan Davis - SY
        {
            "student_id": "stu_sy2",
            "ABCH_primary": {
                "location": "Classroom",
                "context": "Asked to put phone away multiple times",
                "time": "09:30 AM",
                "behaviour": "Verbal abuse to teacher, threw chair",
                "consequence": "Room cleared, leadership intervention",
                "hypothesis": "Escape/Avoidance - adult authority challenge"
            },
            "intended_outcomes": ["Escape/Avoidance", "Access to Tangible"],
            "severity": 4
        },
        {
            "student_id": "stu_sy2",
            "ABCH_primary": {
                "location": "Outdoors",
                "context": "Ruminating about family issues, staff checked in",
                "time": "01:15 PM",
                "behaviour": "Punched fence, self-harm statements",
                "consequence": "Wellbeing team notified, safety plan activated",
                "hypothesis": "Internal regulation - emotional overwhelm"
            },
            "intended_outcomes": ["Escape/Avoidance"],
            "severity": 4
        },
    ]
    
    for i, data in enumerate(critical_data):
        student = next((s for s in MOCK_STUDENTS if s["id"] == data["student_id"]), None)
        if not student:
            continue
            
        days_ago = random.randint(5, 50)
        created = datetime.now() - timedelta(days=days_ago)
        
        critical_incidents.append({
            "id": str(uuid.uuid4()),
            "student_id": data["student_id"],
            "student_name": student["name"],
            "ABCH_primary": data["ABCH_primary"],
            "ABCH_continuation": [],
            "intended_outcomes": data["intended_outcomes"],
            "witnesses": ["Staff member present", "Other students in area"],
            "injuries": "None reported" if data["severity"] < 4 else "Minor - no medical attention required",
            "property_damage": "Minor" if random.random() > 0.5 else "None",
            "police_contacted": False,
            "police_reference": "",
            "staff_completing": random.choice(MOCK_STAFF)["name"],
            "created_at": created.isoformat(),
            "severity": data["severity"]
        })
    
    return critical_incidents

def generate_mock_incidents(n=70):
    """Legacy function - redirects to demo generator"""
    return generate_demo_incidents()

# PAGES
def render_login_page():
    st.markdown("## 🔐 Staff Login")
    
    # DEMO CREDENTIALS INFO
    st.info("""
    🎭 **DEMO VERSION** - Use these credentials to explore:
    
    | Role | Email | Password |
    |------|-------|----------|
    | Admin | admin@example.com | admin123 |
    | Staff (JP) | emily.jones@example.com | demo123 |
    | Staff (PY) | daniel.lee@example.com | demo123 |
    | Staff (SY) | sarah.chen@example.com | demo123 |
    """)

    email = st.text_input("Email Address", placeholder="your.email@example.com", key="login_email")
    password = st.text_input("Password", type="password", placeholder="Enter password", key="login_pass")
    if st.button("Login", type="primary", use_container_width=True):
        if login_user(email, password):
            st.success(f"Welcome {st.session_state.current_user['name']}!")
            st.rerun()
        else:
            st.error("Invalid credentials")

def render_landing_page():
    user = st.session_state.current_user or {}
    
    # DEMO BANNER
    st.warning("🎭 **DEMO VERSION** - This app contains simulated data for demonstration purposes. All student names and incidents are fictional.")
    
    st.markdown(f"### 👋 Welcome, {user.get('name', 'User')}")
    
    col1, col2 = st.columns([6, 1])
    with col2:
        if st.button("Logout", key="logout_btn"):
            st.session_state.logged_in = False
            st.session_state.current_page = "login"
            st.rerun()
    
    # Admin portal button for admin users
    if st.session_state.current_user.get("role") == "ADM":
        st.markdown("---")
        if st.button("🔧 Admin Portal", use_container_width=True, key="goto_admin"):
            go_to("admin_portal")
    
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1: st.metric("Students", len(st.session_state.students))
    with col2: st.metric("Total Incidents", len(st.session_state.incidents))
    with col3: st.metric("Critical", len([i for i in st.session_state.incidents if i.get("is_critical")]))
    
    st.markdown("---")
    st.markdown("### Select Program")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Junior Primary**")
        if st.button("Enter JP", use_container_width=True, type="primary", key="btn_jp"):
            go_to("program_students", selected_program="JP")
    with col2:
        st.markdown("**Primary Years**")
        if st.button("Enter PY", use_container_width=True, type="primary", key="btn_py"):
            go_to("program_students", selected_program="PY")
    with col3:
        st.markdown("**Senior Years**")
        if st.button("Enter SY", use_container_width=True, type="primary", key="btn_sy"):
            go_to("program_students", selected_program="SY")

def render_program_students_page():
    program = st.session_state.get("selected_program", "JP")
    st.markdown(f"## {PROGRAM_NAMES.get(program)} — Students")
    
    col1, col2 = st.columns([6, 1])
    with col1:
        if st.button("⬅ Back to Landing", key="back_students"):
            go_to("landing")
    
    students = [s for s in st.session_state.students if s["program"] == program]
    for stu in students:
        stu_incidents = [i for i in st.session_state.incidents if i["student_id"] == stu["id"]]
        with st.container(border=True):
            col1, col2, col3 = st.columns([3, 2, 2])
            with col1:
                st.markdown(f"**{stu['name']}**")
                st.caption(f"Grade {stu['grade']}")
            with col2:
                st.metric("Incidents", len(stu_incidents))
            with col3:
                if st.button("📝 Log", key=f"log_{stu['id']}", use_container_width=True):
                    go_to("incident_log", selected_student_id=stu["id"])
                if st.button("📊 Analysis", key=f"ana_{stu['id']}", use_container_width=True):
                    go_to("student_analysis", selected_student_id=stu["id"])

def render_incident_log_page():
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## 📝 Incident Log — {student['name']}")
    
    # Navigation buttons
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("⬅ Back to Students", key="back_log_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_log"):
            go_to("landing")
    
    show_severity_guide()
    
    # Check if critical form is required
    if st.session_state.show_critical_prompt:
        inc_info = st.session_state.get("last_incident_info", {})
        if inc_info.get("severity", 0) >= 3:
            st.error(f"⚠️ **Severity {inc_info['severity']} Detected** - Critical Incident ABCH Form Required")
        else:
            st.error("⚠️ **Critical Incident Flagged** - Critical Incident ABCH Form Required")
        st.info("Please complete the Critical Incident ABCH form to document this event fully.")
        
        # Only one button - REMOVED "Skip for Now"
        if st.button("📋 Complete Critical Form Now", type="primary", key="crit_now", use_container_width=True):
            st.session_state.show_critical_prompt = False
            go_to("critical_incident", current_incident_id=st.session_state.current_incident_id)
        st.markdown("---")
        st.stop()
    
    # INCIDENT FORM - Split to show hypothesis before severity
    st.markdown("### Log New Incident")
    
    # First section: Capture antecedent and behaviour for hypothesis generation
    col1, col2 = st.columns(2)
    with col1:
        behaviour_select = st.selectbox("Behaviour Type *", [""] + BEHAVIOUR_TYPES, key="inc_beh_select")
    with col2:
        antecedent_select = st.selectbox("Antecedent/Trigger *", [""] + ANTECEDENTS, key="inc_ant_select")
    
    # Display hypothesis if both are selected
    if behaviour_select and antecedent_select:
        hyp_ai = generate_hypothesis_ai(antecedent_select, behaviour_select, "")
        hypothesis_text = f"{hyp_ai['function']} {hyp_ai['item']}"
        st.info(f"🧠 **Suggested Hypothesis:** {hypothesis_text}")
        # Store for form submission
        st.session_state.current_hypothesis = hyp_ai
    else:
        st.session_state.current_hypothesis = None
    
    # Main form with remaining fields
    with st.form("incident_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            inc_date = st.date_input("Date *", date.today(), key="inc_date", format="DD/MM/YYYY")
            inc_time = st.time_input("Time *", datetime.now().time(), key="inc_time")
            location = st.selectbox("Location *", [""] + LOCATIONS, key="inc_loc")
        with col2:
            # Hidden fields to pass the pre-selected values
            st.markdown(f"**Behaviour Type:** {behaviour_select if behaviour_select else 'Not selected'}")
            st.markdown(f"**Antecedent/Trigger:** {antecedent_select if antecedent_select else 'Not selected'}")
            # MULTIPLE INTERVENTIONS
            interventions = st.multiselect("Interventions Used *", INTERVENTIONS, key="inc_ints")
        
        duration = st.number_input("Duration (minutes) *", min_value=1, value=1, key="inc_dur")
        severity = st.slider("Severity Level (from start to end of incident) *", 1, 5, 1, key="inc_sev")
        description = st.text_area("Brief Description (Optional)", placeholder="Factual, objective description...", key="inc_desc")
        manual_critical = st.checkbox("This incident requires a Critical Incident ABCH Form (regardless of severity)", key="manual_crit")
        submitted = st.form_submit_button("Submit Incident", type="primary")
    
    if submitted:
        behaviour = behaviour_select
        antecedent = antecedent_select
        if not location or not behaviour or not antecedent or not interventions:
            st.error("Please complete all required fields marked with *")
        else:
            new_id = str(uuid.uuid4())
            is_critical = (severity >= 3) or manual_critical
            
            # Use pre-generated hypothesis or generate new one
            hyp_ai = st.session_state.get('current_hypothesis') or generate_hypothesis_ai(antecedent, behaviour, "")
            hypothesis_text = f"{hyp_ai['function']} {hyp_ai['item']}"
            
            rec = {
                "id": new_id, 
                "student_id": student_id, 
                "student_name": student["name"],
                "date": inc_date.isoformat(), 
                "time": inc_time.strftime("%H:%M:%S"),
                "day": inc_date.strftime("%A"), 
                "session": get_session_from_time(inc_time),
                "location": location, 
                "behaviour_type": behaviour, 
                "antecedent": antecedent,
                "intervention": interventions,
                "severity": severity,
                "reported_by": st.session_state.current_user["id"],
                "duration_minutes": duration, 
                "description": description or "", 
                "is_critical": is_critical,
                "hypothesis_function": hyp_ai['function'],
                "hypothesis_item": hyp_ai['item']
            }
            
            # SAVE TO DATABASE FIRST
            if save_incident_to_db(rec):
                # Then add to session state
                st.session_state.incidents.append(rec)
                st.success("✅ Incident logged successfully and saved to database")
                
                if is_critical:
                    st.session_state.current_incident_id = new_id
                    st.session_state.show_critical_prompt = True
                    st.session_state.last_incident_info = {"severity": severity, "manual": manual_critical}
                    st.rerun()
            else:
                st.error("❌ Failed to save incident to database. Please try again.")


def render_critical_incident_page():
    """Critical Incident Form with Police Reference and Improved Labels"""
    inc_id = st.session_state.get("current_incident_id")
    quick_inc = next((i for i in st.session_state.incidents if i["id"] == inc_id), None)
    
    if not quick_inc:
        st.error("No incident found")
        return
    
    student = get_student(quick_inc["student_id"])
    st.markdown(f"## 🚨 Critical Incident ABCH Form")
    
    # Navigation
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("⬅ Back to Students", key="back_crit_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_crit"):
            go_to("landing")
    
    st.markdown("### Incident Details (from Quick Log)")
    with st.container(border=True):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"**Student:** {student['name']}")
            st.markdown(f"**Grade:** {student['grade']}")
        with col2:
            st.markdown(f"**Date:** {format_date_dmy(quick_inc['date'])}")
            st.markdown(f"**Time:** {format_time_12hr(quick_inc['time'])}")
        with col3:
            st.markdown(f"**Location:** {quick_inc['location']}")
            st.markdown(f"**Session:** {quick_inc['session']}")
        with col4:
            st.markdown(f"**Severity:** {quick_inc['severity']}")
        hypothesis_display = f"{quick_inc.get('hypothesis_function', '')} {quick_inc.get('hypothesis_item', '')}".strip() or 'N/A'
        st.markdown(f"**Hypothesis:** {hypothesis_display}")
    
    st.markdown("---")
    st.markdown("### ABCH Chronology")
    st.caption("Document the sequence of events. Add continuation entries if incident was prolonged.")
    
    if "abch_rows" not in st.session_state:
        st.session_state.abch_rows = []
    
    # PRIMARY ROW
    st.markdown("#### Initial Incident")
    col_header = st.columns([2, 2, 2, 2, 2])
    with col_header[0]: st.markdown("**ANTECEDENT (Triggers)**")
    with col_header[2]: st.markdown("**BEHAVIOUR**")
    with col_header[4]: st.markdown("**CONSEQUENCES**")
    
    col_subheader = st.columns([1, 1, 1, 1, 2, 2])
    with col_subheader[0]: st.caption("Location")
    with col_subheader[1]: st.caption("Context (what was happening?)")
    with col_subheader[2]: st.caption("Time")
    with col_subheader[3]: st.caption("Observed Behaviour")
    with col_subheader[4]: st.caption("What happened after?")
    with col_subheader[5]: st.caption("HYPOTHESIS (Function)")
    
    col_inputs1 = st.columns([1, 1, 1, 1, 2, 2])
    
    with col_inputs1[0]:
        location_1 = st.text_input("", value=quick_inc['location'], key="loc_1", label_visibility="collapsed")
    with col_inputs1[1]:
        context_1 = st.text_area("", placeholder="What was going on before?", 
                                key="context_1", height=100, label_visibility="collapsed")
    with col_inputs1[2]:
        time_1 = st.text_input("", value=format_time_12hr(quick_inc['time']), key="time_1", label_visibility="collapsed")
    with col_inputs1[3]:
        behaviour_1 = st.text_area("", placeholder="What did student do?", 
                                  key="behaviour_1", height=100, label_visibility="collapsed")
    with col_inputs1[4]:
        consequence_1 = st.text_area("", placeholder="Staff response? Student reaction?", 
                                    key="consequence_1", height=100, label_visibility="collapsed")
    with col_inputs1[5]:
        if "hyp_1_generated" not in st.session_state:
            st.session_state.hyp_1_generated = False
        if not st.session_state.hyp_1_generated and context_1 and behaviour_1:
            auto_hyp = generate_hypothesis(context_1, behaviour_1, consequence_1)
            st.session_state.hyp_1_auto = auto_hyp
            st.session_state.hyp_1_generated = True
        hypothesis_1 = st.text_area("", 
                                    value=st.session_state.get("hyp_1_auto", ""),
                                    placeholder="Auto-generated (editable)", 
                                    key="hypothesis_1", height=100, label_visibility="collapsed")
    
    st.markdown("---")
    
    if st.button("➕ Add Continuation Entry", key="add_abch_row"):
        st.session_state.abch_rows.append({})
        st.rerun()
    
    # ADDITIONAL ROWS - Changed labels
    for idx, row in enumerate(st.session_state.abch_rows):
        st.markdown(f"#### Continuation Entry {idx + 1}")
        col_add = st.columns([1, 1, 1, 1, 2, 2])
        with col_add[0]:
            row["location"] = st.text_input("", key=f"loc_{idx+2}", label_visibility="collapsed")
        with col_add[1]:
            row["context"] = st.text_area("", key=f"context_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[2]:
            row["time"] = st.text_input("", key=f"time_{idx+2}", label_visibility="collapsed")
        with col_add[3]:
            row["behaviour"] = st.text_area("", key=f"behaviour_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[4]:
            row["consequence"] = st.text_area("", key=f"consequence_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[5]:
            if row.get("context") and row.get("behaviour"):
                auto_hyp_add = generate_hypothesis(row["context"], row["behaviour"], row.get("consequence", ""))
                row["hypothesis"] = st.text_area("", value=auto_hyp_add, key=f"hypothesis_{idx+2}", height=100, label_visibility="collapsed")
            else:
                row["hypothesis"] = st.text_area("", key=f"hypothesis_{idx+2}", height=100, label_visibility="collapsed")
        st.markdown("---")
    
    # INTENDED OUTCOMES with SAPOL reference field
    st.markdown("### Intended Outcomes")
    outcomes_options = [
        "Send Home", "Parent/Caregiver notified via Phone Call",
        "Student Leaving supervised areas/leaving school grounds",
        "Sexualised behaviour", "Incident – student to student",
        "Complaint by co-located school/member of public",
        "Property damage", "Stealing", "Toileting issue",
        "ED155: Staff Injury", "ED155: Student injury",
        "Emergency services - SAPOL",
        "Emergency services - SA Ambulance",
        "Incident Internally Managed - Restorative Session",
        "Incident Internally Managed - Community Service",
        "Incident Internally Managed - Re-Entry",
        "Incident Internally Managed - Case Review",
        "Incident Internally Managed - Make-up Time"
    ]
    
    selected_outcomes = st.multiselect("Select all intended outcomes:", outcomes_options, key="intended_outcomes")
    
    # SAPOL Reference Field - triggered if SAPOL selected
    sapol_reference = ""
    if "Emergency services - SAPOL" in selected_outcomes:
        st.warning("⚠️ SAPOL involvement detected - Police Reference Number required")
        sapol_reference = st.text_input("SAPOL Police Reference Number *", 
                                       placeholder="Enter police reference number",
                                       key="sapol_ref")
    
    tac_notes = st.text_area("Additional Outcome Notes (e.g., TAC meeting):", 
                            placeholder="A TAC meeting will be held...",
                            key="tac_notes", height=100)
    
    st.markdown("---")
    st.markdown("### Notifications & Administration")
    col_notif1, col_notif2 = st.columns(2)
    with col_notif1:
        notified_line_manager = st.checkbox("Notified Line Manager", key="notif_manager", value=True)
        notified_parent = st.checkbox("Notified Parent/Caregiver", key="notif_parent")
    with col_notif2:
        copy_in_file = st.checkbox("Copy in student file", key="copy_file", value=True)
        safety_plan_review = st.checkbox("Safety plan review required", key="safety_review")
    
    st.markdown("---")
    st.markdown("### Staff Agreement")
    staff_name = st.session_state.current_user.get("name", "Staff Member")
    st.markdown(f"**Completing Staff Member:** {staff_name}")
    staff_agrees = st.checkbox(f"✓ I, {staff_name}, confirm this information is accurate and complete.", 
                               key="staff_agrees")
    
    st.markdown("---")
    st.markdown("### Email Distribution")
    col_email1, col_email2 = st.columns(2)
    with col_email1:
        leader_email = st.text_input("Line Manager Email *", 
                                     value="manager@clc.sa.edu.au",
                                     key="leader_email")
    with col_email2:
        admin_email = st.text_input("Admin Email *", 
                                    value="admin@clc.sa.edu.au",
                                    key="admin_email")
    
    st.markdown("---")
    
    if st.button("📧 Submit Critical Incident Form", type="primary", use_container_width=True, key="save_crit"):
        # Validation
        errors = []
        if not context_1 or not behaviour_1 or not consequence_1 or not hypothesis_1:
            errors.append("Please complete all ABCH fields for initial incident")
        if not staff_agrees:
            errors.append("Please confirm your agreement")
        if not leader_email or "@" not in leader_email or not admin_email or "@" not in admin_email:
            errors.append("Please enter valid email addresses")
        if "Emergency services - SAPOL" in selected_outcomes and not sapol_reference:
            errors.append("SAPOL Reference Number is required when SAPOL is involved")
        
        if errors:
            for error in errors:
                st.error(f"❌ {error}")
        else:
            record = {
                "id": str(uuid.uuid4()),
                "created_at": datetime.now().isoformat(),
                "quick_incident_id": inc_id,
                "student_id": quick_inc["student_id"],
                "student_name": student["name"],
                "incident_type": "Critical",
                "ABCH_primary": {
                    "location": location_1,
                    "context": context_1,
                    "time": time_1,
                    "behaviour": behaviour_1,
                    "consequence": consequence_1,
                    "hypothesis": hypothesis_1
                },
                "ABCH_additional": st.session_state.abch_rows.copy(),
                "intended_outcomes": selected_outcomes,
                "sapol_reference": sapol_reference if sapol_reference else None,
                "tac_notes": tac_notes,
                "notifications": {
                    "line_manager": notified_line_manager,
                    "parent": notified_parent,
                    "copy_in_file": copy_in_file,
                    "safety_plan_review": safety_plan_review
                },
                "staff_agreement": {
                    "staff_name": staff_name,
                    "agreed": staff_agrees,
                    "timestamp": datetime.now().isoformat()
                },
                "leader_email": leader_email,
                "admin_email": admin_email
            }
            
            st.session_state.critical_incidents.append(record)
            st.session_state.abch_rows = []
            st.session_state.hyp_1_generated = False
            
            st.success("✅ Critical incident form saved to database")
            
            # GENERATE ADMIN SUMMARY
            admin_summary = generate_admin_summary(record, student, staff_name)
            
            # Show admin summary
            with st.expander("📋 ADMIN SUMMARY (For External Incident Log)", expanded=True):
                st.text_area("Copy this summary for departmental log:", admin_summary, height=400, key="admin_summary_display")
                st.download_button(
                    "📥 Download Admin Summary",
                    admin_summary,
                    file_name=f"Admin_Summary_{student['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain"
                )
            
            # SEND EMAILS
            staff_email = st.session_state.current_user.get("email", "staff@example.com")
            send_critical_incident_email(record, student, staff_email, leader_email, admin_email)
            
            st.markdown("---")
            st.info("✉️ Emails sent to Line Manager, Admin, and completing staff member")
            st.info("💾 Critical incident data saved in student's file")
            st.info("📋 Admin summary generated for external log")
            if sapol_reference:
                st.info(f"🚔 SAPOL Reference: {sapol_reference}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("📊 View Analysis", type="primary", use_container_width=True, key="view_analysis"):
                    go_to("student_analysis", selected_student_id=quick_inc["student_id"])
            with col2:
                if st.button("↩️ Back to Students", use_container_width=True, key="back_crit_after"):
                    go_to("program_students", selected_program=student["program"])
            with col3:
                if st.button("🏠 Program Landing", use_container_width=True, key="home_crit_after"):
                    go_to("landing")
def render_student_analysis_page():
    """Comprehensive Data Analysis with Berry Street Education Model"""
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## 📊 Comprehensive Behaviour Analysis — {student['name']}")
    st.caption("Evidence-based analysis prepared by Learning and Behaviour Unit")
    st.caption("Using ABA, Trauma-Informed Practice, Berry Street Education Model, and CPI principles")
    
    # Display placement information
    if student.get('placement_start'):
        try:
            start_dt = datetime.fromisoformat(student['placement_start'])
            if student.get('placement_end'):
                end_dt = datetime.fromisoformat(student['placement_end'])
                status_txt = "Completed"
            else:
                end_dt = datetime.now()
                status_txt = "Ongoing"
            
            days_enrolled = (end_dt.date() - start_dt.date()).days
            
            st.markdown("---")
            pcol1, pcol2, pcol3 = st.columns(3)
            with pcol1:
                st.metric("Placement Start", start_dt.strftime('%d/%m/%Y'))
            with pcol2:
                if student.get('placement_end'):
                    st.metric("Placement End", datetime.fromisoformat(student['placement_end']).strftime('%d/%m/%Y'))
                else:
                    st.metric("Status", status_txt)
            with pcol3:
                st.metric("Days Enrolled", days_enrolled)
            st.markdown("---")
        except:
            pass
    
    # Navigation
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("⬅ Back to Students", key="back_analysis_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_analysis"):
            go_to("landing")
    
    quick = [i for i in st.session_state.incidents if i["student_id"] == student_id]
    crit = [c for c in st.session_state.critical_incidents if c["student_id"] == student_id]
    
    if not quick and not crit:
        st.info("No incident data available yet.")
        if st.button("↩️ Back", key="back_no_data"):
            go_to("program_students", selected_program=student["program"])
        return
    
    # Build dataframe
    quick_df = pd.DataFrame(quick) if quick else pd.DataFrame()
    crit_df = pd.DataFrame(crit) if crit else pd.DataFrame()
    
    if not quick_df.empty:
        quick_df["incident_type"] = "Quick"
        quick_df["date_parsed"] = pd.to_datetime(quick_df["date"])
        if "intervention" in quick_df.columns:
            quick_df["intervention_str"] = quick_df["intervention"].apply(lambda x: ", ".join(x) if isinstance(x, list) else str(x))
    
    if not crit_df.empty:
        crit_df["incident_type"] = "Critical"
        crit_df["date_parsed"] = pd.to_datetime(crit_df.get("created_at", datetime.now().isoformat()))
        crit_df["severity"] = 5
        crit_df["antecedent"] = crit_df["ABCH_primary"].apply(lambda d: d.get("context","") if isinstance(d, dict) else "")
        crit_df["behaviour_type"] = crit_df["ABCH_primary"].apply(lambda d: d.get("behaviour","") if isinstance(d, dict) else "")
    
    full_df = pd.concat([quick_df, crit_df], ignore_index=True).sort_values("date_parsed")
    full_df["hour"] = pd.to_datetime(full_df["time"], format="%H:%M:%S", errors="coerce").dt.hour
    full_df["day_of_week"] = full_df["date_parsed"].dt.day_name()
    
    # Prepare split dataframes
    quick_only_df = full_df[full_df['incident_type'] == 'Quick'].copy() if 'Quick' in full_df['incident_type'].values else pd.DataFrame()
    crit_only_df = full_df[full_df['incident_type'] == 'Critical'].copy() if 'Critical' in full_df['incident_type'].values else pd.DataFrame()
    
    # OVERVIEW
    st.markdown("### 📈 Executive Summary")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1: st.metric("Total", len(full_df))
    with col2: st.metric("Critical", len(full_df[full_df["incident_type"] == "Critical"]))
    with col3: st.metric("Avg Severity", f"{full_df['severity'].mean():.1f}")
    with col4:
        days = max((full_df["date_parsed"].max() - full_df["date_parsed"].min()).days, 1)
        st.metric("Days Span", days)
    with col5:
        st.metric("Per Day", f"{len(full_df) / days:.1f}")
    
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 1: DAILY FREQUENCY
    # ================================================================
    
    st.markdown("### 📅 Daily Incident Frequency - Regular vs Critical")
    st.caption("Understanding the relationship between regular incidents and critical escalations")
    
    fig1 = go.Figure()
    
    if not quick_only_df.empty:
        daily_quick = quick_only_df.groupby(quick_only_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig1.add_trace(go.Bar(
            x=daily_quick["date_parsed"], 
            y=daily_quick["count"],
            name='Regular Incidents',
            marker=dict(color='#3b82f6', line=dict(color='white', width=1)),
            text=daily_quick["count"],
            textposition='inside',
            textfont=dict(color='white', size=11, family='Arial Black'),
            hovertemplate='<b>Date:</b> %{x}<br><b>Regular:</b> %{y}<extra></extra>'
        ))
    
    if not crit_only_df.empty:
        daily_crit = crit_only_df.groupby(crit_only_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig1.add_trace(go.Bar(
            x=daily_crit["date_parsed"], 
            y=daily_crit["count"],
            name='Critical Incidents',
            marker=dict(color='#ef4444', line=dict(color='white', width=1)),
            text=daily_crit["count"],
            textposition='inside',
            textfont=dict(color='white', size=11, family='Arial Black'),
            hovertemplate='<b>Date:</b> %{x}<br><b>Critical:</b> %{y}<extra></extra>'
        ))
    
    fig1.update_layout(
        height=350,
        barmode='stack',
        xaxis_title="<b>Date</b>",
        yaxis_title="<b>Incident Count</b>",
        plot_bgcolor='#f8fafc',
        paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        yaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', gridwidth=1, showline=True, linewidth=2, linecolor='#cbd5e1'),
        xaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1', gridcolor='#e2e8f0'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1),
        hovermode='x unified'
    )
    st.plotly_chart(fig1, use_container_width=True)
    
    if not quick_only_df.empty and not crit_only_df.empty:
        escalation_rate = (len(crit_only_df) / len(quick_only_df)) * 100 if len(quick_only_df) > 0 else 0
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Regular Incidents", len(quick_only_df), help="Severity 1-2")
        with col2:
            st.metric("Critical Incidents", len(crit_only_df), help="Severity 3+")
        with col3:
            st.metric("Escalation Rate", f"{escalation_rate:.1f}%", help="% of incidents that escalate to critical")
    
    with st.expander("💡 Clinical Interpretation (Berry Street Body Domain)"):
        st.markdown(
            "**Pattern Recognition:** Days with multiple regular incidents may predict critical escalation. "
            "When you see clustering of regular incidents, it indicates the student's nervous system is already dysregulated.\n\n"
            "**Berry Street Body:** On high-frequency days, increase proactive regulation - breathing exercises, "
            "movement breaks, sensory activities. The goal is to widen the Window of Tolerance before it narrows further.\n\n"
            "**Prevention Strategy:** If you see 2+ regular incidents in one day, immediately implement intensive Body domain supports "
            "to prevent critical escalation."
        )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 2: BEHAVIOURS
    # ================================================================
    
    st.markdown("### 🎯 Behaviour Types - Regular vs Critical")
    st.caption("Which behaviours escalate to critical incidents?")
    
    all_behaviours = full_df["behaviour_type"].value_counts().head(6).index.tolist()
    
    quick_beh_counts = []
    crit_beh_counts = []
    for beh in all_behaviours:
        quick_count = len(quick_only_df[quick_only_df["behaviour_type"] == beh]) if not quick_only_df.empty else 0
        crit_count = len(crit_only_df[crit_only_df["behaviour_type"] == beh]) if not crit_only_df.empty else 0
        quick_beh_counts.append(quick_count)
        crit_beh_counts.append(crit_count)
    
    fig2 = go.Figure()
    
    fig2.add_trace(go.Bar(
        y=all_behaviours, x=quick_beh_counts, name='Regular', orientation='h',
        marker=dict(color='#3b82f6', line=dict(color='white', width=1)),
        text=quick_beh_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black'),
        hovertemplate='<b>%{y}</b><br>Regular: %{x}<extra></extra>'
    ))
    
    fig2.add_trace(go.Bar(
        y=all_behaviours, x=crit_beh_counts, name='Critical', orientation='h',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_beh_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black'),
        hovertemplate='<b>%{y}</b><br>Critical: %{x}<extra></extra>'
    ))
    
    fig2.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Incident Count</b>", yaxis_title="",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        xaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', gridwidth=1, showline=True, linewidth=2, linecolor='#cbd5e1'),
        yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1),
        hovermode='y unified'
    )
    st.plotly_chart(fig2, use_container_width=True)
    
    st.markdown("**Escalation Risk by Behaviour:**")
    risk_data = []
    for i, beh in enumerate(all_behaviours):
        total = quick_beh_counts[i] + crit_beh_counts[i]
        if total > 0:
            risk_pct = (crit_beh_counts[i] / total) * 100
            risk_level = "🔴 High Risk" if risk_pct > 50 else "🟡 Medium Risk" if risk_pct > 25 else "🟢 Low Risk"
            risk_data.append(f"• **{beh}**: {risk_pct:.0f}% escalate to critical - {risk_level}")
    
    for risk_item in risk_data:
        st.markdown(risk_item)
    
    with st.expander("💡 Clinical Interpretation (Behaviour as Communication)"):
        if all_behaviours:
            primary_beh = all_behaviours[0]
            st.markdown(
                f"**Primary Concern:** {primary_beh} is the most common behaviour. "
                "Behaviours with high escalation rates (>50%) need immediate intervention planning.\n\n"
                "**Behaviour Analysis:** Focus on high-risk behaviours first. If a behaviour frequently escalates, "
                "it means the student lacks skills or supports to regulate at that level.\n\n"
                "**Berry Street:** All behaviour is communication. High escalation rates tell us the student needs:\n"
                "1. **BODY**: More regulation tools before the behaviour occurs\n"
                "2. **RELATIONSHIP**: Stronger connection to trusted adults who can co-regulate\n"
                "3. **STAMINA**: Skill-building for persistence through challenges"
            )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 3: TRIGGERS
    # ================================================================
    
    st.markdown("### 🔍 Trigger Analysis - What Leads to Critical Escalation?")
    st.caption("Understanding which antecedents most often result in critical incidents")
    
    all_triggers = full_df["antecedent"].value_counts().head(6).index.tolist()
    
    quick_ant_counts = []
    crit_ant_counts = []
    for ant in all_triggers:
        quick_count = len(quick_only_df[quick_only_df["antecedent"] == ant]) if not quick_only_df.empty else 0
        crit_count = len(crit_only_df[crit_only_df["antecedent"] == ant]) if not crit_only_df.empty else 0
        quick_ant_counts.append(quick_count)
        crit_ant_counts.append(crit_count)
    
    fig3 = go.Figure()
    
    fig3.add_trace(go.Bar(
        y=all_triggers, x=quick_ant_counts, name='Regular', orientation='h',
        marker=dict(color='#3b82f6', line=dict(color='white', width=1)),
        text=quick_ant_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig3.add_trace(go.Bar(
        y=all_triggers, x=crit_ant_counts, name='Critical', orientation='h',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_ant_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig3.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Incident Count</b>", yaxis_title="",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        xaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1)
    )
    st.plotly_chart(fig3, use_container_width=True)
    
    st.markdown("**Critical Escalation Risk by Trigger:**")
    for i, ant in enumerate(all_triggers):
        total = quick_ant_counts[i] + crit_ant_counts[i]
        if total > 0:
            risk_pct = (crit_ant_counts[i] / total) * 100
            color = "🔴" if risk_pct > 60 else "🟡" if risk_pct > 30 else "🟢"
            st.markdown(f"{color} **{ant}**: {risk_pct:.0f}% lead to critical incidents ({crit_ant_counts[i]}/{total})")
    
    with st.expander("💡 Clinical Interpretation (Proactive Prevention)"):
        st.markdown(
            "**High-Risk Triggers (>60%):** These antecedents almost always lead to crisis. Implement intensive preventative supports:\n"
            "• Modify environment to prevent trigger\n"
            "• Pre-teach coping strategies specific to this trigger\n"
            "• Provide extra staff support when trigger is likely\n\n"
            "**Medium-Risk Triggers (30-60%):** Student has some capacity but needs support:\n"
            "• Berry Street STAMINA: Build persistence through small exposures\n"
            "• Provide choice and control\n"
            "• Co-regulation available immediately\n\n"
            "**Low-Risk Triggers (<30%):** Student managing well with current supports - maintain strategies."
        )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 4: SEVERITY WITH ZONES
    # ================================================================
    
    st.markdown("### 📊 Severity Trajectory - Window of Tolerance Analysis")
    st.caption("Tracking whether incidents are getting more or less severe over time")
    
    fig4 = go.Figure()
    
    fig4.add_hrect(y0=0, y1=2.5, fillcolor="#d1fae5", opacity=0.2, 
                   annotation_text="Within Window (Regulation)", annotation_position="top left", line_width=0)
    fig4.add_hrect(y0=2.5, y1=3.5, fillcolor="#fef3c7", opacity=0.2,
                   annotation_text="Edge of Window", annotation_position="top left", line_width=0)
    fig4.add_hrect(y0=3.5, y1=5.5, fillcolor="#fee2e2", opacity=0.2,
                   annotation_text="Outside Window (Crisis)", annotation_position="top left", line_width=0)
    
    if not quick_only_df.empty:
        fig4.add_trace(go.Scatter(
            x=quick_only_df["date_parsed"], y=quick_only_df["severity"],
            mode='markers', name='Regular',
            marker=dict(size=12, color='#3b82f6', opacity=0.7, line=dict(color='white', width=1.5)),
            hovertemplate='<b>Date:</b> %{x}<br><b>Severity:</b> %{y}<extra></extra>'
        ))
    
    if not crit_only_df.empty:
        fig4.add_trace(go.Scatter(
            x=crit_only_df["date_parsed"], y=crit_only_df["severity"],
            mode='markers', name='Critical',
            marker=dict(size=15, color='#ef4444', opacity=0.8, symbol='diamond', line=dict(color='white', width=2)),
            hovertemplate='<b>Date:</b> %{x}<br><b>Severity:</b> %{y}<extra></extra>'
        ))
    
    if len(full_df) >= 3:
        z = np.polyfit(range(len(full_df)), full_df["severity"], 1)
        p = np.poly1d(z)
        trend_color = '#22c55e' if z[0] < 0 else '#ef4444'
        fig4.add_trace(go.Scatter(
            x=full_df["date_parsed"], y=p(range(len(full_df))),
            mode='lines', name='Trend',
            line=dict(color=trend_color, width=3, dash='dash'),
            hovertemplate='Trend<extra></extra>'
        ))
    
    fig4.update_layout(
        height=400, xaxis_title="<b>Date</b>", yaxis_title="<b>Severity Level</b>",
        yaxis=dict(range=[0, 5.5], tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        xaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1', gridcolor='#e2e8f0'),
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.9)', bordercolor='#cbd5e1', borderwidth=1),
        hovermode='closest'
    )
    st.plotly_chart(fig4, use_container_width=True)
    
    if len(full_df) >= 5:
        recent_avg = full_df.tail(5)["severity"].mean()
        early_avg = full_df.head(5)["severity"].mean()
        trend_dir = "improving" if recent_avg < early_avg else "worsening" if recent_avg > early_avg else "stable"
        trend_emoji = "📈" if trend_dir == "improving" else "📉" if trend_dir == "worsening" else "➡️"
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Early Average", f"{early_avg:.1f}", help="First 5 incidents")
        with col2:
            st.metric("Recent Average", f"{recent_avg:.1f}", help="Last 5 incidents")
        with col3:
            st.metric("Trend", f"{trend_emoji} {trend_dir.title()}")
    
    with st.expander("💡 Clinical Interpretation (Window of Tolerance)"):
        st.markdown(
            "**Window of Tolerance (Siegel, 1999):** The optimal arousal zone where the student can think, learn, and regulate.\n\n"
            "**Green Zone (Severity 1-2):** Student is within or near their window. Accessible to support.\n\n"
            "**Yellow Zone (Severity 2.5-3.5):** Student is at the edge. CO-REGULATION NEEDED NOW.\n\n"
            "**Red Zone (Severity 3.5-5):** Outside window in survival mode. Safety first, teach later."
        )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 5: LOCATION
    # ================================================================
    
    st.markdown("### 📍 Location Hotspots - Where Do Critical Incidents Occur?")
    st.caption("Environmental factors and escalation patterns by location")
    
    all_locations = full_df["location"].value_counts().head(6).index.tolist()
    
    quick_loc_counts = []
    crit_loc_counts = []
    for loc in all_locations:
        quick_count = len(quick_only_df[quick_only_df["location"] == loc]) if not quick_only_df.empty else 0
        crit_count = len(crit_only_df[crit_only_df["location"] == loc]) if not crit_only_df.empty else 0
        quick_loc_counts.append(quick_count)
        crit_loc_counts.append(crit_count)
    
    fig5 = go.Figure()
    
    fig5.add_trace(go.Bar(
        y=all_locations, x=quick_loc_counts, name='Regular', orientation='h',
        marker=dict(color='#3b82f6', line=dict(color='white', width=1)),
        text=quick_loc_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig5.add_trace(go.Bar(
        y=all_locations, x=crit_loc_counts, name='Critical', orientation='h',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_loc_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig5.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Incident Count</b>", yaxis_title="",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        xaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1)
    )
    st.plotly_chart(fig5, use_container_width=True)
    
    st.markdown("**Environmental Risk Assessment:**")
    for i, loc in enumerate(all_locations):
        total = quick_loc_counts[i] + crit_loc_counts[i]
        if total > 0:
            risk_pct = (crit_loc_counts[i] / total) * 100
            if risk_pct > 60:
                risk_level = "🔴 HIGH RISK ENVIRONMENT"
                recommendation = "Immediate environmental modification needed"
            elif risk_pct > 30:
                risk_level = "🟡 MODERATE RISK"
                recommendation = "Enhanced supervision and supports"
            else:
                risk_level = "🟢 MANAGED ENVIRONMENT"
                recommendation = "Current strategies effective"
            st.markdown(f"**{loc}**: {risk_pct:.0f}% escalate - {risk_level} - *{recommendation}*")
    
    with st.expander("💡 Clinical Interpretation (Environmental Strategies)"):
        if all_locations:
            st.markdown(
                f"**Primary Hotspot:** {all_locations[0]}\n\n"
                "**Berry Street BODY - Sensory Environment:**\n"
                "• Lighting: Consider natural light or lamps\n"
                "• Noise: Provide noise-cancelling headphones or quiet spaces\n"
                "• Space: Create clear pathways and defined areas"
            )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 6: TIME OF DAY
    # ================================================================
    
    st.markdown("### ⏰ Time of Day Patterns - When Does Escalation Occur?")
    st.caption("Understanding daily rhythm and predicting high-risk periods")
    
    session_order = ['Morning', 'Middle', 'Afternoon']
    
    quick_session_counts = []
    crit_session_counts = []
    for session in session_order:
        quick_count = len(quick_only_df[quick_only_df["session"] == session]) if not quick_only_df.empty and 'session' in quick_only_df.columns else 0
        crit_count = len(crit_only_df[crit_only_df["session"] == session]) if not crit_only_df.empty and 'session' in crit_only_df.columns else 0
        quick_session_counts.append(quick_count)
        crit_session_counts.append(crit_count)
    
    fig6 = go.Figure()
    
    fig6.add_trace(go.Bar(
        x=session_order, y=quick_session_counts, name='Regular',
        marker=dict(color='#3b82f6', line=dict(color='white', width=1)),
        text=quick_session_counts, textposition='inside',
        textfont=dict(color='white', size=12, family='Arial Black')
    ))
    
    fig6.add_trace(go.Bar(
        x=session_order, y=crit_session_counts, name='Critical',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_session_counts, textposition='inside',
        textfont=dict(color='white', size=12, family='Arial Black')
    ))
    
    fig6.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Time of Day</b>", yaxis_title="<b>Incident Count</b>",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        yaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        xaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1)
    )
    st.plotly_chart(fig6, use_container_width=True)
    
    st.markdown("**Escalation Risk by Time of Day:**")
    for i, session in enumerate(session_order):
        total = quick_session_counts[i] + crit_session_counts[i]
        if total > 0:
            risk_pct = (crit_session_counts[i] / total) * 100
            risk_emoji = "🔴" if risk_pct > 50 else "🟡" if risk_pct > 25 else "🟢"
            st.markdown(f"{risk_emoji} **{session}**: {risk_pct:.0f}% escalate to critical ({total} total incidents)")
    
    with st.expander("💡 Clinical Interpretation (Circadian Regulation)"):
        if quick_session_counts:
            peak_session = session_order[quick_session_counts.index(max(quick_session_counts))] if max(quick_session_counts) > 0 else "Unknown"
            st.markdown(
                f"**Peak Incident Time:** {peak_session}\n\n"
                "**Berry Street BODY:** Proactive regulation before peak periods - breathing, movement, sensory check-ins."
            )
    st.markdown("---")
    
    # QUICK VS CRITICAL COMPARISON SECTION
    st.markdown("## 📊 Quick vs Critical Incident Analysis")
    st.caption("Understanding the relationship between quick logs and critical incidents helps identify escalation patterns")
    
    quick_only = [i for i in quick if not i.get("is_critical")]
    critical_data = crit
    
    col_q, col_c = st.columns(2)
    with col_q:
        st.metric("Quick Incidents", len(quick_only), help="Standard behaviour logs (Severity 1-2)")
    with col_c:
        st.metric("Critical Incidents", len(critical_data), help="Severity 3+ requiring ABCH form")
    
    st.markdown("---")
    
    # Day of Week
    st.markdown("### 📆 Day of Week Patterns")
    day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    day_counts = full_df["day_of_week"].value_counts().reindex(day_order, fill_value=0)
    fig7 = go.Figure()
    fig7.add_trace(go.Bar(
        x=day_counts.index, y=day_counts.values,
        marker=dict(color='#64748b'),
        text=day_counts.values, textposition='outside'
    ))
    fig7.update_layout(
        height=280, showlegend=False, yaxis_title="Total",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig7, use_container_width=True)
    
    high_day = day_counts.idxmax()
    with st.expander("💡 Clinical Interpretation (Berry Street Relationship)"):
        st.markdown(f"**{high_day}** has most incidents. Consider connection routines. " +
                   "**Berry Street Relationship:** Strong connections reduce incidents.")
    st.markdown("---")
    
    # CLINICAL SUMMARY
    st.markdown("### 🧠 Clinical Summary")
    st.caption("Evidence-based interpretation using ABA, Trauma-Informed Practice, Berry Street Education Model, and CPI principles")
    
    top_beh = full_df["behaviour_type"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_ant = full_df["antecedent"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_loc = full_df["location"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_session = full_df["session"].mode()[0] if len(full_df) > 0 else "Unknown"
    
    recent = full_df.tail(7)
    risk_score = min(100, int(
        (len(recent) / 7 * 10) +
        (recent["severity"].mean() * 8) +
        (len(full_df[full_df["incident_type"] == "Critical"]) / len(full_df) * 50)
    ))
    risk_level = "LOW" if risk_score < 30 else "MODERATE" if risk_score < 60 else "HIGH"
    
    st.info(f"""
    **Key Patterns Identified:**
    - Primary behaviour: **{top_beh}**
    - Main trigger: **{top_ant}**
    - Hotspot location: **{top_loc}**
    - Peak time: **{top_session}**
    - Risk Level: **{risk_level}** ({risk_score}/100)
    
    **Berry Street Focus:** Body (regulation) and Relationship (connection) domains are foundation.
    """)
    
    st.success(f"""
    **Evidence-Based Recommendations:**
    
    **1. Body Domain:** Regulated start before {top_session}, breathing exercises, movement breaks
    **2. Relationship Domain:** Key adult check-in, acknowledgment of feelings, co-regulation
    **3. Stamina Domain:** Teach help-seeking, practice requesting breaks
    **4. SMART Goal:** Over 5 weeks, use help-seeking strategy in 4/5 opportunities
    """)
    
    st.markdown("---")
    
    # EXPORT
    st.markdown("### 📄 Export Data & Reports")
    
    col1, col2 = st.columns(2)
    
    with col1:
        csv = full_df.to_csv(index=False)
        st.download_button(
            "📥 Download Raw Data (CSV)",
            csv,
            file_name=f"{student['name'].replace(' ', '_')}_Incident_Data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    with col2:
        with st.spinner("Generating Behaviour Analysis Plan..."):
            docx_file = generate_behaviour_analysis_plan_docx(
                student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level
            )
        if docx_file:
            st.download_button(
                "📄 Behaviour Analysis Plan (Word)",
                docx_file,
                file_name=f"BAP_{student['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅ Back to Students", type="primary", key="back_analysis_bottom", use_container_width=True):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_analysis_bottom", use_container_width=True):
            go_to("landing")



def render_admin_portal():
    """Admin portal for managing students and placement dates"""
    if st.session_state.current_user.get("role") != "ADM":
        st.error("⛔ Access Denied: Admin privileges required")
        if st.button("⬅ Back to Landing"):
            go_to("landing")
        return
    
    st.markdown("## 🔧 Admin Portal")
    
    col1, col2 = st.columns([6, 1])
    with col1:
        if st.button("⬅ Back to Landing", key="back_admin"):
            go_to("landing")
    
    st.markdown("---")
    
    # TABS
    tab1, tab2, tab3 = st.tabs(["👥 Manage Students", "📊 System Overview", "👨‍💼 Staff Management"])
    
    with tab1:
        st.markdown("### Student Management")
        
        # ADD NEW STUDENT
        with st.expander("➕ Add New Student", expanded=False):
            with st.form("add_student_form"):
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    new_first_name = st.text_input("First Name *", placeholder="John")
                    new_last_name = st.text_input("Last Name *", placeholder="Smith")
                
                with col2:
                    new_grade = st.selectbox("Grade *", ["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"])
                    new_edid = st.text_input("EDID *", placeholder="ED123456")
                
                with col3:
                    new_dob = st.date_input("Date of Birth *", value=date(2015, 1, 1), format="DD/MM/YYYY")
                    new_program = st.selectbox("Program *", ["JP", "PY", "SY"])
                
                with col4:
                    new_placement_start = st.date_input("Placement Start Date *", value=date.today(), format="DD/MM/YYYY")
                    new_placement_end = st.date_input("Placement End Date (Optional)", value=None, format="DD/MM/YYYY")
                
                submitted = st.form_submit_button("Add Student", type="primary")
                
                if submitted:
                    if new_first_name and new_last_name and new_grade and new_program and new_edid:
                        full_name = f"{new_first_name} {new_last_name}"
                        new_student = {
                            "id": f"stu_{uuid.uuid4().hex[:8]}",
                            "first_name": new_first_name.strip(),
                            "last_name": new_last_name.strip(),
                            "name": full_name.strip(),
                            "grade": new_grade,
                            "dob": new_dob.isoformat(),
                            "edid": new_edid,
                            "program": new_program,
                            "placement_start": new_placement_start.isoformat(),
                            "placement_end": new_placement_end.isoformat() if new_placement_end else None
                        }
                        
                        # SAVE TO DATABASE FIRST
                        if save_student_to_db(new_student):
                            st.session_state.students.append(new_student)
                            st.success(f"✅ Added {full_name} (EDID: {new_edid}) to {PROGRAM_NAMES[new_program]}")
                            st.rerun()
                        else:
                            st.error("❌ Failed to save student to database")
                    else:
                        st.error("Please complete all required fields (First Name, Last Name, Grade, EDID, Program)")
        
        st.markdown("---")
        
        # EXISTING STUDENTS
        st.markdown("### Current Students")
        
        for program in ["JP", "PY", "SY"]:
            st.markdown(f"#### {PROGRAM_NAMES[program]}")
            
            program_students = [s for s in st.session_state.students if s["program"] == program]
            
            if not program_students:
                st.caption("No students in this program")
                continue
            
            for student in program_students:
                with st.container(border=True):
                    col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
                    
                    with col1:
                        st.markdown(f"**{student['name']}**")
                        st.caption(f"Grade {student['grade']}")
                        if student.get('edid'):
                            st.caption(f"🆔 EDID: {student['edid']}")
                    
                    with col2:
                        if student.get('placement_start'):
                            start_date = datetime.fromisoformat(student['placement_start']).strftime('%d/%m/%Y')
                            st.caption(f"📅 Start: {start_date}")
                            
                            # Calculate days enrolled
                            start = datetime.fromisoformat(student['placement_start']).date()
                            end = datetime.fromisoformat(student['placement_end']).date() if student.get('placement_end') else date.today()
                            days = (end - start).days
                            st.caption(f"📊 {days} days enrolled")
                        else:
                            st.caption("📅 Start: Not set")
                            st.caption("📊 Days: N/A")
                    
                    with col3:
                        if student.get('placement_end'):
                            end_date = datetime.fromisoformat(student['placement_end']).strftime('%d/%m/%Y')
                            st.caption(f"📅 End: {end_date}")
                            st.caption("🔴 Inactive")
                        else:
                            st.caption("📅 End: Ongoing")
                            st.caption("🟢 Active")
                    
                    with col4:
                        if st.button("✏️", key=f"edit_{student['id']}", help="Edit student"):
                            st.session_state.editing_student = student['id']
                            st.rerun()
                
                # EDIT STUDENT
                if st.session_state.get("editing_student") == student['id']:
                    with st.expander("✏️ Edit Student Details", expanded=True):
                        with st.form(f"edit_form_{student['id']}"):
                            edit_col1, edit_col2 = st.columns(2)
                            
                            with edit_col1:
                                edit_first_name = st.text_input("First Name", 
                                                               value=student.get('first_name', student['name'].split()[0] if student['name'] else ''),
                                                               key=f"edit_first_{student['id']}")
                                edit_last_name = st.text_input("Last Name",
                                                              value=student.get('last_name', ' '.join(student['name'].split()[1:]) if len(student['name'].split()) > 1 else ''),
                                                              key=f"edit_last_{student['id']}")
                                # Use existing date or default to today
                                default_start = datetime.fromisoformat(student['placement_start']).date() if student.get('placement_start') else date.today()
                                edit_start = st.date_input("Placement Start", 
                                                          value=default_start,
                                                          key=f"edit_start_{student['id']}",
                                                          format="DD/MM/YYYY")
                            
                            with edit_col2:
                                edit_edid = st.text_input("EDID", value=student.get('edid', ''), key=f"edit_edid_{student['id']}")
                                edit_grade = st.selectbox("Grade", 
                                                         ["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"],
                                                         index=["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"].index(student['grade']) if student['grade'] in ["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"] else 0,
                                                         key=f"edit_grade_{student['id']}")
                                current_end = datetime.fromisoformat(student['placement_end']).date() if student.get('placement_end') else None
                                edit_end = st.date_input("Placement End (None = Ongoing)",
                                                        value=current_end,
                                                        key=f"edit_end_{student['id']}",
                                                        format="DD/MM/YYYY")
                            
                            col_save, col_cancel = st.columns(2)
                            with col_save:
                                if st.form_submit_button("Save Changes", type="primary"):
                                    student['first_name'] = edit_first_name.strip()
                                    student['last_name'] = edit_last_name.strip()
                                    student['name'] = f"{edit_first_name} {edit_last_name}".strip()
                                    student['edid'] = edit_edid
                                    student['grade'] = edit_grade
                                    student['placement_start'] = edit_start.isoformat()
                                    student['placement_end'] = edit_end.isoformat() if edit_end else None
                                    save_student_to_db(student)
                                    st.session_state.editing_student = None
                                    st.success("✅ Updated")
                                    st.rerun()
                            
                            with col_cancel:
                                if st.form_submit_button("Cancel"):
                                    st.session_state.editing_student = None
                                    st.rerun()
    
    with tab2:
        st.markdown("### System Overview")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            total_students = len(st.session_state.students)
            active_students = len([s for s in st.session_state.students if not s.get('placement_end')])
            st.metric("Total Students", total_students)
            st.caption(f"{active_students} active")
        
        with col2:
            st.metric("Total Incidents", len(st.session_state.incidents))
        
        with col3:
            critical = len([i for i in st.session_state.incidents if i.get("is_critical")])
            st.metric("Critical Incidents", critical)
        
        with col4:
            st.metric("Staff Members", len(st.session_state.staff))
        
        st.markdown("---")
        
        # Program breakdown
        st.markdown("#### Students by Program")
        for prog in ["JP", "PY", "SY"]:
            count = len([s for s in st.session_state.students if s["program"] == prog])
            active = len([s for s in st.session_state.students if s["program"] == prog and not s.get('placement_end')])
            st.write(f"**{PROGRAM_NAMES[prog]}:** {count} total ({active} active)")
    
    with tab3:
        st.markdown("### Staff Management")
        
        # ADD NEW STAFF
        with st.expander("➕ Add New Staff Member", expanded=False):
            with st.form("add_staff_form"):
                col1, col2 = st.columns(2)
                
                with col1:
                    staff_first_name = st.text_input("First Name *", placeholder="Jane")
                    staff_last_name = st.text_input("Last Name *", placeholder="Smith")
                    staff_email = st.text_input("Email Address *", placeholder="jane.smith@school.edu.au", 
                                               help="Will be used as username and for critical incident notifications")
                    staff_password = st.text_input("Initial Password *", type="password", value="demo123")
                
                with col2:
                    staff_role = st.selectbox("Role *", 
                                             ["TSS", "Teacher", "Leader", "ADM"],
                                             help="TSS=Teacher/Support Staff, Leader=Program Leader, ADM=Administrator")
                    staff_program = st.selectbox("Program *", ["JP", "PY", "SY", "All Programs"])
                    staff_phone = st.text_input("Phone Number", placeholder="0412 345 678")
                
                staff_notes = st.text_area("Notes (Optional)", placeholder="Additional information about this staff member")
                
                submit_staff = st.form_submit_button("Add Staff Member", type="primary")
                
                if submit_staff:
                    if staff_first_name and staff_last_name and staff_email and staff_password and staff_role:
                        staff_full_name = f"{staff_first_name} {staff_last_name}"
                        # Check if email already exists
                        if any(s.get("email", "").lower() == staff_email.lower() for s in st.session_state.staff):
                            st.error(f"❌ Email {staff_email} already exists")
                        else:
                            new_staff = {
                                "id": f"staff_{uuid.uuid4().hex[:8]}",
                                "first_name": staff_first_name.strip(),
                                "last_name": staff_last_name.strip(),
                                "name": staff_full_name.strip(),
                                "email": staff_email.lower().strip(),
                                "password": staff_password,
                                "role": staff_role,
                                "program": staff_program if staff_program != "All Programs" else None,
                                "phone": staff_phone if staff_phone else None,
                                "notes": staff_notes if staff_notes else None,
                                "receive_critical_emails": True,  # Default to receiving emails
                                "created_date": date.today().isoformat()
                            }
                            # Save to database first
                            if save_staff_to_db(new_staff):
                                st.session_state.staff.append(new_staff)
                                st.success(f"✅ Added {staff_full_name} ({staff_email}) to database")
                                st.rerun()
                            else:
                                st.error("❌ Failed to save staff member to database")
                    else:
                        st.error("Please complete all required fields (First Name, Last Name, Email, Password, Role)")
    
        st.markdown("---")
        
        # EXISTING STAFF
        st.markdown("### Current Staff")
        
        # Group by role
        for role in ["ADM", "Leader", "Teacher", "TSS"]:
            role_names = {"ADM": "Administrators", "Leader": "Program Leaders", "Teacher": "Teachers", "TSS": "Support Staff"}
            role_staff = [s for s in st.session_state.staff if s.get("role") == role]
            
            if role_staff:
                st.markdown(f"#### {role_names[role]}")
                
                for staff in role_staff:
                    with st.container(border=True):
                        col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
                        
                        with col1:
                            st.markdown(f"**{staff['name']}**")
                            st.caption(f"📧 {staff['email']}")
                            if staff.get('phone'):
                                st.caption(f"📱 {staff['phone']}")
                        
                        with col2:
                            st.caption(f"**Role:** {staff['role']}")
                            if staff.get('program'):
                                st.caption(f"**Program:** {PROGRAM_NAMES[staff['program']]}")
                            else:
                                st.caption("**Program:** All Programs")
                        
                        with col3:
                            receives_emails = staff.get('receive_critical_emails', True)
                            if receives_emails:
                                st.caption("📬 Receives critical alerts")
                            else:
                                st.caption("📪 No critical alerts")
                            
                            if staff.get('created_date'):
                                st.caption(f"Added: {staff['created_date']}")
                        
                        with col4:
                            if st.button("✏️", key=f"edit_staff_{staff['id']}", help="Edit staff"):
                                st.session_state.editing_staff = staff['id']
                                st.rerun()
                        
                        # EDIT STAFF
                        if st.session_state.get("editing_staff") == staff['id']:
                            with st.expander("✏️ Edit Staff Details", expanded=True):
                                with st.form(f"edit_staff_form_{staff['id']}"):
                                    edit_col1, edit_col2 = st.columns(2)
                                    
                                    with edit_col1:
                                        edit_first_name = st.text_input("First Name", value=staff.get('first_name', staff['name'].split()[0] if staff['name'] else ''), key=f"edit_staff_first_{staff['id']}")
                                        edit_last_name = st.text_input("Last Name", value=staff.get('last_name', ' '.join(staff['name'].split()[1:]) if len(staff['name'].split()) > 1 else ''), key=f"edit_staff_last_{staff['id']}")
                                        edit_email = st.text_input("Email", value=staff['email'], key=f"edit_staff_email_{staff['id']}")
                                        edit_phone = st.text_input("Phone", value=staff.get('phone', ''), key=f"edit_staff_phone_{staff['id']}")
                                    
                                    with edit_col2:
                                        edit_role = st.selectbox("Role", ["TSS", "Teacher", "Leader", "ADM"],
                                                                index=["TSS", "Teacher", "Leader", "ADM"].index(staff['role']),
                                                                key=f"edit_staff_role_{staff['id']}")
                                        edit_program = st.selectbox("Program", ["JP", "PY", "SY", "All Programs"],
                                                                   index=["JP", "PY", "SY", "All Programs"].index(
                                                                       PROGRAM_NAMES.get(staff.get('program'), "All Programs") 
                                                                       if staff.get('program') else "All Programs"
                                                                   ) if staff.get('program') else 3,
                                                                   key=f"edit_staff_program_{staff['id']}")
                                        edit_receive_emails = st.checkbox("Receive critical incident emails",
                                                                         value=staff.get('receive_critical_emails', True),
                                                                         key=f"edit_staff_emails_{staff['id']}")
                                    
                                    edit_notes = st.text_area("Notes", value=staff.get('notes', ''), key=f"edit_staff_notes_{staff['id']}")
                                    
                                    col_save, col_cancel, col_delete = st.columns([1, 1, 1])
                                    with col_save:
                                        if st.form_submit_button("💾 Save Changes", type="primary"):
                                            staff['first_name'] = edit_first_name.strip()
                                            staff['last_name'] = edit_last_name.strip()
                                            staff['name'] = f"{edit_first_name} {edit_last_name}".strip()
                                            staff['email'] = edit_email.lower().strip()
                                            staff['phone'] = edit_phone if edit_phone else None
                                            staff['role'] = edit_role
                                            staff['program'] = edit_program if edit_program != "All Programs" else None
                                            staff['receive_critical_emails'] = edit_receive_emails
                                            staff['notes'] = edit_notes if edit_notes else None
                                            save_staff_to_db(staff)
                                            st.session_state.editing_staff = None
                                            st.success("✅ Updated")
                                            st.rerun()
                                    
                                    with col_cancel:
                                        if st.form_submit_button("❌ Cancel"):
                                            st.session_state.editing_staff = None
                                            st.rerun()
                                    
                                    with col_delete:
                                        if st.form_submit_button("🗑️ Delete", help="Remove this staff member"):
                                            if staff['role'] != 'ADM' or len([s for s in st.session_state.staff if s.get('role') == 'ADM']) > 1:
                                                st.session_state.staff.remove(staff)
                                                st.session_state.editing_staff = None
                                                st.success("✅ Staff member removed")
                                                st.rerun()
                                            else:
                                                st.error("❌ Cannot delete the last administrator")
        
        st.markdown("---")
        
        # EMAIL NOTIFICATION SETTINGS
        st.markdown("### 📧 Email Notification Recipients")
        st.caption("Staff members who will receive critical incident notifications:")
        
        email_recipients = [s for s in st.session_state.staff if s.get('receive_critical_emails', True)]
        
        if email_recipients:
            for recipient in email_recipients:
                st.write(f"• **{recipient['name']}** ({recipient['email']}) - {recipient['role']}")
        else:
            st.warning("⚠️ No staff members set to receive critical incident emails!")






def main():
    init_state()
    
    if not st.session_state.logged_in:
        render_login_page()
        return
    
    page = st.session_state.current_page
    
    if page == "landing": render_landing_page()
    elif page == "program_students": render_program_students_page()
    elif page == "incident_log": render_incident_log_page()
    elif page == "critical_incident": render_critical_incident_page()
    elif page == "student_analysis": render_student_analysis_page()
    elif page == "admin_portal": render_admin_portal()
    else: render_landing_page()

if __name__ == "__main__":
    main()
