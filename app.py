import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime, date, time, timedelta
import uuid
import random
from io import BytesIO
import base64

st.set_page_config(page_title="CLC Behaviour Support", page_icon="📊", layout="wide", initial_sidebar_state="collapsed")

# MINIMALIST PROFESSIONAL STYLING
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    * { font-family: 'Inter', sans-serif; }
    .stApp { background: #f8fafc; }
    .stButton>button {
        background: #334155 !important; color: white !important;
        border: none !important; border-radius: 6px !important;
        padding: 0.5rem 1.2rem !important; font-weight: 600 !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        transition: all 0.2s !important;
    }
    .stButton>button:hover { background: #1e293b !important; transform: translateY(-1px) !important; }
    button[kind="primary"] { background: #0ea5e9 !important; color: white !important; }
    button[kind="primary"]:hover { background: #0284c7 !important; }
    [data-testid="stVerticalBlock"] > div[style*="border"] {
        background: white !important; border-radius: 8px !important;
        padding: 1.5rem !important; box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        border: 1px solid #e2e8f0 !important;
    }
    [data-testid="stMetricValue"] {
        font-size: 1.875rem !important; font-weight: 700 !important; color: #0f172a !important;
    }
    [data-testid="stMetricLabel"] {
        color: #64748b !important; font-weight: 600 !important; font-size: 0.875rem !important;
        text-transform: uppercase !important; letter-spacing: 0.05em !important;
    }
    .stTextInput>div>div>input, .stSelectbox>div>div>select, 
    .stTextArea>div>div>textarea, .stNumberInput>div>div>input,
    .stDateInput>div>div>input, .stTimeInput>div>div>input {
        border: 1px solid #cbd5e1 !important; background: white !important;
        color: #0f172a !important; font-weight: 500 !important; border-radius: 6px !important;
    }
    h1 { color: #0f172a !important; font-weight: 700 !important; }
    h2 { color: #0f172a !important; font-weight: 700 !important; }
    h3 { color: #0f172a !important; font-weight: 600 !important; }
    label { color: #334155 !important; font-weight: 600 !important; font-size: 0.875rem !important; }
    .stSuccess { background: #ecfdf5 !important; color: #065f46 !important; 
                 border-left: 4px solid #10b981 !important; }
    .stInfo { background: #f0f9ff !important; color: #075985 !important; 
              border-left: 4px solid #0ea5e9 !important; }
    .stWarning { background: #fffbeb !important; color: #92400e !important; 
                 border-left: 4px solid #f59e0b !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div style='background: white; padding: 1.25rem; border-radius: 8px; margin-bottom: 1.5rem;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1); border-left: 4px solid #0ea5e9;'>
    <div style='color: #0f172a; font-weight: 700; font-size: 1.125rem; margin-bottom: 0.25rem;'>
        🎭 SANDBOX MODE
    </div>
    <div style='color: #64748b; font-size: 0.875rem; font-weight: 500;'>
        This demonstration uses synthetic data only. No real student information is included.
    </div>
</div>
""", unsafe_allow_html=True)

# MOCK DATA
MOCK_STAFF = [
    {"id": "s1", "name": "Emily Jones", "role": "JP", "email": "emily.jones@example.com", "password": "demo123"},
    {"id": "s2", "name": "Daniel Lee", "role": "PY", "email": "daniel.lee@example.com", "password": "demo123"},
    {"id": "s3", "name": "Sarah Chen", "role": "SY", "email": "sarah.chen@example.com", "password": "demo123"},
    {"id": "s4", "name": "Admin User", "role": "ADM", "email": "admin@example.com", "password": "admin123"},
]

MOCK_STUDENTS = [
    {"id": "stu_jp1", "name": "Emma T.", "grade": "R", "dob": "2018-05-30", "program": "JP"},
    {"id": "stu_jp2", "name": "Oliver S.", "grade": "Y1", "dob": "2017-09-12", "program": "JP"},
    {"id": "stu_jp3", "name": "Sophie M.", "grade": "Y2", "dob": "2016-03-20", "program": "JP"},
    {"id": "stu_py1", "name": "Liam C.", "grade": "Y3", "dob": "2015-06-15", "program": "PY"},
    {"id": "stu_py2", "name": "Ava R.", "grade": "Y4", "dob": "2014-11-08", "program": "PY"},
    {"id": "stu_py3", "name": "Noah B.", "grade": "Y6", "dob": "2012-02-28", "program": "PY"},
    {"id": "stu_sy1", "name": "Isabella G.", "grade": "Y7", "dob": "2011-04-17", "program": "SY"},
    {"id": "stu_sy2", "name": "Ethan D.", "grade": "Y9", "dob": "2009-12-03", "program": "SY"},
    {"id": "stu_sy3", "name": "Mia A.", "grade": "Y11", "dob": "2007-08-20", "program": "SY"},
]

PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}
BEHAVIOUR_TYPES = ["Verbal Refusal", "Elopement", "Property Destruction", "Aggression (Peer)", 
                   "Aggression (Adult)", "Self-Harm", "Verbal Aggression", "Other"]
ANTECEDENTS = ["Requested to transition", "Given instruction/demand", "Peer conflict", 
               "Staff attention shifted", "Unstructured time", "Sensory overload", 
               "Access denied", "Change in routine", "Difficult task"]
INTERVENTIONS = ["CPI Supportive stance", "Offered break", "Reduced demand", "Provided choices", 
                "Removed audience", "Visual supports", "Co-regulation", "Prompted coping skill", "Redirection"]
LOCATIONS = ["JP Classroom", "PY Classroom", "SY Classroom", "Playground", "Library", "Admin", "Gate", "Toilets"]
VALID_PAGES = ["login", "landing", "program_students", "incident_log", "critical_incident", "student_analysis"]

def format_time_12hr(time_str):
    """Convert 24hr time string to 12hr format"""
    try:
        if isinstance(time_str, str):
            dt = datetime.strptime(time_str, "%H:%M:%S")
        else:
            dt = datetime.combine(date.today(), time_str)
        return dt.strftime("%I:%M %p")
    except:
        return time_str

def generate_hypothesis(antecedent, behaviour, consequence):
    """Auto-generate hypothesis based on ABC data"""
    hypotheses = []
    antecedent_lower = antecedent.lower()
    behaviour_lower = behaviour.lower()
    
    if any(word in antecedent_lower for word in ["instruction", "demand", "task", "transition", "work"]):
        hypotheses.append("To avoid or escape the demand/task")
    if any(word in antecedent_lower for word in ["attention", "shifted", "ignored", "alone"]):
        hypotheses.append("To gain staff/peer attention")
    if any(word in antecedent_lower for word in ["sensory", "loud", "noise", "bright", "touch"]):
        hypotheses.append("To escape sensory discomfort or seek sensory input")
    if any(word in antecedent_lower for word in ["denied", "can't have", "no", "wait"]):
        hypotheses.append("To gain access to preferred item/activity")
    if any(word in behaviour_lower for word in ["refusal", "defiance", "left", "ran"]):
        hypotheses.append("To assert control or autonomy")
    
    if not hypotheses:
        hypotheses.append("Function requires further analysis")
    
    return " / ".join(hypotheses[:2])

def generate_admin_summary(incident_data, student, staff_name):
    """Generate brief summary for external incident log - FOR ADMIN USE ONLY"""
    abch_primary = incident_data.get("ABCH_primary", {})
    intended = incident_data.get("intended_outcomes", [])
    
    date_time = incident_data.get("created_at", datetime.now().isoformat())
    dt = datetime.fromisoformat(date_time)
    
    location = abch_primary.get("location", "Unknown location")
    time_str = abch_primary.get("time", "Unknown time")
    behaviour = abch_primary.get("behaviour", "Behaviour not specified")
    context = abch_primary.get("context", "")
    consequence = abch_primary.get("consequence", "")
    hypothesis = abch_primary.get("hypothesis", "")
    
    summary = f"""CRITICAL INCIDENT SUMMARY - FOR ADMIN USE ONLY
    
Student: {student['name']} (Grade {student['grade']})
Date/Time: {dt.strftime('%d/%m/%Y')} at {time_str}
Location: {location}

INCIDENT DESCRIPTION:
{behaviour[:200]}{'...' if len(behaviour) > 200 else ''}

CONTEXT/ANTECEDENT:
{context[:200]}{'...' if len(context) > 200 else ''}

IMMEDIATE STAFF RESPONSE:
{consequence[:200]}{'...' if len(consequence) > 200 else ''}

BEHAVIOURAL FUNCTION:
{hypothesis}

OUTCOMES IMPLEMENTED:
{', '.join(intended[:5]) if intended else 'See full form for details'}

REPORTED BY: {staff_name}
FORM COMPLETED: {dt.strftime('%d/%m/%Y %H:%M')}

** This summary is for external departmental incident log entry purposes only **
** Full critical incident form has been saved and distributed to relevant parties **
"""
    
    return summary

def show_severity_guide():
    st.markdown("""
    <div style='background: white; padding: 1.25rem; border-radius: 8px; margin: 1rem 0; 
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1); border: 1px solid #e2e8f0;'>
        <div style='color: #0f172a; font-weight: 700; margin-bottom: 1rem; font-size: 1rem;'>
            📊 Severity Level Guide (from start to end of incident)
        </div>
        <div style='display: grid; grid-template-columns: repeat(5, 1fr); gap: 0.75rem;'>
            <div style='background: #f8fafc; padding: 1rem; border-radius: 6px; border: 2px solid #cbd5e1;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>1 - Low</div>
                <div style='color: #64748b; font-size: 0.8rem;'>Persistent minor</div>
            </div>
            <div style='background: #f1f5f9; padding: 1rem; border-radius: 6px; border: 2px solid #94a3b8;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>2 - Disruptive</div>
                <div style='color: #64748b; font-size: 0.8rem;'>Impacts others</div>
            </div>
            <div style='background: #e2e8f0; padding: 1rem; border-radius: 6px; border: 2px solid #64748b;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>3 - Concerning</div>
                <div style='color: #475569; font-size: 0.8rem;'>Verbal aggression</div>
            </div>
            <div style='background: #cbd5e1; padding: 1rem; border-radius: 6px; border: 2px solid #475569;'>
                <div style='color: #0f172a; font-weight: 700; margin-bottom: 0.5rem;'>4 - Serious</div>
                <div style='color: #334155; font-size: 0.8rem;'>Physical aggression</div>
            </div>
            <div style='background: #94a3b8; padding: 1rem; border-radius: 6px; border: 2px solid #1e293b;'>
                <div style='color: #fff; font-weight: 700; margin-bottom: 0.5rem;'>5 - Critical</div>
                <div style='color: #f1f5f9; font-size: 0.8rem;'>Severe violence</div>
            </div>
        </div>
        <div style='margin-top: 1rem; padding: 0.75rem; background: #fffbeb; border-radius: 6px; border-left: 4px solid #f59e0b;'>
            <div style='color: #92400e; font-weight: 600; font-size: 0.85rem;'>
                ⚠️ Severity 3 or above requires a Critical Incident ABCH Form
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def send_critical_incident_email(incident_data, student, staff_email, leader_email, admin_email):
    """Send email notification to all parties"""
    st.info(f"""📧 **Email Notification Sent**
    
**To:** {leader_email}, {admin_email}, {staff_email}  
**Subject:** CRITICAL INCIDENT - {student['name']}

**Student:** {student['name']} | **Programme:** {student['program']} | **Grade:** {student['grade']}  

Critical Incident Form completed and saved.
Admin summary included for departmental log.

*(In production, this sends via SMTP)*
    """)


def generate_behaviour_analysis_plan_docx(student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level):
    """Generate comprehensive BAP with embedded graphs - Berry Street & Learning and Behaviour Unit"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import plotly.graph_objects as go
        
        doc = Document()
        
        # TITLE PAGE
        title = doc.add_heading('Behaviour Analysis Plan', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph('Evidence-Based Analysis & Recommendations')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 116, 139)
        
        doc.add_paragraph()
        
        branding = doc.add_paragraph('Prepared by: Learning and Behaviour Unit')
        branding.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in branding.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(14, 165, 233)
        
        doc.add_page_break()
        
        # STUDENT INFO
        doc.add_heading('Student Information', 1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.rows[0].cells[0].text = 'Student:'
        info_table.rows[0].cells[1].text = student['name']
        info_table.rows[1].cells[0].text = 'Program:'
        info_table.rows[1].cells[1].text = PROGRAM_NAMES.get(student['program'], student['program'])
        info_table.rows[2].cells[0].text = 'Grade:'
        info_table.rows[2].cells[1].text = student['grade']
        info_table.rows[3].cells[0].text = 'Analysis Date:'
        info_table.rows[3].cells[1].text = datetime.now().strftime('%d %B %Y')
        info_table.rows[4].cells[0].text = 'Data Period:'
        info_table.rows[4].cells[1].text = f"{full_df['date_parsed'].min().strftime('%d/%m/%Y')} - {full_df['date_parsed'].max().strftime('%d/%m/%Y')}"
        
        doc.add_paragraph()
        
        # EXECUTIVE SUMMARY
        doc.add_heading('Executive Summary', 1)
        summary = doc.add_paragraph()
        summary.add_run('Total Incidents: ').bold = True
        summary.add_run(f"{len(full_df)}\n")
        summary.add_run('Critical Incidents: ').bold = True
        summary.add_run(f"{len(full_df[full_df['incident_type'] == 'Critical'])}\n")
        summary.add_run('Average Severity: ').bold = True
        summary.add_run(f"{full_df['severity'].mean():.2f}/5\n")
        summary.add_run('Risk Level: ').bold = True
        summary.add_run(f"{risk_level} ({risk_score}/100)")
        
        doc.add_paragraph()
        
        # KEY FINDINGS
        doc.add_heading('Key Findings', 1)
        findings = doc.add_paragraph()
        findings.add_run('Primary Behaviour: ').bold = True
        findings.add_run(f"{top_beh}\n\n")
        findings.add_run('Most Common Trigger: ').bold = True
        findings.add_run(f"{top_ant}\n\n")
        findings.add_run('Hotspot Location: ').bold = True
        findings.add_run(f"{top_loc}\n\n")
        findings.add_run('Peak Time: ').bold = True
        findings.add_run(f"{top_session}")
        
        doc.add_page_break()
        
        # VISUAL ANALYTICS
        doc.add_heading('Visual Analytics', 1)
        doc.add_paragraph('The following graphs provide visual representation of incident patterns and trends.')
        
        # GRAPH 1: Daily Frequency
        doc.add_heading('1. Daily Incident Frequency', 2)
        daily = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig1 = go.Figure()
        fig1.add_trace(go.Scatter(
            x=daily["date_parsed"], y=daily["count"],
            mode='lines+markers', line=dict(color='#334155', width=3),
            marker=dict(size=8), fill='tozeroy', fillcolor='rgba(51, 65, 85, 0.15)'
        ))
        fig1.update_layout(
            width=700, height=350, showlegend=False,
            plot_bgcolor='white', paper_bgcolor='white',
            margin=dict(l=60, r=60, t=40, b=60),
            xaxis_title="Date", yaxis_title="Incident Count",
            font=dict(size=12)
        )
        img1 = "/tmp/graph1_daily.png"
        fig1.write_image(img1, width=700, height=350, scale=2)
        doc.add_picture(img1, width=Inches(6))
        doc.add_paragraph("Daily incident frequency shows temporal patterns and potential correlations with external factors.")
        doc.add_paragraph()
        
        # GRAPH 2: Top Behaviours
        doc.add_heading('2. Most Common Behaviours', 2)
        beh_counts = full_df["behaviour_type"].value_counts().head(5)
        fig2 = go.Figure()
        fig2.add_trace(go.Bar(
            y=beh_counts.index, x=beh_counts.values,
            orientation='h', marker=dict(color='#334155'),
            text=beh_counts.values, textposition='outside',
            textfont=dict(size=14)
        ))
        fig2.update_layout(
            width=700, height=350, showlegend=False,
            plot_bgcolor='white', paper_bgcolor='white',
            margin=dict(l=150, r=60, t=40, b=60),
            xaxis_title="Frequency", font=dict(size=12)
        )
        img2 = "/tmp/graph2_behaviours.png"
        fig2.write_image(img2, width=700, height=350, scale=2)
        doc.add_picture(img2, width=Inches(6))
        doc.add_paragraph(f"Primary behaviour: {beh_counts.index[0]} ({beh_counts.values[0]} incidents). Intervention planning should prioritize this behaviour.")
        doc.add_paragraph()
        
        # GRAPH 3: Top Triggers
        doc.add_heading('3. Most Common Triggers', 2)
        ant_counts = full_df["antecedent"].value_counts().head(5)
        fig3 = go.Figure()
        fig3.add_trace(go.Bar(
            y=ant_counts.index, x=ant_counts.values,
            orientation='h', marker=dict(color='#475569'),
            text=ant_counts.values, textposition='outside',
            textfont=dict(size=14)
        ))
        fig3.update_layout(
            width=700, height=350, showlegend=False,
            plot_bgcolor='white', paper_bgcolor='white',
            margin=dict(l=150, r=60, t=40, b=60),
            xaxis_title="Frequency", font=dict(size=12)
        )
        img3 = "/tmp/graph3_triggers.png"
        fig3.write_image(img3, width=700, height=350, scale=2)
        doc.add_picture(img3, width=Inches(6))
        doc.add_paragraph(f"Key trigger: {ant_counts.index[0]}. Proactive strategies should address this antecedent.")
        doc.add_paragraph()
        
        # GRAPH 4: Severity Trend
        doc.add_heading('4. Severity Over Time', 2)
        fig4 = go.Figure()
        fig4.add_trace(go.Scatter(
            x=full_df["date_parsed"], y=full_df["severity"],
            mode='markers', marker=dict(size=10, color='#334155', opacity=0.6)
        ))
        if len(full_df) >= 2:
            z = np.polyfit(range(len(full_df)), full_df["severity"], 1)
            p = np.poly1d(z)
            fig4.add_trace(go.Scatter(
                x=full_df["date_parsed"], y=p(range(len(full_df))),
                mode='lines', line=dict(color='#94a3b8', width=3, dash='dash'),
                name='Trend'
            ))
        fig4.update_layout(
            width=700, height=350, yaxis=dict(range=[0, 6]),
            plot_bgcolor='white', paper_bgcolor='white',
            margin=dict(l=60, r=60, t=40, b=60),
            xaxis_title="Date", yaxis_title="Severity (1-5)",
            font=dict(size=12), showlegend=False
        )
        img4 = "/tmp/graph4_severity.png"
        fig4.write_image(img4, width=700, height=350, scale=2)
        doc.add_picture(img4, width=Inches(6))
        doc.add_paragraph("Severity trend analysis indicates pattern trajectory. Increasing trend requires immediate intervention adjustment.")
        doc.add_paragraph()
        
        # GRAPH 5: Location Hotspots
        doc.add_heading('5. Location Hotspots', 2)
        loc_counts = full_df["location"].value_counts().head(5)
        fig5 = go.Figure()
        fig5.add_trace(go.Bar(
            y=loc_counts.index, x=loc_counts.values,
            orientation='h', marker=dict(color='#64748b'),
            text=loc_counts.values, textposition='outside',
            textfont=dict(size=14)
        ))
        fig5.update_layout(
            width=700, height=350, showlegend=False,
            plot_bgcolor='white', paper_bgcolor='white',
            margin=dict(l=150, r=60, t=40, b=60),
            xaxis_title="Frequency", font=dict(size=12)
        )
        img5 = "/tmp/graph5_locations.png"
        fig5.write_image(img5, width=700, height=350, scale=2)
        doc.add_picture(img5, width=Inches(6))
        doc.add_paragraph(f"Most incidents occur in: {loc_counts.index[0]}. Environmental modifications and increased support recommended.")
        doc.add_paragraph()
        
        # GRAPH 6: Time of Day
        doc.add_heading('6. Time of Day Patterns', 2)
        session_counts = full_df["session"].value_counts()
        fig6 = go.Figure()
        fig6.add_trace(go.Bar(
            x=session_counts.index, y=session_counts.values,
            marker=dict(color='#475569'),
            text=session_counts.values, textposition='outside',
            textfont=dict(size=14)
        ))
        fig6.update_layout(
            width=700, height=350, showlegend=False,
            plot_bgcolor='white', paper_bgcolor='white',
            margin=dict(l=60, r=60, t=40, b=60),
            yaxis_title="Frequency", font=dict(size=12)
        )
        img6 = "/tmp/graph6_time.png"
        fig6.write_image(img6, width=700, height=350, scale=2)
        doc.add_picture(img6, width=Inches(6))
        doc.add_paragraph(f"Peak incident time: {session_counts.index[0]}. Schedule proactive regulation supports before this period.")
        
        doc.add_page_break()
        
        # CLINICAL INTERPRETATION
        doc.add_heading('Clinical Interpretation', 1)
        doc.add_paragraph('Based on Applied Behaviour Analysis (ABA), Trauma-Informed Practice, Berry Street Education Model, and CPI principles:')
        
        interp = doc.add_paragraph()
        interp.add_run('Pattern Analysis: ').bold = True
        interp.add_run(f"Data indicates {student['name']} is most vulnerable when '{top_ant}' occurs in {top_loc} during {top_session}. ")
        interp.add_run("This behaviour pattern serves as a safety strategy and communication method.\n\n")
        
        interp.add_run('Trauma-Informed & Berry Street Lens: ').bold = True
        interp.add_run("Behaviours represent adaptive responses to perceived threat. The student's nervous system is responding to environmental cues. ")
        interp.add_run("Berry Street Education Model emphasizes strengthening Body (self-regulation and wellbeing), Relationship (positive connections), ")
        interp.add_run("Stamina (persistence and engagement), Engagement (learning readiness), and Character (values and agency). ")
        interp.add_run("Focus on Body and Relationship domains first to build foundation for learning.\n\n")
        
        interp.add_run('CPI Alignment: ').bold = True
        interp.add_run("Crisis Prevention Institute principles emphasize Supportive Stance, understanding behaviour as communication, ")
        interp.add_run("and maintaining dignity throughout the intervention process. Use non-restrictive approaches.")
        
        doc.add_paragraph()
        
        # EVIDENCE-BASED RECOMMENDATIONS
        doc.add_heading('Evidence-Based Recommendations', 1)
        
        doc.add_heading('1. Proactive Strategies (Prevention) - Berry Street Body Domain', 2)
        doc.add_paragraph(f"• Regulated start to {top_session}: breathing, movement, sensory breaks", style='List Bullet')
        doc.add_paragraph(f"• Visual check-in before '{top_ant}' occurs", style='List Bullet')
        doc.add_paragraph(f"• Environmental modification in {top_loc}", style='List Bullet')
        doc.add_paragraph("• Predictable routines with visual supports", style='List Bullet')
        doc.add_paragraph("• Sensory regulation opportunities (zones of regulation)", style='List Bullet')
        
        doc.add_heading('2. Co-Regulation Strategies (CPI-Aligned) - Berry Street Relationship Domain', 2)
        doc.add_paragraph("• Maintain Supportive Stance: low, slow voice; non-threatening posture", style='List Bullet')
        doc.add_paragraph("• Reduce audience and environmental stimulation", style='List Bullet')
        doc.add_paragraph("• One key adult maintains connection (relationship is foundation)", style='List Bullet')
        doc.add_paragraph("• Acknowledge feelings: 'I can see you're feeling...'", style='List Bullet')
        doc.add_paragraph("• Offer choices to restore sense of control and agency", style='List Bullet')
        
        doc.add_heading('3. Teaching Replacement Skills - Berry Street Stamina & Character Domains', 2)
        doc.add_paragraph("• Link to Personal & Social Capability curriculum", style='List Bullet')
        doc.add_paragraph("• Teach help-seeking routines with visual cues", style='List Bullet')
        doc.add_paragraph("• Practice requesting breaks before escalation (self-advocacy)", style='List Bullet')
        doc.add_paragraph("• Emotional literacy: naming and understanding feelings", style='List Bullet')
        doc.add_paragraph("• Build persistence and coping strategies (Stamina domain)", style='List Bullet')
        
        doc.add_heading('4. SMART Goal', 2)
        goal = doc.add_paragraph()
        goal.add_run('Measurable Outcome: ').bold = True
        goal.add_run("Over the next 5 weeks, student will use taught help-seeking strategy ")
        goal.add_run("(break card/verbal request) in 4 out of 5 opportunities when experiencing ")
        goal.add_run("escalation triggers, with staff prompting as needed. This supports Berry Street Relationship and Stamina domains.")
        
        doc.add_paragraph()
        doc.add_paragraph('Review Date: ' + (datetime.now() + timedelta(weeks=5)).strftime('%d %B %Y'))
        
        doc.add_page_break()
        
        # FOOTER
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('\n\nPrepared by Learning and Behaviour Unit\n')
        footer_run.font.size = Pt(10)
        footer_run.font.bold = True
        footer_run.font.color.rgb = RGBColor(14, 165, 233)
        
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer2_run = footer2.add_run('Evidence-based analysis using ABA, Trauma-Informed, Berry Street, and CPI principles\n')
        footer2_run.font.size = Pt(9)
        footer2_run.font.color.rgb = RGBColor(100, 116, 139)
        
        footer3 = doc.add_paragraph()
        footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer3_run = footer3.add_run(datetime.now().strftime('%d %B %Y'))
        footer3_run.font.size = Pt(9)
        footer3_run.font.color.rgb = RGBColor(100, 116, 139)
        
        # Save to bytes
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
        
    except Exception as e:
        st.error(f"Error generating BAP: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None


def init_state():
    ss = st.session_state
    if "logged_in" not in ss: ss.logged_in = False
    if "current_user" not in ss: ss.current_user = None
    if "current_page" not in ss: ss.current_page = "login"
    if "students" not in ss: ss.students = MOCK_STUDENTS
    if "staff" not in ss: ss.staff = MOCK_STAFF
    if "incidents" not in ss: ss.incidents = generate_mock_incidents(70)
    if "critical_incidents" not in ss: ss.critical_incidents = []
    if "selected_program" not in ss: ss.selected_program = "JP"
    if "selected_student_id" not in ss: ss.selected_student_id = None
    if "current_incident_id" not in ss: ss.current_incident_id = None
    if "abch_rows" not in ss: ss.abch_rows = []
    if "show_critical_prompt" not in ss: ss.show_critical_prompt = False

def login_user(email: str, password: str) -> bool:
    email = (email or "").strip().lower()
    password = (password or "").strip()
    if not email or not password: return False
    for staff in st.session_state.staff:
        if staff.get("email", "").lower() == email and staff.get("password", "") == password:
            st.session_state.logged_in = True
            st.session_state.current_user = staff
            st.session_state.current_page = "landing"
            return True
    return False

def go_to(page: str, **kwargs):
    if page not in VALID_PAGES: return
    st.session_state.current_page = page
    for k, v in kwargs.items():
        setattr(st.session_state, k, v)
    st.rerun()

def get_student(sid): 
    return next((s for s in st.session_state.students if s["id"] == sid), None)

def get_session_from_time(t): 
    return "Morning" if t.hour < 11 else "Middle" if t.hour < 13 else "Afternoon"

def generate_mock_incidents(n=70):
    incidents = []
    weights = {"stu_sy1": 12, "stu_py1": 10, "stu_sy2": 9, "stu_jp1": 8, "stu_py2": 7}
    pool = []
    for stu in MOCK_STUDENTS:
        pool.extend([stu] * weights.get(stu["id"], 5))
    for _ in range(n):
        stu = random.choice(pool)
        sev = random.choices([1, 2, 3, 4, 5], weights=[20, 35, 25, 15, 5])[0]
        dt = datetime.now() - timedelta(days=random.randint(0, 90))
        dt = dt.replace(hour=random.choices([9,10,11,12,13,14,15], weights=[10,15,12,8,12,18,10])[0], 
                       minute=random.randint(0,59), second=0)
        incidents.append({
            "id": str(uuid.uuid4()), "student_id": stu["id"], "student_name": stu["name"],
            "date": dt.date().isoformat(), "time": dt.time().strftime("%H:%M:%S"),
            "day": dt.strftime("%A"), "session": get_session_from_time(dt.time()),
            "location": random.choice(LOCATIONS), "behaviour_type": random.choice(BEHAVIOUR_TYPES),
            "antecedent": random.choice(ANTECEDENTS), 
            "intervention": [random.choice(INTERVENTIONS)],  # Changed to list
            "severity": sev, "reported_by": random.choice(MOCK_STAFF)["name"],
            "description": "Mock incident", "is_critical": sev >= 3, "duration_minutes": random.randint(2, 25)
        })
    return incidents

# PAGES
def render_login_page():
    st.markdown("## 🔐 Staff Login")
    with st.container(border=True):
        st.markdown("**Demo Credentials:**")
        st.code("Email: emily.jones@example.com\nPassword: demo123")
    email = st.text_input("Email Address", placeholder="your.email@example.com", key="login_email")
    password = st.text_input("Password", type="password", placeholder="Enter password", key="login_pass")
    if st.button("Login", type="primary", use_container_width=True):
        if login_user(email, password):
            st.success(f"Welcome {st.session_state.current_user['name']}!")
            st.rerun()
        else:
            st.error("Invalid credentials")

def render_landing_page():
    user = st.session_state.current_user or {}
    st.markdown(f"### 👋 Welcome, {user.get('name', 'User')}")
    
    col1, col2 = st.columns([6, 1])
    with col2:
        if st.button("Logout", key="logout_btn"):
            st.session_state.logged_in = False
            st.session_state.current_page = "login"
            st.rerun()
    
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1: st.metric("Students", len(st.session_state.students))
    with col2: st.metric("Total Incidents", len(st.session_state.incidents))
    with col3: st.metric("Critical", len([i for i in st.session_state.incidents if i.get("is_critical")]))
    
    st.markdown("---")
    st.markdown("### Select Program")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Junior Primary**")
        if st.button("Enter JP", use_container_width=True, type="primary", key="btn_jp"):
            go_to("program_students", selected_program="JP")
    with col2:
        st.markdown("**Primary Years**")
        if st.button("Enter PY", use_container_width=True, type="primary", key="btn_py"):
            go_to("program_students", selected_program="PY")
    with col3:
        st.markdown("**Senior Years**")
        if st.button("Enter SY", use_container_width=True, type="primary", key="btn_sy"):
            go_to("program_students", selected_program="SY")

def render_program_students_page():
    program = st.session_state.get("selected_program", "JP")
    st.markdown(f"## {PROGRAM_NAMES.get(program)} — Students")
    
    col1, col2 = st.columns([6, 1])
    with col1:
        if st.button("⬅ Back to Landing", key="back_students"):
            go_to("landing")
    
    students = [s for s in st.session_state.students if s["program"] == program]
    for stu in students:
        stu_incidents = [i for i in st.session_state.incidents if i["student_id"] == stu["id"]]
        with st.container(border=True):
            col1, col2, col3 = st.columns([3, 2, 2])
            with col1:
                st.markdown(f"**{stu['name']}**")
                st.caption(f"Grade {stu['grade']}")
            with col2:
                st.metric("Incidents", len(stu_incidents))
            with col3:
                if st.button("📝 Log", key=f"log_{stu['id']}", use_container_width=True):
                    go_to("incident_log", selected_student_id=stu["id"])
                if st.button("📊 Analysis", key=f"ana_{stu['id']}", use_container_width=True):
                    go_to("student_analysis", selected_student_id=stu["id"])

def render_incident_log_page():
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## 📝 Incident Log — {student['name']}")
    
    # Navigation buttons
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("⬅ Back to Students", key="back_log_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_log"):
            go_to("landing")
    
    show_severity_guide()
    
    # Check if critical form is required
    if st.session_state.show_critical_prompt:
        inc_info = st.session_state.get("last_incident_info", {})
        if inc_info.get("severity", 0) >= 3:
            st.error(f"⚠️ **Severity {inc_info['severity']} Detected** - Critical Incident ABCH Form Required")
        else:
            st.error("⚠️ **Critical Incident Flagged** - Critical Incident ABCH Form Required")
        st.info("Please complete the Critical Incident ABCH form to document this event fully.")
        
        # Only one button - REMOVED "Skip for Now"
        if st.button("📋 Complete Critical Form Now", type="primary", key="crit_now", use_container_width=True):
            st.session_state.show_critical_prompt = False
            go_to("critical_incident", current_incident_id=st.session_state.current_incident_id)
        st.markdown("---")
        st.stop()
    
    # INCIDENT FORM
    with st.form("incident_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            inc_date = st.date_input("Date *", date.today(), key="inc_date")
            inc_time = st.time_input("Time *", datetime.now().time(), key="inc_time")
            location = st.selectbox("Location *", [""] + LOCATIONS, key="inc_loc")
        with col2:
            behaviour = st.selectbox("Behaviour Type *", [""] + BEHAVIOUR_TYPES, key="inc_beh")
            antecedent = st.selectbox("Antecedent/Trigger *", [""] + ANTECEDENTS, key="inc_ant")
            # MULTIPLE INTERVENTIONS
            interventions = st.multiselect("Interventions Used *", INTERVENTIONS, key="inc_ints")
        
        duration = st.number_input("Duration (minutes) *", min_value=1, value=1, key="inc_dur")
        severity = st.slider("Severity Level (from start to end of incident) *", 1, 5, 1, key="inc_sev")
        description = st.text_area("Brief Description (Optional)", placeholder="Factual, objective description...", key="inc_desc")
        manual_critical = st.checkbox("This incident requires a Critical Incident ABCH Form (regardless of severity)", key="manual_crit")
        submitted = st.form_submit_button("Submit Incident", type="primary")
    
    if submitted:
        if not location or not behaviour or not antecedent or not interventions:
            st.error("Please complete all required fields marked with *")
        else:
            new_id = str(uuid.uuid4())
            is_critical = (severity >= 3) or manual_critical
            rec = {
                "id": new_id, "student_id": student_id, "student_name": student["name"],
                "date": inc_date.isoformat(), "time": inc_time.strftime("%H:%M:%S"),
                "day": inc_date.strftime("%A"), "session": get_session_from_time(inc_time),
                "location": location, "behaviour_type": behaviour, "antecedent": antecedent,
                "intervention": interventions,  # Save as list
                "severity": severity,
                "reported_by": st.session_state.current_user["name"],
                "duration_minutes": duration, "description": description or "", 
                "is_critical": is_critical
            }
            st.session_state.incidents.append(rec)
            st.success("✅ Incident logged successfully")
            
            if is_critical:
                st.session_state.current_incident_id = new_id
                st.session_state.show_critical_prompt = True
                st.session_state.last_incident_info = {"severity": severity, "manual": manual_critical}
                st.rerun()
            else:
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("↩️ Back to Students", key="back_after_log"):
                        go_to("program_students", selected_program=student["program"])
                with col2:
                    if st.button("🏠 Program Landing", key="home_after_log"):
                        go_to("landing")


def render_critical_incident_page():
    """Critical Incident Form with Police Reference and Improved Labels"""
    inc_id = st.session_state.get("current_incident_id")
    quick_inc = next((i for i in st.session_state.incidents if i["id"] == inc_id), None)
    
    if not quick_inc:
        st.error("No incident found")
        return
    
    student = get_student(quick_inc["student_id"])
    st.markdown(f"## 🚨 Critical Incident ABCH Form")
    
    # Navigation
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("⬅ Back to Students", key="back_crit_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_crit"):
            go_to("landing")
    
    st.markdown("### Incident Details (from Quick Log)")
    with st.container(border=True):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"**Student:** {student['name']}")
            st.markdown(f"**Grade:** {student['grade']}")
        with col2:
            st.markdown(f"**Date:** {quick_inc['date']}")
            st.markdown(f"**Time:** {format_time_12hr(quick_inc['time'])}")
        with col3:
            st.markdown(f"**Location:** {quick_inc['location']}")
            st.markdown(f"**Session:** {quick_inc['session']}")
        with col4:
            st.markdown(f"**Severity:** {quick_inc['severity']}")
            st.markdown(f"**Behaviour:** {quick_inc['behaviour_type']}")
    
    st.markdown("---")
    st.markdown("### ABCH Chronology")
    st.caption("Document the sequence of events. Add continuation entries if incident was prolonged.")
    
    if "abch_rows" not in st.session_state:
        st.session_state.abch_rows = []
    
    # PRIMARY ROW
    st.markdown("#### Initial Incident")
    col_header = st.columns([2, 2, 2, 2, 2])
    with col_header[0]: st.markdown("**ANTECEDENT (Triggers)**")
    with col_header[2]: st.markdown("**BEHAVIOUR**")
    with col_header[4]: st.markdown("**CONSEQUENCES**")
    
    col_subheader = st.columns([1, 1, 1, 1, 2, 2])
    with col_subheader[0]: st.caption("Location")
    with col_subheader[1]: st.caption("Context (what was happening?)")
    with col_subheader[2]: st.caption("Time")
    with col_subheader[3]: st.caption("Observed Behaviour")
    with col_subheader[4]: st.caption("What happened after?")
    with col_subheader[5]: st.caption("HYPOTHESIS (Function)")
    
    col_inputs1 = st.columns([1, 1, 1, 1, 2, 2])
    
    with col_inputs1[0]:
        location_1 = st.text_input("", value=quick_inc['location'], key="loc_1", label_visibility="collapsed")
    with col_inputs1[1]:
        context_1 = st.text_area("", placeholder="What was going on before?", 
                                key="context_1", height=100, label_visibility="collapsed")
    with col_inputs1[2]:
        time_1 = st.text_input("", value=format_time_12hr(quick_inc['time']), key="time_1", label_visibility="collapsed")
    with col_inputs1[3]:
        behaviour_1 = st.text_area("", placeholder="What did student do?", 
                                  key="behaviour_1", height=100, label_visibility="collapsed")
    with col_inputs1[4]:
        consequence_1 = st.text_area("", placeholder="Staff response? Student reaction?", 
                                    key="consequence_1", height=100, label_visibility="collapsed")
    with col_inputs1[5]:
        if "hyp_1_generated" not in st.session_state:
            st.session_state.hyp_1_generated = False
        if not st.session_state.hyp_1_generated and context_1 and behaviour_1:
            auto_hyp = generate_hypothesis(context_1, behaviour_1, consequence_1)
            st.session_state.hyp_1_auto = auto_hyp
            st.session_state.hyp_1_generated = True
        hypothesis_1 = st.text_area("", 
                                    value=st.session_state.get("hyp_1_auto", ""),
                                    placeholder="Auto-generated (editable)", 
                                    key="hypothesis_1", height=100, label_visibility="collapsed")
    
    st.markdown("---")
    
    if st.button("➕ Add Continuation Entry", key="add_abch_row"):
        st.session_state.abch_rows.append({})
        st.rerun()
    
    # ADDITIONAL ROWS - Changed labels
    for idx, row in enumerate(st.session_state.abch_rows):
        st.markdown(f"#### Continuation Entry {idx + 1}")
        col_add = st.columns([1, 1, 1, 1, 2, 2])
        with col_add[0]:
            row["location"] = st.text_input("", key=f"loc_{idx+2}", label_visibility="collapsed")
        with col_add[1]:
            row["context"] = st.text_area("", key=f"context_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[2]:
            row["time"] = st.text_input("", key=f"time_{idx+2}", label_visibility="collapsed")
        with col_add[3]:
            row["behaviour"] = st.text_area("", key=f"behaviour_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[4]:
            row["consequence"] = st.text_area("", key=f"consequence_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[5]:
            if row.get("context") and row.get("behaviour"):
                auto_hyp_add = generate_hypothesis(row["context"], row["behaviour"], row.get("consequence", ""))
                row["hypothesis"] = st.text_area("", value=auto_hyp_add, key=f"hypothesis_{idx+2}", height=100, label_visibility="collapsed")
            else:
                row["hypothesis"] = st.text_area("", key=f"hypothesis_{idx+2}", height=100, label_visibility="collapsed")
        st.markdown("---")
    
    # INTENDED OUTCOMES with SAPOL reference field
    st.markdown("### Intended Outcomes")
    outcomes_options = [
        "Send Home", "Parent/Caregiver notified via Phone Call",
        "Student Leaving supervised areas/leaving school grounds",
        "Sexualised behaviour", "Incident – student to student",
        "Complaint by co-located school/member of public",
        "Property damage", "Stealing", "Toileting issue",
        "ED155: Staff Injury", "ED155: Student injury",
        "Emergency services - SAPOL",
        "Emergency services - SA Ambulance",
        "Incident Internally Managed - Restorative Session",
        "Incident Internally Managed - Community Service",
        "Incident Internally Managed - Re-Entry",
        "Incident Internally Managed - Case Review",
        "Incident Internally Managed - Make-up Time"
    ]
    
    selected_outcomes = st.multiselect("Select all intended outcomes:", outcomes_options, key="intended_outcomes")
    
    # SAPOL Reference Field - triggered if SAPOL selected
    sapol_reference = ""
    if "Emergency services - SAPOL" in selected_outcomes:
        st.warning("⚠️ SAPOL involvement detected - Police Reference Number required")
        sapol_reference = st.text_input("SAPOL Police Reference Number *", 
                                       placeholder="Enter police reference number",
                                       key="sapol_ref")
    
    tac_notes = st.text_area("Additional Outcome Notes (e.g., TAC meeting):", 
                            placeholder="A TAC meeting will be held...",
                            key="tac_notes", height=100)
    
    st.markdown("---")
    st.markdown("### Notifications & Administration")
    col_notif1, col_notif2 = st.columns(2)
    with col_notif1:
        notified_line_manager = st.checkbox("Notified Line Manager", key="notif_manager", value=True)
        notified_parent = st.checkbox("Notified Parent/Caregiver", key="notif_parent")
    with col_notif2:
        copy_in_file = st.checkbox("Copy in student file", key="copy_file", value=True)
        safety_plan_review = st.checkbox("Safety plan review required", key="safety_review")
    
    st.markdown("---")
    st.markdown("### Staff Agreement")
    staff_name = st.session_state.current_user.get("name", "Staff Member")
    st.markdown(f"**Completing Staff Member:** {staff_name}")
    staff_agrees = st.checkbox(f"✓ I, {staff_name}, confirm this information is accurate and complete.", 
                               key="staff_agrees")
    
    st.markdown("---")
    st.markdown("### Email Distribution")
    col_email1, col_email2 = st.columns(2)
    with col_email1:
        leader_email = st.text_input("Line Manager Email *", 
                                     value="manager@clc.sa.edu.au",
                                     key="leader_email")
    with col_email2:
        admin_email = st.text_input("Admin Email *", 
                                    value="admin@clc.sa.edu.au",
                                    key="admin_email")
    
    st.markdown("---")
    
    if st.button("📧 Submit Critical Incident Form", type="primary", use_container_width=True, key="save_crit"):
        # Validation
        errors = []
        if not context_1 or not behaviour_1 or not consequence_1 or not hypothesis_1:
            errors.append("Please complete all ABCH fields for initial incident")
        if not staff_agrees:
            errors.append("Please confirm your agreement")
        if not leader_email or "@" not in leader_email or not admin_email or "@" not in admin_email:
            errors.append("Please enter valid email addresses")
        if "Emergency services - SAPOL" in selected_outcomes and not sapol_reference:
            errors.append("SAPOL Reference Number is required when SAPOL is involved")
        
        if errors:
            for error in errors:
                st.error(f"❌ {error}")
        else:
            record = {
                "id": str(uuid.uuid4()),
                "created_at": datetime.now().isoformat(),
                "quick_incident_id": inc_id,
                "student_id": quick_inc["student_id"],
                "student_name": student["name"],
                "incident_type": "Critical",
                "ABCH_primary": {
                    "location": location_1,
                    "context": context_1,
                    "time": time_1,
                    "behaviour": behaviour_1,
                    "consequence": consequence_1,
                    "hypothesis": hypothesis_1
                },
                "ABCH_additional": st.session_state.abch_rows.copy(),
                "intended_outcomes": selected_outcomes,
                "sapol_reference": sapol_reference if sapol_reference else None,
                "tac_notes": tac_notes,
                "notifications": {
                    "line_manager": notified_line_manager,
                    "parent": notified_parent,
                    "copy_in_file": copy_in_file,
                    "safety_plan_review": safety_plan_review
                },
                "staff_agreement": {
                    "staff_name": staff_name,
                    "agreed": staff_agrees,
                    "timestamp": datetime.now().isoformat()
                },
                "leader_email": leader_email,
                "admin_email": admin_email
            }
            
            st.session_state.critical_incidents.append(record)
            st.session_state.abch_rows = []
            st.session_state.hyp_1_generated = False
            
            st.success("✅ Critical incident form saved to database")
            
            # GENERATE ADMIN SUMMARY
            admin_summary = generate_admin_summary(record, student, staff_name)
            
            # Show admin summary
            with st.expander("📋 ADMIN SUMMARY (For External Incident Log)", expanded=True):
                st.text_area("Copy this summary for departmental log:", admin_summary, height=400, key="admin_summary_display")
                st.download_button(
                    "📥 Download Admin Summary",
                    admin_summary,
                    file_name=f"Admin_Summary_{student['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain"
                )
            
            # SEND EMAILS
            staff_email = st.session_state.current_user.get("email", "staff@example.com")
            send_critical_incident_email(record, student, staff_email, leader_email, admin_email)
            
            st.markdown("---")
            st.info("✉️ Emails sent to Line Manager, Admin, and completing staff member")
            st.info("💾 Critical incident data saved in student's file")
            st.info("📋 Admin summary generated for external log")
            if sapol_reference:
                st.info(f"🚔 SAPOL Reference: {sapol_reference}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("📊 View Analysis", type="primary", use_container_width=True, key="view_analysis"):
                    go_to("student_analysis", selected_student_id=quick_inc["student_id"])
            with col2:
                if st.button("↩️ Back to Students", use_container_width=True, key="back_crit_after"):
                    go_to("program_students", selected_program=student["program"])
            with col3:
                if st.button("🏠 Program Landing", use_container_width=True, key="home_crit_after"):
                    go_to("landing")


def render_student_analysis_page():
    """Comprehensive Data Analysis with Berry Street Education Model"""
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## 📊 Comprehensive Behaviour Analysis — {student['name']}")
    st.caption("Evidence-based analysis prepared by Learning and Behaviour Unit")
    st.caption("Using ABA, Trauma-Informed Practice, Berry Street Education Model, and CPI principles")
    
    # Navigation
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("⬅ Back to Students", key="back_analysis_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_analysis"):
            go_to("landing")
    
    quick = [i for i in st.session_state.incidents if i["student_id"] == student_id]
    crit = [c for c in st.session_state.critical_incidents if c["student_id"] == student_id]
    
    if not quick and not crit:
        st.info("No incident data available yet.")
        if st.button("↩️ Back", key="back_no_data"):
            go_to("program_students", selected_program=student["program"])
        return
    
    # Build dataframe
    quick_df = pd.DataFrame(quick) if quick else pd.DataFrame()
    crit_df = pd.DataFrame(crit) if crit else pd.DataFrame()
    
    if not quick_df.empty:
        quick_df["incident_type"] = "Quick"
        quick_df["date_parsed"] = pd.to_datetime(quick_df["date"])
        # Handle intervention as list
        if "intervention" in quick_df.columns:
            quick_df["intervention_str"] = quick_df["intervention"].apply(lambda x: ", ".join(x) if isinstance(x, list) else str(x))
    
    if not crit_df.empty:
        crit_df["incident_type"] = "Critical"
        crit_df["date_parsed"] = pd.to_datetime(crit_df.get("created_at", datetime.now().isoformat()))
        crit_df["severity"] = 5
        crit_df["antecedent"] = crit_df["ABCH_primary"].apply(lambda d: d.get("context","") if isinstance(d, dict) else "")
        crit_df["behaviour_type"] = crit_df["ABCH_primary"].apply(lambda d: d.get("behaviour","") if isinstance(d, dict) else "")
    
    full_df = pd.concat([quick_df, crit_df], ignore_index=True).sort_values("date_parsed")
    full_df["hour"] = pd.to_datetime(full_df["time"], format="%H:%M:%S", errors="coerce").dt.hour
    full_df["day_of_week"] = full_df["date_parsed"].dt.day_name()
    
    # OVERVIEW
    st.markdown("### 📈 Executive Summary")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1: st.metric("Total", len(full_df))
    with col2: st.metric("Critical", len(full_df[full_df["incident_type"] == "Critical"]))
    with col3: st.metric("Avg Severity", f"{full_df['severity'].mean():.1f}")
    with col4:
        days = max((full_df["date_parsed"].max() - full_df["date_parsed"].min()).days, 1)
        st.metric("Days Span", days)
    with col5:
        st.metric("Per Day", f"{len(full_df) / days:.1f}")
    
    st.markdown("---")
    
    # GRAPH 1: Daily Frequency
    st.markdown("### 📅 Daily Incident Frequency")
    daily = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
    fig1 = go.Figure()
    fig1.add_trace(go.Scatter(
        x=daily["date_parsed"], y=daily["count"],
        mode='lines+markers', line=dict(color='#334155', width=2),
        marker=dict(size=7, color='#334155'),
        fill='tozeroy', fillcolor='rgba(51, 65, 85, 0.1)'
    ))
    fig1.update_layout(
        height=280, showlegend=False, xaxis_title="Date", yaxis_title="Incidents",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig1, use_container_width=True)
    with st.expander("💡 Clinical Interpretation (Berry Street Body Domain)"):
        st.markdown("**Pattern Recognition:** Look for patterns (e.g., Mondays, after breaks). " +
                   "**Berry Street Body:** Schedule extra regulation supports during high-frequency periods - breathing, movement breaks, sensory activities. " +
                   "Increasing frequency may indicate student's nervous system is dysregulated and needs Body domain strategies.")
    st.markdown("---")
    
    # GRAPH 2: Top Behaviours
    st.markdown("### 🎯 Most Common Behaviours")
    beh_counts = full_df["behaviour_type"].value_counts().head(5)
    fig2 = go.Figure()
    fig2.add_trace(go.Bar(
        y=beh_counts.index, x=beh_counts.values,
        orientation='h', marker=dict(color='#475569'),
        text=beh_counts.values, textposition='outside'
    ))
    fig2.update_layout(
        height=280, showlegend=False, xaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig2, use_container_width=True)
    with st.expander("💡 Clinical Interpretation (Behaviour as Communication)"):
        st.markdown(f"**Primary:** {beh_counts.index[0]} ({beh_counts.values[0]} incidents). " +
                   "**Behaviour Analysis:** Focus intervention planning on top 2-3 behaviours. " +
                   "**Berry Street:** Behaviours are communication - what is student trying to tell us? Strengthen Relationship domain through connection.")
    st.markdown("---")
    
    # GRAPH 3: Top Triggers
    st.markdown("### 🔍 Most Common Triggers (Antecedents)")
    ant_counts = full_df["antecedent"].value_counts().head(5)
    fig3 = go.Figure()
    fig3.add_trace(go.Bar(
        y=ant_counts.index, x=ant_counts.values,
        orientation='h', marker=dict(color='#64748b'),
        text=ant_counts.values, textposition='outside'
    ))
    fig3.update_layout(
        height=280, showlegend=False, xaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig3, use_container_width=True)
    with st.expander("💡 Clinical Interpretation (Proactive Strategies)"):
        st.markdown(f"**Key trigger:** {ant_counts.index[0]}. " +
                   "**Behaviour Analysis:** Plan proactive supports before this occurs - antecedent manipulation is most effective prevention. " +
                   "**Berry Street Stamina:** Build student's capacity to persist through challenging moments.")
    st.markdown("---")
    
    # GRAPH 4: Severity Trend
    st.markdown("### 📊 Severity Over Time")
    fig4 = go.Figure()
    fig4.add_trace(go.Scatter(
        x=full_df["date_parsed"], y=full_df["severity"],
        mode='markers', marker=dict(size=8, color='#334155', opacity=0.6)
    ))
    if len(full_df) >= 2:
        z = np.polyfit(range(len(full_df)), full_df["severity"], 1)
        p = np.poly1d(z)
        fig4.add_trace(go.Scatter(
            x=full_df["date_parsed"], y=p(range(len(full_df))),
            mode='lines', line=dict(color='#94a3b8', width=2, dash='dash'),
            name='Trend'
        ))
    fig4.update_layout(
        height=280, yaxis=dict(range=[0, 6]), xaxis_title="Date", yaxis_title="Severity",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig4, use_container_width=True)
    
    trend_dir = "increasing" if len(full_df) >= 2 and full_df.tail(5)["severity"].mean() > full_df.head(5)["severity"].mean() else "decreasing"
    with st.expander("💡 Clinical Interpretation (Progress Monitoring)"):
        st.markdown(f"Severity appears **{trend_dir}** over time. " +
                   ("**Action Required:** Review strategies - may need stronger Body and Relationship supports. " if trend_dir == "increasing" 
                    else "**Positive Progress:** Current Berry Street strategies showing effect. Continue Body and Relationship focus. ") +
                   "Monitor for plateaus which may indicate need for strategy adjustment or focus on Engagement domain.")
    st.markdown("---")
    
    # GRAPH 5: Location Hotspots
    st.markdown("### 📍 Location Hotspots")
    loc_counts = full_df["location"].value_counts().head(5)
    fig5 = go.Figure()
    fig5.add_trace(go.Bar(
        y=loc_counts.index, x=loc_counts.values,
        orientation='h', marker=dict(color='#64748b'),
        text=loc_counts.values, textposition='outside'
    ))
    fig5.update_layout(
        height=280, showlegend=False, xaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig5, use_container_width=True)
    with st.expander("💡 Clinical Interpretation (Environmental Strategies)"):
        st.markdown(f"Most incidents in **{loc_counts.index[0]}**. " +
                   "Consider: environmental modifications (lighting, noise, space), increased staff support, " +
                   "**Berry Street Body:** sensory-friendly adjustments, calming spaces, visual supports.")
    st.markdown("---")
    
    # GRAPH 6: Time of Day
    st.markdown("### ⏰ Time of Day Patterns")
    session_counts = full_df["session"].value_counts()
    fig6 = go.Figure()
    fig6.add_trace(go.Bar(
        x=session_counts.index, y=session_counts.values,
        marker=dict(color='#475569'),
        text=session_counts.values, textposition='outside'
    ))
    fig6.update_layout(
        height=280, showlegend=False, yaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig6, use_container_width=True)
    with st.expander("💡 Clinical Interpretation (Regulation Timing)"):
        st.markdown(f"Peak time: **{session_counts.index[0]}**. " +
                   "**Berry Street Body:** Provide proactive regulation before this period - breathing exercises, movement breaks, sensory check-ins. " +
                   "Build student's self-regulation capacity through predictable regulation routines.")
    st.markdown("---")

    
    # GRAPH 7-10 and rest of analysis...
    st.markdown("### 📆 Day of Week Patterns")
    day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    day_counts = full_df["day_of_week"].value_counts().reindex(day_order, fill_value=0)
    fig7 = go.Figure()
    fig7.add_trace(go.Bar(
        x=day_counts.index, y=day_counts.values,
        marker=dict(color='#64748b'),
        text=day_counts.values, textposition='outside'
    ))
    fig7.update_layout(
        height=280, showlegend=False, yaxis_title="Frequency",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig7, use_container_width=True)
    
    high_day = day_counts.idxmax()
    with st.expander("💡 Clinical Interpretation (Berry Street Relationship)"):
        st.markdown(f"**{high_day}** has most incidents. Consider connection routines: Monday welcome/check-in, Friday regulation support. " +
                   "**Berry Street Relationship:** Strong connections reduce incidents. Pattern may indicate when student needs extra relational support.")
    st.markdown("---")
    
    # CLINICAL SUMMARY with Berry Street
    st.markdown("### 🧠 Clinical Summary")
    st.caption("Evidence-based interpretation using ABA, Trauma-Informed Practice, Berry Street Education Model, and CPI principles")
    
    top_beh = full_df["behaviour_type"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_ant = full_df["antecedent"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_loc = full_df["location"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_session = full_df["session"].mode()[0] if len(full_df) > 0 else "Unknown"
    
    # Calculate risk score
    recent = full_df.tail(7)
    risk_score = min(100, int(
        (len(recent) / 7 * 10) +
        (recent["severity"].mean() * 8) +
        (len(full_df[full_df["incident_type"] == "Critical"]) / len(full_df) * 50)
    ))
    risk_level = "LOW" if risk_score < 30 else "MODERATE" if risk_score < 60 else "HIGH"
    risk_color = "#10b981" if risk_score < 30 else "#f59e0b" if risk_score < 60 else "#ef4444"
    
    st.info(f"""
    **Key Patterns Identified:**
    - Primary behaviour: **{top_beh}**
    - Main trigger: **{top_ant}**
    - Hotspot location: **{top_loc}**
    - Peak time: **{top_session}**
    - Risk Level: **{risk_level}** ({risk_score}/100)
    
    **Behaviour Analysis Framework:** {student['name']} is most vulnerable when "{top_ant}" occurs in {top_loc} during {top_session}. 
    This behaviour is a safety strategy and communication method. The behaviour serves a function - likely escape/avoidance or attention-seeking based on patterns.
    
    **Trauma-Informed & Berry Street Lens:** Behaviours represent adaptive responses to perceived threat. Student's nervous system is responding to environmental cues. 
    **Berry Street Education Model** emphasizes five domains:
    - **Body:** Self-regulation, wellbeing, sensory needs
    - **Relationship:** Positive connections with adults and peers
    - **Stamina:** Persistence, engagement, coping with challenges
    - **Engagement:** Readiness for learning, curiosity
    - **Character:** Values, agency, identity
    
    **Foundation First:** Focus on Body (regulation) and Relationship (connection) domains before expecting Engagement or Character development.
    
    **CPI Alignment:** Use Supportive Stance, low slow voice, reduce audience, one key adult maintains connection. 
    Behaviour is communication - understand the message before responding.
    """)
    
    st.success(f"""
    **Evidence-Based Recommendations (Berry Street Framework):**
    
    **1. Body Domain (Regulation):** Regulated start before {top_session}, breathing exercises, movement breaks, sensory check-ins in {top_loc}, zones of regulation
    
    **2. Relationship Domain (Connection):** Key adult check-in before "{top_ant}", relationship-building activities, acknowledgment of feelings, co-regulation strategies
    
    **3. Stamina Domain (Persistence):** Link to Personal & Social Capability, teach help-seeking, practice requesting breaks, build coping strategies
    
    **4. SMART Goal:** Over 5 weeks, use help-seeking strategy in 4/5 opportunities with support (supports Body and Relationship). Review {(datetime.now() + timedelta(weeks=5)).strftime('%d/%m/%Y')}.
    """)
    
    st.markdown("---")
    
    # EXPORT with Berry Street branding
    st.markdown("### 📄 Export Data & Reports")
    st.caption("Professional reports prepared by Learning and Behaviour Unit using Berry Street Education Model")
    
    col1, col2 = st.columns(2)
    
    with col1:
        csv = full_df.to_csv(index=False)
        st.download_button(
            "📥 Download Raw Data (CSV)",
            csv,
            file_name=f"{student['name'].replace(' ', '_')}_Incident_Data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    with col2:
        with st.spinner("Generating Behaviour Analysis Plan with graphs..."):
            docx_file = generate_behaviour_analysis_plan_docx(
                student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level
            )
        if docx_file:
            st.download_button(
                "📄 Behaviour Analysis Plan (Word with Graphs)",
                docx_file,
                file_name=f"BAP_{student['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Professional report with 6 embedded graphs, Berry Street framework, and evidence-based recommendations"
            )
        else:
            st.error("Unable to generate BAP. Please ensure kaleido is installed.")
    
    st.markdown("---")
    
    # Bottom Navigation
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅ Back to Students", type="primary", key="back_analysis_bottom", use_container_width=True):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("🏠 Program Landing", key="home_analysis_bottom", use_container_width=True):
            go_to("landing")

def main():
    init_state()
    
    if not st.session_state.logged_in:
        render_login_page()
        return
    
    page = st.session_state.current_page
    
    if page == "landing": render_landing_page()
    elif page == "program_students": render_program_students_page()
    elif page == "incident_log": render_incident_log_page()
    elif page == "critical_incident": render_critical_incident_page()
    elif page == "student_analysis": render_student_analysis_page()
    else: render_landing_page()

if __name__ == "__main__":
    main()
